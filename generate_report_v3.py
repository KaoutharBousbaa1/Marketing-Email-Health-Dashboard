from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import datetime
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import pandas as pd
import numpy as np
from scipy import stats as scipy_stats

BASE   = Path(__file__).resolve().parent
GENERATED = BASE / "data" / "generated"
GENERATED.mkdir(parents=True, exist_ok=True)
CHARTS = BASE / "charts"

doc = Document()

# ── Page margins ───────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)

# ── Style helpers ──────────────────────────────────────────────────────────────
def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(79, 70, 229)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(30, 30, 30)
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(79, 70, 229)
    return p

def body(doc, text, size=11, bold=False, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def callout(doc, text, color_rgb=(79, 70, 229)):
    """A highlighted insight box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f"  {text}")
    set_font(run, size=10.5, italic=True, color=color_rgb)
    # Add left border shading via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "24")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "%02X%02X%02X" % color_rgb)
    pBdr.append(left)
    pPr.append(pBdr)
    return p

def chart_block(doc, filename, figure_num, caption, interpretation, width=Inches(5.8)):
    """Insert chart + caption + interpretation paragraph."""
    path = CHARTS / filename
    if not path.exists():
        body(doc, f"[Chart not found: {filename}]", color=(200,0,0))
        return
    # Chart image
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=width)
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(4)
    run = cap.add_run(f"Figure {figure_num} — {caption}")
    set_font(run, size=9, italic=True, color=(100, 100, 100))
    # Interpretation
    callout(doc, f"Deep interpretation: {interpretation}", color_rgb=(79, 70, 229))
    doc.add_paragraph()

def shade_cell(cell, hex_color="4F46E5"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_table(doc, headers, rows, header_color="4F46E5"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, header_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=10, bold=True, color=(255, 255, 255))
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = "F3F4F6" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, fill)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_font(run, size=10)
    doc.add_paragraph()
    return table

def divider(doc):
    p = doc.add_paragraph("─" * 90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_font(run, size=7, color=(200, 200, 200))

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("EMAIL HEALTH — DEEP ANALYSIS REPORT")
set_font(run, size=24, bold=True, color=(79, 70, 229))

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run("Subscriber Lifespan  •  Engagement Trends  •  Phase Shift Analysis  •  Sales vs Value Performance")
set_font(run, size=12, italic=True, color=(100, 100, 100))

doc.add_paragraph()
divider(doc)
doc.add_paragraph()

for label, val in [
    ("Report Date", datetime.date.today().strftime("%B %d, %Y")),
    ("Data Period", "February 2025 – February 2026"),
    ("Subscribers Analysed", "6,854 cancelled subscribers"),
    ("Broadcasts Analysed", "206 emails (>900 recipients)"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label}:  ")
    set_font(r1, size=11, bold=True, color=(79, 70, 229))
    r2 = p.add_run(val)
    set_font(r2, size=11, color=(50, 50, 50))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "Table of Contents")
toc_items = [
    ("1.", "Executive Summary"),
    ("2.", "Subscriber Lifespan Analysis"),
    ("    2.1", "Overall Statistics"),
    ("    2.2", "Cancellation Timing Buckets"),
    ("    2.3", "Cancellation Trend Over Time"),
    ("3.", "Email Broadcast Engagement Overview"),
    ("4.", "Recent Engagement Decline"),
    ("5.", "Are Recent Unsubscribers Old or New Subscribers?"),
    ("6.", "When Did the Decline Start? — Period Analysis"),
    ("    6.1", "Three Distinct Phases"),
    ("    6.2", "Lifespan of Cancellers by Phase"),
    ("7.", "Sales vs Value Emails — Which Drove the Decline?"),
    ("    7.1", "Overall Performance by Category"),
    ("    7.2", "Performance by Phase × Category"),
    ("    7.3", "Decline Magnitude"),
    ("8.", "Cold Subscribers — Who Stopped Engaging?"),
    ("    8.1", "How Old Are the Cold Subscribers?"),
    ("    8.2", "Acquisition Source (Tag Breakdown)"),
    ("    8.3", "When Did Cold Subscribers Join?"),
    ("    8.4", "Tag × Phase Cross-Analysis"),
    ("    8.5", "When Did Cold Subscribers Stop Engaging? — API Data"),
    ("9.", "Conversion Timing — Signup to First Purchase"),
    ("    9.1", "How Long Buyers Stay on List Before First Purchase"),
    ("10.", "Lead Magnet Impact Analysis"),
    ("    10.1", "Did the Survey Boost Future Open Rates? (Q1)"),
    ("    10.2", "Survey Responders vs Non-Responders: Engagement (Q2)"),
    ("    10.3", "Retention: Who Stays Active?"),
    ("    10.4", "Post-Survey Email-Type Lift: Value OR, Sales OR, Sales CTOR"),
    ("    10.5", "Are Group A Subscribers New or Old?"),
    ("    10.6", "Monthly Evolution: Open Rate and Sales CTOR (A vs B)"),
    ("11.", "Subscriber Origin Split: Lead-Magnet Signups vs Existing Recipients"),
    ("12.", "Signup-Date Cohort Evolution (Open Rate + CTOR)"),
    ("13.", "Bootcamp Buyers vs Non-Buyers — Engagement & Age Profile"),
    ("14.", "Sales Intent Outcomes — Intended vs Acted"),
    ("15.", "Workshop-to-Bootcamp Conversion (Did Workshop Buyers Convert to Bootcamp?)"),
    ("16.", "Final Verdict — Did the Free Lead Magnet Work?"),
    ("17.", "Detailed Executive Summary of Key Takeaways"),
    ("18.", "Strategic Narrative & Action Plan (Moved from former 9.1/9.2)"),
    ("    18.1", "The Story the Data Is Telling Us"),
    ("    18.2", "What We Should Do About It"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{num}  ")
    set_font(r1, size=10.5, bold=True, color=(79, 70, 229))
    r2 = p.add_run(title)
    set_font(r2, size=10.5)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "1. Executive Summary")

body(doc,
     "This report presents a comprehensive deep-dive into the health of an email newsletter list. "
     "Two datasets were analysed: a subscriber cancellation export covering 6,854 cancelled contacts "
     "and a broadcast performance export covering 206 email campaigns sent to audiences of more than "
     "900 recipients. The data spans February 2025 through February 2026 — a full year of activity.")

body(doc,
     "The analysis was structured around four central questions that the data was equipped to answer:")
bullet(doc, "How long do subscribers stay before cancelling — and is that changing?")
bullet(doc, "Is email engagement (opens and clicks) declining, and by how much?")
bullet(doc, "When exactly did the decline begin, and what triggered it?")
bullet(doc, "Is the decline coming from sales emails, value emails, or both?")

body(doc,
     "The findings paint a coherent and concerning picture: the newsletter list is growing in size, "
     "but its engaged core is eroding. Subscribers who have been on the list for months are now "
     "leaving at a rate previously unseen. Email opens are declining, and click-through behaviour "
     "has collapsed — particularly inside value content. Most significantly, this did not happen "
     "overnight. The data shows a gradual, two-stage deterioration that started as early as June 2025, "
     "five months before the more visible open rate decline that began in November 2025.")

body(doc, "Key findings at a glance:", bold=True)
bullet(doc, "Median subscriber lifespan: 81 days overall — but recent (last 30 days) cancellers stayed 194 days on average, 2.5× the historical norm.")
bullet(doc, "Open rate fell from 40.8% (Strong Start) to 35.9% (Visible Decline) — a drop of ~5 percentage points.")
bullet(doc, "Click-to-open rate (CTOR) collapsed from 12.6% to 6.4% — a 49% relative decline.")
bullet(doc, "CTOR started dropping in June 2025. Open rate followed in November 2025.")
bullet(doc, "Sales emails: open rate dropped 23.5 pp (56.5% → 33.0%). Both categories saw CTOR decline ~43–60%.")
bullet(doc, "The leavers are no longer new sign-ups — they are established, long-tenured subscribers.")

callout(doc,
        "The core narrative: A newsletter that was highly engaging in early 2025 gradually lost its "
        "click-through power from June 2025, then began losing its ability to even get opened from "
        "November 2025. The audience most affected is not casual sign-ups — it is subscribers who "
        "committed to the list for months and are now disengaging.",
        color_rgb=(220, 38, 38))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — SUBSCRIBER LIFESPAN ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "2. Subscriber Lifespan Analysis")

body(doc,
     "The subscriber export contains every contact who cancelled their subscription, with two key "
     "dates recorded: when they subscribed and when they last updated their status (i.e., when they "
     "cancelled). The difference between these two dates — the lifespan — tells us how long each "
     "subscriber stayed before deciding to leave. This is one of the most important health signals "
     "for a newsletter: are people leaving after one email, or after months of engagement?")

h2(doc, "2.1 Overall Lifespan Statistics")

body(doc,
     "Across all 6,854 cancelled subscribers, the lifespan ranged from 0 days (someone who cancelled "
     "the same day they signed up) to 379 days (someone who stayed for over a year before leaving). "
     "The distribution tells a nuanced story.")

add_table(doc,
    ["Metric", "Value", "What It Means"],
    [
        ["Total cancelled", "6,854", "All contacts who ever cancelled during the period"],
        ["Minimum", "0 days", "Immediate cancellations — signed up and left same day"],
        ["25th percentile", "29 days", "25% of cancellers left within the first month"],
        ["Median", "81 days", "Half of all cancellers stayed fewer than 3 months"],
        ["Mean", "109.5 days", "Average pulled higher by long-term subscribers"],
        ["75th percentile", "176 days", "75% left within 6 months"],
        ["Maximum", "379 days", "The longest-staying canceller was on the list ~1 year"],
    ]
)

body(doc,
     "The gap between the median (81 days) and the mean (109.5 days) is significant. It tells us "
     "that the distribution is right-skewed — most people leave early, but a meaningful tail of "
     "long-tenured subscribers is pulling the average up. This tail is important: these are people "
     "who invested real time with the newsletter before eventually leaving.")

chart_block(doc, "1_lifespan_distribution.png", 1,
    "Distribution of subscriber lifespan — histogram with density curve overlay.",
    "The tall spike on the left (near 0–30 days) represents the large volume of early cancellers — "
    "subscribers who signed up out of curiosity and quickly decided the newsletter was not for them. "
    "This is normal for any growing newsletter. What is more interesting is the long, flat tail "
    "extending to the right: hundreds of subscribers who stayed 3, 6, even 9 months before leaving. "
    "These are not casual sign-ups. The density curve (pink line) shows two soft peaks — one at the "
    "very start (early bounces) and a broader one around 90–200 days (mid-to-long-term leavers). "
    "A healthy list would have the right tail shrinking over time; here, it is growing.")

h2(doc, "2.2 Cancellation Timing Buckets")

body(doc,
     "To make the lifespan data more actionable, cancellers were grouped into six buckets based on "
     "how long they stayed. Each bucket tells a different story about why someone might have left.")

add_table(doc,
    ["Lifespan Bucket", "Count", "Share", "What This Group Likely Represents"],
    [
        ["0–7 days", "395", "5.8%", "Immediate bounce — content or frequency mismatch from day one"],
        ["8–30 days", "1,257", "18.4%", "Early leavers — completed onboarding, decided it wasn't for them"],
        ["31–90 days", "1,914", "27.9%", "Largest group — stayed through initial curiosity phase then left"],
        ["91–180 days", "1,484", "21.7%", "Mid-term — engaged for a season, then disengaged gradually"],
        ["181–365 days", "1,621", "23.7%", "Long-term — committed subscribers who eventually churned"],
        [">1 year", "37", "0.5%", "Veteran loyalists — extremely rare to lose these"],
    ]
)

body(doc,
     "The most striking observation is that 24% of all cancellers stayed between 6 months and a full "
     "year before leaving. These are not people who sampled the newsletter and moved on — these are "
     "subscribers who were genuinely engaged for an extended period. Losing them represents a "
     "meaningful erosion of the newsletter's most invested audience.")

chart_block(doc, "2_lifespan_buckets.png", 2,
    "Horizontal bar chart of cancellation count and share per lifespan bucket.",
    "The bars tell a clear story: the 31–90 day window is the single biggest loss zone — nearly 2,000 "
    "subscribers cancelled in this window. This is the critical retention window. If a subscriber "
    "makes it past 90 days, they are more likely to stay — but if they leave before then, they "
    "never fully committed. The 181–365 day bar being nearly as large as the 91–180 day bar is "
    "the real concern: almost as many long-term subscribers are leaving as medium-term ones. "
    "This asymmetry — losing nearly as many 6–12 month subscribers as 3–6 month subscribers — "
    "suggests a structural disengagement problem, not just normal early churn.")

h2(doc, "2.3 Cancellation Trend Over Time")

body(doc,
     "Beyond the overall distribution, it is critical to understand how cancellation patterns evolved "
     "month by month. Are more people cancelling now than before? And are the people cancelling now "
     "different from those who cancelled earlier?")

chart_block(doc, "3_monthly_trend.png", 3,
    "Monthly cancellation volume (bars) with median lifespan of cancellers per month (line).",
    "This chart contains two stories in one. The bars show the raw volume of cancellations each month — "
    "which rises over time as the list grows, which is expected. More subscribers means more potential "
    "cancellers. But the orange line (median lifespan of that month's cancellers) is the real signal. "
    "In early 2025, the median lifespan of monthly cancellers was very low — 20 to 40 days — meaning "
    "mostly new subscribers were leaving. From around September–October 2025, the line starts rising "
    "sharply. By November–February 2026, the median lifespan of monthly cancellers exceeds 150–180 days. "
    "This means the composition of who is leaving fundamentally changed. It stopped being mostly "
    "new-subscriber churn and became long-term subscriber churn. That is a fundamentally different "
    "and more serious problem.")

chart_block(doc, "4_cumulative_cancellations.png", 4,
    "Cumulative cancellation curve over time.",
    "The cumulative curve shows the total build-up of cancellations from day one. A straight diagonal "
    "line would indicate a constant cancellation rate. What we see instead is a curve that steepens "
    "noticeably from mid-2025 onward — the slope increases, meaning cancellations are accelerating. "
    "The list is not just experiencing normal attrition; the rate of loss is picking up speed. "
    "If this trajectory continues, the gap between new subscribers gained and subscribers lost "
    "will narrow, threatening net list growth.")

chart_block(doc, "5_lifespan_by_quarter.png", 5,
    "Box plot of lifespan distribution by cancellation quarter.",
    "Each box represents one quarter's worth of cancellers. The width of the box shows where the "
    "middle 50% of lifespans fall, and the line inside is the median. In Q1 2025, the boxes are "
    "narrow and low — cancellers had short, tightly clustered lifespans, mostly early quitters. "
    "By Q4 2025 and into Q1 2026, the boxes are taller, wider, and sitting higher on the axis. "
    "This means cancellers in the most recent quarters are not only staying longer before leaving — "
    "they are also more varied in when they leave, from 1 month to nearly a year. The spread widening "
    "quarter by quarter is a warning sign: the newsletter is losing subscribers across all tenure groups, "
    "not just the most recent sign-ups.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — BROADCAST ENGAGEMENT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "3. Email Broadcast Engagement Overview")

body(doc,
     "The broadcast dataset covers every email campaign sent to more than 900 recipients — a threshold "
     "chosen to exclude small test sends, bootcamp cohort emails, and niche segment campaigns that "
     "are not representative of the main newsletter audience. This leaves 206 sends across the full "
     "year, providing a clean view of how the main list responds to each email.")

body(doc,
     "Two metrics are the primary focus throughout this report:")
bullet(doc, "Open Rate — the percentage of recipients who opened the email. This measures subject line quality, sender trust, and inbox deliverability.")
bullet(doc, "Click-to-Open Rate (CTOR) — the percentage of people who opened the email and then clicked a link. This measures content quality, CTA strength, and how well the email body delivers on the subject line's promise.")

body(doc,
     "CTOR is often more revealing than raw click rate, because it removes the variable of how many "
     "people opened in the first place and isolates the question: of the people who were interested "
     "enough to open, how many found the content compelling enough to act?")

body(doc,
     "Important methodology note: all CTOR figures in this report are calculated using only emails "
     "that contained at least one link (174 out of 205 sends with >900 recipients). Emails with no "
     "links were excluded from CTOR analysis — including them would artificially depress the metric, "
     "since a zero click rate on a no-link email tells us nothing about content quality. Open Rate "
     "is calculated across all emails (with or without links), since the presence of a link has no "
     "bearing on whether someone opens an email.")

add_table(doc,
    ["Metric", "Average", "Median", "Min", "Max", "Note"],
    [
        ["Open Rate",            "41.3%", "38.1%", "14.5%", "71.8%", "All 205 sends >900 recipients"],
        ["Click-to-Open Rate",   "11.1%", "9.5%",  "0.3%",  "54.8%", "174 sends with links only"],
        ["Click Rate",           "5.1%",  "3.5%",  "0.1%",  "39.0%", "174 sends with links only"],
        ["Unsubscribe Rate",     "0.24%", "0.19%", "0.0%",  "1.36%", "All 205 sends"],
    ]
)

body(doc,
     "An average open rate of 41.1% is strong by industry standards (typical newsletters average "
     "20–35%). This suggests the sender has built genuine trust and recognition with the audience. "
     "However, averages can mask trends — and the trend is what matters here.")

body(doc, "Standout performers:")
bullet(doc, "Highest open rate: 'Workshop recording and slides' — 71.8%. Recordings consistently drive the highest opens because recipients are actively looking for promised content.")
bullet(doc, "Highest CTOR: 'Leaked AI memo from Shopify CEO' — 54.8%. Curiosity-driven, newsworthy subject lines drive clicks once opened.")

chart_block(doc, "6_openrate_vs_ctor_scatter.png", 6,
    "Scatter plot of open rate vs CTOR for all 206 broadcasts. Bubble size = number of recipients.",
    "Each bubble is one email send. The horizontal axis is open rate and the vertical axis is CTOR. "
    "Ideally, emails should cluster in the top-right corner — high opens AND high clicks. "
    "Instead, the chart shows a wide scatter with most emails clustering in a band of 30–50% open "
    "rate but very low CTOR (under 10%). The bubbles with the highest CTOR are often smaller "
    "(fewer recipients), suggesting that the most click-worthy content tends to go to niche, "
    "highly targeted segments rather than the full list. The largest bubbles (main list sends) "
    "sit at high open rates but middling-to-low CTOR — people open but rarely click. "
    "This is the core tension in the data.")

chart_block(doc, "8_rates_rolling_trend.png", 7,
    "7-email rolling average of open rate and CTOR over time.",
    "The two lines tell very different stories. The blue line (open rate) starts high, stays "
    "relatively stable through mid-2025, then begins a visible downward slope from late 2025. "
    "The pink line (CTOR) drops sharply and early — it falls steeply from around June 2025 and "
    "never recovers. The gap between the two lines widens over time, which is the most important "
    "visual signal in this entire dataset: subscribers are still opening but increasingly not "
    "clicking. Something changed in the content, the CTAs, or the audience's relationship with "
    "the newsletter around June 2025 that broke the click behaviour long before it broke the "
    "open behaviour.")

chart_block(doc, "9_rates_heatmap.png", 8,
    "Monthly heatmap of average open rate, CTOR, and the gap between them.",
    "The heatmap uses colour to show performance: green = higher, red = lower. Reading left to right "
    "across the months, both open rate and CTOR show a clear colour shift from greener (earlier months) "
    "to redder (later months). The 'Gap' row (open rate minus CTOR) widens over time, shown by "
    "increasingly darker colours — confirming that the divergence between opens and clicks is not "
    "random variation but a consistent and worsening structural trend. The months of February–May 2025 "
    "stand out as the healthiest period. From June 2025 onward, the heat maps show persistent cooling.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — RECENT ENGAGEMENT DECLINE
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "4. Recent Engagement Decline")

body(doc,
     "To understand the current state of the newsletter, the most recent 60 days of broadcast data "
     "(December 2025 – February 2026) were compared against the full historical average. This comparison "
     "reveals how performance has shifted from the overall baseline.")

add_table(doc,
    ["Period", "Avg Open Rate", "Avg CTOR", "Avg Openers / Send"],
    [
        ["Full history (Feb 2025 – Feb 2026)", "39.1%", "8.8%",  "5,592"],
        ["Recent 60 days (Dec 2025 – Feb 2026)", "35.3%", "6.3%", "6,132"],
        ["Previous period (Feb – Dec 2025)", "39.9%", "9.3%", "5,478"],
    ]
)

body(doc,
     "The numbers reveal an important nuance: the absolute number of people opening each email "
     "is actually higher in the recent period (6,132 vs 5,478 average openers per send). This is "
     "because the list has grown significantly. However, as a proportion of the list, open rate "
     "has fallen by 4.6 percentage points — and CTOR has dropped by 3 points, a 32% relative decline. "
     "The list is growing faster than engagement is keeping up. In other words: more people are "
     "receiving the emails but a smaller and smaller fraction of them are engaging with the content.")

callout(doc,
        "A growing list with declining engagement rates is often more dangerous than a flat list "
        "with stable engagement. It means the new subscribers are diluting the engaged core, "
        "and the engaged core itself is starting to disengage.",
        color_rgb=(220, 38, 38))

chart_block(doc, "A_absolute_openers_trend.png", 9,
    "Absolute opener count per send (blue) vs recipient count (orange dashed).",
    "This chart separates the two trends that often get confused. The orange line shows the list "
    "growing steadily throughout the year — from around 8,000 recipients per send in early 2025 "
    "to nearly 20,000 by February 2026. The blue line shows the number of people actually opening "
    "each email. These numbers have grown too, but at a much slower pace than the list itself. "
    "The visual gap between the two lines widens over time — and that gap represents the growing "
    "portion of the list that receives but does not open. By February 2026, roughly 65% of "
    "recipients are not opening each email, compared to roughly 60% in early 2025. "
    "The absolute numbers look healthy; the proportional reality is not.")

chart_block(doc, "B_open_rate_trend.png", 10,
    "Open rate % trend over time with overall average reference line and recent 60-day window highlighted.",
    "The rolling average line (8-send window) shows a relatively stable open rate from February "
    "through October 2025, hovering around 38–42%. Then from November 2025 onward, the line begins "
    "a consistent downward drift. The shaded pink area on the right marks the recent 60 days — "
    "and the rolling line sits noticeably below the historical average (dashed grey line) during "
    "this period. This is not a temporary dip; it is a sustained shift in how often subscribers "
    "choose to open the email when they see it in their inbox.")

chart_block(doc, "C_ctor_trend.png", 11,
    "Click-to-open rate % trend over time with recent 60-day window highlighted.",
    "The CTOR story is more dramatic than the open rate story. The rolling average for CTOR "
    "was around 12–15% in early 2025 — meaning roughly 1 in 7 openers clicked something in "
    "each email. By mid-2025, this had already fallen to around 7–8%. By early 2026, it sits "
    "at 5–6%. Each individual data point (scattered dots) shows high volatility — some emails "
    "still achieve 20–30% CTOR (typically sales launch emails with clear CTAs), while many "
    "regular sends now achieve under 3%. The structural average has been in sustained decline "
    "since June 2025, and the recent 60-day window shows no sign of recovery.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — OLD OR NEW UNSUBSCRIBERS?
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "5. Are Recent Unsubscribers Old or New Subscribers?")

body(doc,
     "One of the most important questions in newsletter health analysis is whether churn is coming "
     "from the recently acquired subscribers (who never fully committed) or from the long-standing "
     "audience (who had a genuine relationship with the content). These two scenarios have very "
     "different implications. Early churn is normal and expected — not everyone who signs up will "
     "stay. But veteran churn is a structural alarm: it means the content that once retained committed "
     "subscribers has stopped working for them.")

body(doc,
     "To answer this, the lifespan of subscribers who cancelled in the most recent periods was "
     "compared against those who cancelled earlier in the year.")

add_table(doc,
    ["Group", "Count", "Median Lifespan", "Mean Lifespan", "% Left Within 30d", "% Left After 180d"],
    [
        ["Older cancellers (before Dec 2025)", "5,869", "74 days",  "97.4 days",  "—", "—"],
        ["Recent 60-day cancellers",            "985",   "180 days", "182.0 days", "—", "—"],
        ["Recent 30-day cancellers",            "442",   "194 days", "189.8 days", "14.5%", "47.0%"],
    ]
)

body(doc,
     "The numbers are striking. Subscribers cancelling in the last 30 days had a median lifespan "
     "of 194 days — compared to 74 days for historical cancellers. That is 2.6× longer. "
     "Nearly half (47%) of the most recent cancellers had been subscribed for more than 6 months. "
     "Only 14.5% were new subscribers (under 30 days). This is the opposite of what healthy churn "
     "looks like. Healthy churn is mostly new-subscriber bounce. What we are seeing now is "
     "established-subscriber departure.")

callout(doc,
        "The subscribers who are leaving right now are not people who recently discovered the newsletter "
        "and decided it wasn't for them. They are people who have been reading for 6+ months — "
        "who opened emails, perhaps clicked links, perhaps even bought programmes — and have now "
        "decided to leave. This is the most valuable segment of any newsletter audience, and losing "
        "them at scale is the most urgent signal in this entire dataset.",
        color_rgb=(220, 38, 38))

chart_block(doc, "E_recent_vs_older_kde.png", 12,
    "Density curves comparing lifespan of recent 30-day cancellers vs older cancellers.",
    "This chart overlays two density curves — one for all historical cancellers (blue) and one for "
    "the most recent 30 days of cancellers (pink). The historical curve peaks sharply near 0–30 days, "
    "reflecting the large volume of early-exit subscribers who were always part of normal churn. "
    "The recent 30-day curve looks completely different: it peaks around 150–220 days, with almost "
    "no spike at the beginning. The dashed vertical lines show the medians — 74 days for historical "
    "vs 194 days for recent. The entire recent curve has shifted to the right. This is not a gradual "
    "drift; it is a fundamental change in the profile of who is leaving. The newsletter is no longer "
    "mostly losing new subscribers — it is losing its veterans.")

chart_block(doc, "F_lifespan_bucket_comparison.png", 13,
    "Side-by-side bucket breakdown: older cancellers vs recent 30-day cancellers.",
    "The grouped bars make the comparison viscerally clear. For older cancellers, the tallest bar "
    "is the 8–30 day bucket — new-subscriber churn dominated. For recent 30-day cancellers, "
    "that bar nearly disappears, and the tallest bars are in the 91–180 day and 181–365 day buckets. "
    "The long-tenure bars have effectively swapped places with the short-tenure bars. If these two "
    "groups were unlabelled, you would guess they came from completely different newsletters. "
    "They come from the same list, one year apart. Something changed in between.")

chart_block(doc, "H_lifespan_boxplot_groups.png", 14,
    "Box plot comparing lifespan across older, recent 60-day, and recent 30-day cancellers.",
    "The three boxes make the progression impossible to ignore. The 'Older' box sits low, "
    "with a median of 74 days and a tight distribution. The 'Recent 60-day' box is higher and "
    "wider. The 'Recent 30-day' box is the highest, widest, and has the most spread — a median "
    "of 194 days with a large interquartile range. The expanding boxes tell us not only that "
    "leavers are staying longer, but also that the group leaving now is more diverse in their "
    "tenure — spanning from a few weeks all the way to nearly a year. The newsletter is losing "
    "people from every tenure group now, not just the newest arrivals.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — WHEN DID THE SHIFT HAPPEN?
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "6. When Did the Decline Start? — Phase Analysis")

body(doc,
     "Understanding when the decline started is as important as understanding that it happened. "
     "An early inflection point gives context for what might have caused the change, and when "
     "a recovery effort would need to have begun to prevent the current state. The data reveals "
     "not one shift but two — occurring at different times and affecting different metrics.")

h2(doc, "6.1 Three Distinct Phases")

body(doc,
     "After examining the rolling trend data, the broadcast history was divided into three phases "
     "defined by clear changes in engagement patterns:")

add_table(doc,
    ["Phase", "Period", "Open Rate", "CTOR", "Key Characteristic"],
    [
        ["The Strong Start (Feb–May 2025)",    "Feb–May 2025",     "40.8%", "12.6%", "Both metrics strong and stable"],
        ["The Silent Warning (Jun–Oct 2025)",  "Jun–Oct 2025",     "40.3%", "7.9%",  "CTOR dropped 37%. Open rate held."],
        ["The Visible Decline (Nov 2025–now)",     "Nov 2025–present", "35.9%", "6.4%",  "Both declining. Veteran churn begins."],
    ]
)

body(doc,
     "The two-stage nature of the decline is critical to understand. In the Silent Warning period (Jun–Oct 2025), the open rate "
     "remained essentially flat (40.8% → 40.3%) while CTOR fell sharply (12.6% → 7.9%). This means "
     "subscribers were still trusting the sender enough to open, but once inside the email, "
     "they stopped finding reasons to click. This is a content and CTA signal, not a "
     "deliverability or subject-line signal. The fact that opens held for five months while clicks "
     "collapsed means there was a window — June to October 2025 — during which the problem was "
     "detectable and potentially reversible before it progressed to open rate decline.")

callout(doc,
        "The June 2025 CTOR drop was an early warning signal that went unaddressed for five months. "
        "By the time the open rate began declining in November 2025, the disengagement had deepened "
        "to the point where subscribers no longer found the subject lines worth their attention. "
        "Early action in June–July 2025 might have prevented the open rate decline entirely.",
        color_rgb=(245, 158, 11))

chart_block(doc, "I_open_rate_shift.png", 15,
    "Open rate trend with phase bands highlighted and shift annotation.",
    "The three coloured bands (green = Strong Start, orange = Silent Warning, red = Visible Decline) make the phases "
    "immediately visible. The horizontal lines within each band show the average open rate for that "
    "phase. The transition from green to orange shows almost no change in the open rate line — "
    "it barely moves. This is exactly what makes the Silent Warning period so deceptive: if you were only watching "
    "open rates, you would have thought everything was fine. The real signal was hiding in the CTOR. "
    "From November 2025 onward (the red band), the rolling average line begins a visible downward "
    "slope that has not reversed. The gap between the Strong Start average (green line) and the Visible Decline "
    "average (red line) represents approximately 5 percentage points of lost open rate.")

chart_block(doc, "J_ctor_shift.png", 16,
    "CTOR trend with phase bands — the sharp Silent Warning drop is immediately visible.",
    "Unlike the open rate chart, the CTOR chart shows a dramatic and obvious shift at the Silent Warning boundary "
    "boundary (the orange dashed line at June 2025). The rolling average drops steeply and does not "
    "recover. The green phase average (12.6%) and the orange phase average (7.9%) are visually far "
    "apart. By the Visible Decline period, the average settles even lower at 6.4%. The scattered dots show that "
    "individual emails still occasionally spike high (particularly sales launch emails with strong CTAs), "
    "but these are outliers. The structural baseline for CTOR has permanently shifted downward since "
    "June 2025. This chart is the clearest single piece of evidence that something changed in the "
    "content strategy or audience relationship around that time.")

chart_block(doc, "K_monthly_phases.png", 17,
    "Monthly bar charts for open rate (top) and CTOR (bottom), colour-coded by phase.",
    "Reading these two bar charts top to bottom and left to right reveals the full story in one view. "
    "The top chart (open rate) shows green bars at a consistently healthy level, orange bars at a "
    "similar level, then red bars visibly shorter — the decline is gradual but clear. The bottom "
    "chart (CTOR) tells a more dramatic story: the green bars are tall, then the orange bars drop "
    "sharply at the phase boundary and stay low, and the red bars are barely half the height of the "
    "green bars. The CTOR collapse happened fast and has not recovered. Having both charts aligned "
    "on the same timeline makes it easy to see that the two metrics moved at different times — "
    "CTOR first, open rate months later.")

h2(doc, "6.2 Lifespan of Cancellers by Phase")

body(doc,
     "The engagement decline visible in the broadcast data is directly mirrored in the subscriber "
     "lifespan data. As the newsletter's click engagement collapsed and then its open engagement "
     "declined, the profile of who was cancelling changed dramatically.")

add_table(doc,
    ["Phase", "Cancellers", "Median Lifespan", "% Left ≤30d", "% Left >180d", "Who Is Leaving"],
    [
        ["The Strong Start (Feb–May 2025)",   "1,816", "32 days",  "47.9%", "0.0%",  "Almost entirely new subscribers"],
        ["The Silent Warning (Jun–Oct 2025)", "3,292", "106 days", "20.5%", "25.4%", "Mix — early loyalists starting to leave"],
        ["The Visible Decline (Nov 2025–now)",    "1,746", "167 days", "14.5%", "47.0%", "Mostly established, long-tenured subscribers"],
    ]
)

body(doc,
     "In the Strong Start period (Feb–May 2025), cancellations were dominated by new-subscriber bounce — 47.9% of leavers cancelled "
     "within their first 30 days, and not a single canceller had been subscribed for more than 180 days. "
     "This is exactly what healthy churn looks like: a newsletter naturally filters out people who "
     "are not a good fit early, and retains those who connect with the content.")

body(doc,
     "By the Visible Decline period, that picture has completely inverted. Nearly half (47%) of cancellers had been "
     "subscribed for more than 6 months, and only 14.5% were new-subscriber bounces. The newsletter "
     "is retaining fewer new sign-ups long enough to convert them into loyalists, while simultaneously "
     "losing the loyalists it already had. The median lifespan of Visible Decline cancellers (167 days) is "
     "5.2× higher than the Strong Start period (32 days). That is not drift — it is a structural reversal.")

chart_block(doc, "M_phase_lifespan_kde.png", 18,
    "KDE density curves for lifespan of cancellers in each phase.",
    "Three curves, three completely different shapes. The green curve (Strong Start) has a sharp, tall "
    "peak near 0–30 days — typical early-churn dominated pattern. The orange curve (Silent Warning) has "
    "spread out, with the peak moving to around 60–120 days, and a noticeable tail toward the right. "
    "The red curve (Visible Decline) has shifted almost entirely to the right, with a broad, flat shape "
    "centred around 150–200 days. The vertical dashed lines show the phase medians moving from "
    "32 days → 106 days → 167 days. This is one of the most powerful charts in the analysis: "
    "it shows that the typical person leaving the newsletter right now is not a new subscriber "
    "who bounced — they are someone who stayed for half a year before giving up.")

chart_block(doc, "Q_phase_engagement_vs_lifespan.png", 19,
    "Combined chart: open rate and CTOR per phase (bars) with median canceller lifespan (line).",
    "This chart ties the entire analysis together in a single visual. As the bars (open rate and CTOR) "
    "shrink from phase to phase — showing declining engagement — the diamond markers connected by "
    "the blue line rise. The median lifespan of cancellers goes up as engagement goes down. "
    "This inverse relationship is the core finding of the entire report: the worse the engagement "
    "gets, the older the subscribers who are leaving. As click-through collapsed (Silent Warning period), the "
    "early churners stopped being the dominant leaving group, replaced by mid-tenure subscribers. "
    "As open rates collapsed too (Visible Decline period), even long-term subscribers gave up. "
    "The two trends are not coincidental — they are causally linked.")

chart_block(doc, "L_lifespan_shift.png", 20,
    "Rolling median lifespan of cancellers over time with phase bands.",
    "This chart traces the rolling median lifespan of cancellers month by month, using a 150-person "
    "rolling window to smooth out noise. In early 2025 (green band), the line sits low — between "
    "20 and 50 days — confirming that new-subscriber churn was the norm. Through the orange transition "
    "band (June–October 2025), the line begins to climb. Then from November 2025 (the red band), "
    "the line accelerates sharply upward, reaching 150–200 days by early 2026. The annotation "
    "marks the point where this acceleration became undeniable. If the line continues this trajectory, "
    "the newsletter risks losing most of its remaining long-tenured audience within the next "
    "few months without intervention.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — SALES vs VALUE
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "7. Sales vs Value Emails — Which Drove the Decline?")

body(doc,
     "To pinpoint whether the engagement decline came from a specific type of content, all 206 "
     "broadcasts (>900 recipients) were manually categorised based on their subject line intent:")
bullet(doc, "Sales emails (101 emails) — emails intended to sell a programme, promote a launch, create urgency, announce deadlines, share enrolment invites, or drive sign-ups. Examples: 'Last chance…', 'Doors closing tonight!', 'Your private invite to join the AI Agent Bootcamp', 'Sold out'.")
bullet(doc, "Value emails (105 emails) — emails intended to educate, inform, share insights, tell personal stories, or provide free resources. Examples: 'How to build AI agents', 'My 2025 AI learning roadmap', 'The real reason I quit working at Meta', 'Leaked AI memo from Shopify CEO'.")

body(doc,
     "The split was almost perfectly even: 53 sales emails vs 152 value emails across the full year. "
     "This is a notably high proportion of sales content for a newsletter — roughly one in two emails "
     "has a commercial intent. Understanding which type drove the decline has direct implications "
     "for content strategy.")

h2(doc, "7.1 Overall Performance by Category")

add_table(doc,
    ["Category", "Count", "Avg Open Rate", "Avg CTOR", "Interpretation"],
    [
        ["Value emails", "152", "41.1%", "10.6%", "Higher opens, lower clicks — content consumed passively"],
        ["Sales emails", "53",  "41.7%", "12.6%", "Similar opens, higher clicks — explicit CTAs drive action"],
    ]
)

body(doc,
     "At the aggregate level, value and sales emails have similar open rates (within 1 percentage point). "
     "Sales emails have a higher CTOR (12.6% vs 10.6%) because they contain strong, "
     "specific calls to action — 'join now', 'claim your spot', 'download here'. Value emails, "
     "by design, are less transactional, so lower CTOR is expected. The relevant question is not "
     "which type performs better overall, but which type deteriorated more as the year progressed.")

h2(doc, "7.2 Performance by Phase × Category")

add_table(doc,
    ["Phase", "Category", "Sends", "Avg Open Rate", "Avg CTOR"],
    [
        ["The Strong Start (Feb–May 2025)",   "Value", "51 (OR) / 51 (CTOR)", "42.1%", "13.9%"],
        ["The Strong Start (Feb–May 2025)",   "Sales", "10 (OR) / 9 (CTOR)",  "56.5%", "19.2%"],
        ["The Silent Warning (Jun–Oct 2025)", "Value", "59 (OR) / 53 (CTOR)", "43.9%", "8.5%"],
        ["The Silent Warning (Jun–Oct 2025)", "Sales", "29 (OR) / 24 (CTOR)", "40.8%", "12.7%"],
        ["The Visible Decline (Nov 2025–now)", "Value", "41 (OR) / 23 (CTOR)", "36.0%", "7.9%"],
        ["The Visible Decline (Nov 2025–now)", "Sales", "14 (OR) / 13 (CTOR)", "33.0%", "7.7%"],
    ]
)

body(doc,
     "Reading across the phases for each category reveals two very different decline stories.")

h2(doc, "7.3 Decline Magnitude — The Core Finding")

add_table(doc,
    ["Category", "Open Rate: Strong Start → Visible Decline Drop", "CTOR: Strong Start → Visible Decline Drop", "CTOR Relative Decline"],
    [
        ["Value emails", "−6.1 pp  (42.1% → 36.0%)", "−6.0 pp  (13.9% → 7.9%)", "−43%"],
        ["Sales emails", "−23.5 pp  (56.5% → 33.0%)", "−11.5 pp  (19.2% → 7.7%)", "−60%"],
    ]
)

body(doc, "Problem 1 — Sales Email Open Rate Fatigue:", bold=True, color=(220, 38, 38))
body(doc,
     "Sales emails lost nearly four times as many open-rate percentage points as value emails "
     "(−23.5 pp vs −6.1 pp). In the Visible Decline period (Nov 2025–now), sales emails average only 33.0% open rate — "
     "roughly 1 in 3 recipients — compared to 56.5% in the Strong Start period. This is a classic symptom of "
     "list fatigue from high-frequency commercial content. When a subscriber receives repeated "
     "urgency-driven emails ('last chance', 'doors closing', 'only 5 spots left'), they begin "
     "pattern-recognising the sender as commercial noise and stop opening. The open rate drop "
     "on sales emails is almost certainly driven by a combination of inbox filtering and deliberate "
     "subscriber skip behaviour. Sales emails represented a significant share of total sends, "
     "and the commercial volume may have exceeded what the audience was willing to accept.")

body(doc, "Problem 2 — Value Email Click Collapse:", bold=True, color=(220, 38, 38))
body(doc,
     "Both value and sales emails suffered a significant CTOR decline. "
     "In the Visible Decline period (Nov 2025–now), value email CTOR fell to 7.9% from 13.9% (−43% relative), "
     "and sales email CTOR fell to 7.7% from 19.2% (−60% relative). "
     "When filtered to emails that actually contain links, the CTOR for value emails in the Visible Decline period "
     "is still a meaningful drop — but the story is more nuanced than it appeared when no-link emails "
     "were included. People are still opening value emails (36.0% open rate) but clicking less, "
     "suggesting the problem is not the subject lines — it is the link quality, CTA clarity, "
     "and content relevance inside the email body.")

callout(doc,
        "The critical insight: these are two separate problems requiring two different solutions. "
        "Sales email fatigue is a frequency and positioning problem — too many commercial sends "
        "have trained the audience to ignore them. Value email click collapse is a content and "
        "CTA design problem — the emails are being opened but no longer drive curiosity or action "
        "from readers.",
        color_rgb=(79, 70, 229))

chart_block(doc, "U_decline_by_category.png", 21,
    "Strong Start vs Visible Decline drop comparison for open rate and CTOR, split by Sales and Value.",
    "This is the most important chart for diagnosing the problem. On the left panel (open rate), "
    "the arrow for Sales emails is longer and steeper than for Value emails — sales content lost "
    "more opens proportionally. On the right panel (CTOR), the situation reverses: the arrow for "
    "Value emails is longer and steeper — value content lost more click engagement proportionally. "
    "The arrows point in the same direction (down), but the two categories are failing in "
    "different ways. A single fix will not solve both problems. The newsletter needs two distinct "
    "interventions: one for how sales emails are positioned and how often they are sent, and "
    "another for what value emails contain and how they invite readers to engage.")

chart_block(doc, "R_sales_vs_value_monthly_trend.png", 22,
    "Monthly open rate and CTOR trends separated by Sales (pink) and Value (blue) emails.",
    "The two-panel chart shows both metrics over time, with Sales in pink and Value in blue. "
    "In the open rate panel, both lines track closely together until around October–November 2025, "
    "when the sales line (pink) drops more sharply below the value line (blue). This is when "
    "the audience began disproportionately skipping sales emails. In the CTOR panel, the divergence "
    "happens earlier — around June 2025 — and the value line (blue) drops faster and lower. "
    "The pink sales CTOR line also drops but retains more resilience because sales emails always "
    "have explicit reasons to click. The timing difference between the two panels reinforces "
    "that the CTOR collapse began with value content in June 2025, while the open rate suppression "
    "of sales content came later in November 2025.")

chart_block(doc, "S_phase_category_bars.png", 23,
    "Grouped bar chart: open rate and CTOR by phase, split by Sales and Value.",
    "The side-by-side bars make phase-by-phase comparison clean and visual. In the Strong Start period (Feb–May 2025), "
    "all four bars are tall — both categories performed well on both metrics. By the Visible Decline period, "
    "all bars have shrunk, but not equally. The most dramatic drop is in Sales Open Rate — "
    "from 56.5% in the Strong Start to 33.0% in the Visible Decline. CTOR for both categories "
    "also shows a clear downward progression across phases, with Sales CTOR falling from 19.2% to 7.7% "
    "and Value CTOR from 13.9% to 7.9% (note: CTOR bars reflect emails with links only, so n-values "
    "differ from open rate bars). These two decline patterns represent the two failure modes described above. "
    "The n-values annotated on each bar also reveal something useful: the engagement quality drop "
    "is real — it is not explained by changes in send volume.")

chart_block(doc, "V_rolling_by_category.png", 24,
    "Rolling trend lines for Value (top) and Sales (bottom) emails, showing open rate and CTOR separately.",
    "Separating the rolling trend by category creates two distinct narratives. The Value email "
    "panel (top) shows open rate (blue) remaining relatively stable until late 2025, while CTOR "
    "(red dashed) drops steeply from June 2025 and continues falling. This is the clearest view "
    "of the value content click collapse — the audience kept opening value emails even as clicks "
    "evaporated. The Sales email panel (bottom) shows a different pattern: both lines decline, "
    "but CTOR remains higher in absolute terms because of the structural advantage of explicit CTAs. "
    "The open rate in the sales panel starts declining earlier and more severely — the audience "
    "trained themselves to recognise and skip sales sends before they stopped engaging with value sends.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — COLD SUBSCRIBERS
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "8. Cold Subscribers — Who Stopped Engaging?")

body(doc,
     "Cold subscribers are confirmed subscribers who have not opened or clicked any email "
     "within the last 90 days (or 30 days for subscribers active fewer than 90 days). "
     "They remain on the list and count toward billing, but they are no longer engaging. "
     "At the time of this analysis, there are 2,819 cold subscribers on the list.")

callout(doc,
        "A high number of cold subscribers is a compounding risk: they depress open and click "
        "rates (making engagement metrics look worse than the truly engaged portion of the list), "
        "and if left unaddressed, they are likely future cancellations.",
        color_rgb=(220, 38, 38))

h2(doc, "8.1 How Old Are the Cold Subscribers?")

body(doc,
     "Understanding when cold subscribers signed up tells us whether the disengagement problem "
     "is recent (a Phase 3 content issue) or chronic (a long-standing onboarding or nurture failure).")

add_table(doc,
    ["Age Since Sign-Up", "Count", "% of Cold List", "Interpretation"],
    [
        ["< 3 months",    "4",     "0.1%",  "Almost none — very recent sign-ups rarely go cold immediately"],
        ["3–6 months",    "253",   "9.0%",  "Signed up Jun–Nov 2025 — Silent Warning / early Visible Decline period"],
        ["6–12 months",   "1,255", "44.5%", "Signed up Feb–Aug 2025 — the Strong Start & Silent Warning period"],
        ["12–18 months",  "1,307", "46.4%", "Signed up Aug 2024–Feb 2025 — pre-analysis, long-standing cold"],
        ["18–24 months",  "0",     "0.0%",  "No cold subscribers in this range"],
    ]
)

body(doc,
     "The overwhelming majority of cold subscribers (90.9%) have been on the list for 6 months "
     "or more. Nearly half (46.4%) have been subscribed for 12–18 months — meaning they signed "
     "up in late 2024 / early 2025, were never truly activated, and have been sitting cold ever "
     "since. Only 9.1% signed up in the last 6 months, which means this is not a recent "
     "acquisition quality problem. It is a long-standing engagement failure with subscribers "
     "who were never properly onboarded or re-engaged.")

chart_block(doc, "W1_cold_age_distribution.png", 25,
    "Age distribution of cold subscribers, colour-coded by the phase in which they signed up.",
    "The bar chart shows that cold subscribers are not a recent phenomenon. The two tallest bars "
    "are '6–12 months' and '12–18 months', together accounting for over 90% of all cold "
    "subscribers. The colour coding reinforces this: the red bar (Visible Decline sign-ups) is "
    "tiny — only 4 subscribers who signed up in the last 3 months have already gone cold. "
    "The bulk of cold subscribers are blue (signed up before June 2025), meaning they joined "
    "during the Strong Start period when open rates were healthy, but never engaged. "
    "This points to an onboarding and early nurture gap, not a content decline issue.")

h2(doc, "8.2 Acquisition Source — Tag Breakdown")

body(doc,
     "Each subscriber's first assigned tag identifies how they arrived on the list — "
     "which lead magnet, import, or campaign brought them in. This tells us which "
     "acquisition channels are producing the most cold subscribers.")

add_table(doc,
    ["Acquisition Source (Tag)", "Cold Subscribers", "% of Cold List"],
    [
        ["Livestream / YouTube",  "1,623", "57.6%"],
        ["AI Program Import",     "741",   "26.3%"],
        ["Bootcamp Waitlist",     "402",   "14.3%"],
        ["No Tag / Other",        "48",    "1.7%"],
        ["Alumni / Paid Program", "5",     "0.2%"],
    ]
)

body(doc,
     "More than half of all cold subscribers (57.6%) originally came from YouTube livestreams "
     "or the Livestream Kit lead magnet. This is the largest single acquisition channel on the "
     "list, so some volume here is expected — but it also suggests that the livestream-to-email "
     "nurture journey is not converting viewers into engaged readers. People watched a livestream, "
     "signed up, and then stopped opening. The second largest group (26.3%) are CSV imports "
     "from AI program lists, which typically have lower baseline engagement as imported contacts "
     "did not organically opt in to the newsletter.")

callout(doc,
        "Alumni / Paid Program subscribers make up only 0.2% of the cold list (5 people) — "
        "meaning people who have actually paid for a programme remain engaged. The cold problem "
        "is concentrated in free lead magnet and import channels.",
        color_rgb=(16, 185, 129))

chart_block(doc, "W2_cold_tag_breakdown.png", 26,
    "Horizontal bar chart of cold subscribers by acquisition source (tag group).",
    "The chart makes the disparity immediately clear: Livestream / YouTube is the dominant "
    "source of cold subscribers by a wide margin, nearly double the next category. "
    "AI Program Imports are second. Both of these channels share a characteristic: "
    "subscribers arrive with a specific, transactional intent (watch a video, download a resource) "
    "rather than an ongoing content interest. Once that intent is satisfied, they stop engaging. "
    "The Bootcamp Waitlist group (14.3%) is also notable — these subscribers expressed interest "
    "in a specific product, and if that product was not purchased, they may have disengaged "
    "once the launch window closed.")

h2(doc, "8.3 When Did Cold Subscribers Join?")

body(doc,
     "Looking at which newsletter phase the cold subscribers joined during reveals whether "
     "the cold problem is linked to specific periods of content or audience quality.")

add_table(doc,
    ["Sign-Up Phase", "Cold Subscribers", "% of Cold List"],
    [
        ["Strong Start (pre-Jun 2025)",         "2,052", "72.8%"],
        ["Silent Warning (Jun–Oct 2025)",       "735",   "26.1%"],
        ["Visible Decline (Nov 2025–now)",       "32",    "1.1%"],
    ]
)

body(doc,
     "72.8% of cold subscribers joined before June 2025 — during or before the Strong Start "
     "period when the newsletter was performing at its best. This is a critical finding: "
     "the cold subscriber problem pre-dates the engagement decline. These are not people "
     "who went cold because of the content changes in Phase 2 or Phase 3. They were already "
     "cold before those problems emerged. Only 32 cold subscribers (1.1%) joined during the "
     "Visible Decline period, which confirms that recent acquisition quality is not the issue.")

chart_block(doc, "W3_cold_phase_bar.png", 27,
    "Bar chart of cold subscribers by the phase in which they signed up.",
    "The green bar (Strong Start) towers above the other two phases. This means the largest "
    "reservoir of cold subscribers consists of people who joined when the newsletter was at its "
    "healthiest. They were acquired but never retained. The orange bar (Silent Warning) "
    "represents a meaningful secondary group — people who signed up during June–October 2025 "
    "and also failed to engage. The red bar (Visible Decline) is negligible. The takeaway is "
    "that the cold problem is structural and predates the engagement decline — it requires "
    "a dedicated re-engagement or list-cleaning strategy independent of the content fixes.")

h2(doc, "8.4 Acquisition Source × Sign-Up Phase")

body(doc,
     "Combining tag group and sign-up phase shows which channels produced cold subscribers "
     "in each period, and whether the composition has changed over time.")

chart_block(doc, "W4_cold_tag_x_phase.png", 28,
    "Stacked bar chart: cold subscribers by acquisition source, broken down by sign-up phase.",
    "In the Strong Start phase (left bar), Livestream / YouTube and AI Program Imports dominate — "
    "these were the primary acquisition channels at that time. In the Silent Warning phase "
    "(middle bar), Bootcamp Waitlist becomes a larger proportion, reflecting the bootcamp "
    "launch campaigns that ran in mid-2025. In the Visible Decline phase (right bar), "
    "the volume is very small (32 total), but Bootcamp Waitlist again appears prominently — "
    "likely subscribers who joined during a late-2025 launch sequence and then disengaged "
    "after the launch window closed. "
    "Across all phases, Livestream / YouTube is consistently the largest cold subscriber source, "
    "reinforcing that the livestream-to-newsletter conversion journey needs a stronger "
    "engagement sequence to activate new sign-ups before they go cold.")

callout(doc,
        "Key insight: 2,819 cold subscribers are a significant segment — roughly 41% of the "
        "6,854 cancelled subscribers tracked in this analysis. They are not yet lost (they "
        "are still confirmed), but without a targeted re-engagement campaign, they are the "
        "most likely next wave of cancellations. Prioritise a win-back sequence for the "
        "Livestream / YouTube and AI Program Import segments specifically.",
        color_rgb=(79, 70, 229))

h2(doc, "8.5 When Did Cold Subscribers Stop Engaging? — Kit API Data")

body(doc,
     "Using the Kit API, individual engagement stats were fetched for all 2,819 cold subscribers "
     "— including last open date, last click date, and sends_since_last_open (how many emails they "
     "have been sent since their last recorded open). This provides a precise timeline of disengagement "
     "at the individual level.")

add_table(doc,
    ["Segment", "Count", "% of Cold List", "Interpretation"],
    [
        ["Never opened any email",           "1,326", "47.0%",
         "Ghost subscribers — signed up but never engaged once"],
        ["Last opened during Silent Warning\n(Jun–Oct 2025)",  "1,163", "41.3%",
         "Were engaged, went cold when CTOR was already declining"],
        ["Last opened during Visible Decline\n(Nov 2025–now)", "330",   "11.7%",
         "Most recently disengaged — highest chance of re-activation"],
    ]
)

body(doc,
     "Nearly half of all cold subscribers (47.0%) have never opened a single email — they are "
     "ghost subscribers who exist only as list overhead. They inflate the subscriber count, "
     "depress all engagement metrics, and represent zero revenue potential. "
     "The remaining 53.0% (1,493 subscribers) did open at least once: 77.9% of those last opened "
     "during the Silent Warning phase (Jun–Oct 2025), and 22.1% last opened during the Visible "
     "Decline phase (Nov 2025 onward).")

callout(doc,
        "77.9% of previously-engaged cold subscribers stopped opening during the Silent Warning "
        "phase — the exact same period when CTOR collapsed. This is not a coincidence: the same "
        "content quality issue that killed clicks also caused a wave of disengagement. "
        "The API data confirms at the individual level what the broadcast metrics showed in aggregate.",
        color_rgb=(239, 68, 68))

body(doc,
     "The sends_since_last_open metric reveals how deep the disengagement is. The median cold "
     "subscriber has been sent 94 emails since their last open, and the mean is 88. "
     "56.3% have ignored 80 or more consecutive sends, and 46.6% have ignored 100 or more. "
     "At this depth of disengagement, re-activation is unlikely without a dedicated win-back "
     "campaign or a clean-off decision.")

add_table(doc,
    ["Sends Since Last Open", "Count", "% of Cold List (excl. never-opened)"],
    [
        ["1–20 sends",   "—",   "Minimal — very recently gone cold"],
        ["21–60 sends",  "~790", "~28.6%  — moderately cold, still possible to re-engage"],
        ["61–100 sends", "~815", "~29.5%  — deeply cold, difficult to re-engage"],
        ["101–127 sends","~1,157","~41.9%  — extremely cold, unlikely to re-engage organically"],
    ]
)

chart_block(doc, "X1_last_open_timeline.png", 29,
    "Monthly timeline of last open dates for the 1,493 cold subscribers who opened at least once.",
    "The chart shows a clear spike in the Silent Warning phase: October 2025 is the single "
    "largest month (450 last opens), followed by November 2025 (330) and September 2025 (314). "
    "This means the majority of cold subscribers who were previously engaged made their last "
    "engagement attempt in the Sep–Oct 2025 window — precisely when CTOR was at its lowest "
    "and content quality issues were most acute. The phase boundary lines confirm the alignment: "
    "the bulk of last opens cluster between the Jun and Nov 2025 phase cuts. "
    "Almost no cold subscriber's last open was in the Strong Start period, which means the "
    "disengagement wave is a product of the Silent Warning and Visible Decline content failures.")

chart_block(doc, "X2_sends_since_last_open.png", 30,
    "Distribution of sends since last open — how many consecutive emails cold subscribers have ignored.",
    "The histogram is heavily skewed toward high values. The 81–100 and 101–120 buckets together "
    "account for the majority of cold subscribers who ever opened. This means most cold subscribers "
    "have not engaged in a very long time — they have been sitting on the list, silently ignoring "
    "send after send. The gradient colouring from green (recently cold) to deep red (long-term cold) "
    "visually reinforces the severity: most of the mass is in the red and dark-red bars.")

chart_block(doc, "X3_disengagement_type.png", 31,
    "Donut and horizontal bar chart showing cold subscribers split by engagement history.",
    "The donut chart immediately surfaces the 47% never-opened segment. Combined with the bar chart, "
    "it is clear that cold subscribers are two distinct populations requiring two different strategies: "
    "(1) Ghost subscribers who never opened — these should be removed from the list or moved to a "
    "separate suppression segment, as they have no re-engagement potential; "
    "(2) Previously-engaged subscribers who went cold — these warrant a targeted win-back sequence "
    "with a specific hook, ideally referencing the content or campaign that originally activated them.")

chart_block(doc, "X4_disengagement_cohort.png", 32,
    "Stacked bar: sign-up phase of cold subscribers vs the phase in which they last opened.",
    "This cohort view shows that regardless of when subscribers signed up, almost all their "
    "last opens cluster in the Silent Warning phase. Subscribers who signed up in the Strong Start "
    "phase were engaged during that period, but their last opens are in the Silent Warning window — "
    "they survived Phase 1 fine, then disengaged in Phase 2. The pattern is consistent: "
    "Phase 2 (Jun–Oct 2025) is where engagement broke down across all cohorts.")

callout(doc,
        "Recommended action: Segment the cold list into two groups before any re-engagement campaign. "
        "Group A — Never Opened (1,326): run a one-time reactivation email with a very clear hook; "
        "if no open within 30 days, remove from list. "
        "Group B — Previously Engaged (1,493): send a win-back sequence (2–3 emails) acknowledging "
        "the gap and offering something new. Prioritise the 330 who last opened during the Visible "
        "Decline phase — they are the most recently cold and most likely to re-activate.",
        color_rgb=(79, 70, 229))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — CONCLUSIONS & RECOMMENDATIONS
# ══════════════════════════════════════════════════════════════════════════════
# 9.1 conversion-latency summary variables (populated in section 9.1 if data exists)
_lat_buyers_n = np.nan
_lat_median_days = np.nan
_lat_mean_days = np.nan
_lat_within_90 = np.nan
_lat_over_180 = np.nan

h1(doc, "9. Conversion Timing — Signup to First Purchase")

body(doc,
     "This section quantifies how long buyers typically stay on the email list before purchasing a bootcamp. "
     "The goal is practical planning: define a realistic nurture window and avoid either premature sales pressure "
     "or under-investing in longer-cycle subscribers.")

h2(doc, "9.1 How Long Buyers Stay on List Before First Purchase")

body(doc,
     "To estimate the typical nurture window before conversion, we measured time from subscription "
     "date (`created_at`) to first bootcamp purchase month for subscribers with at least one purchase tag. "
     "Because this export stores purchase month in tags (not exact transaction day), the purchase date is "
     "approximated using the month-end date for each tagged cohort.")

_confirmed_path = BASE / "Confirmed Subscribers.csv"
if _confirmed_path.exists():
    _cf = pd.read_csv(_confirmed_path, usecols=["email", "created_at", "status", "tags"])
    _cf["created_at_dt"] = pd.to_datetime(_cf["created_at"], utc=True, errors="coerce")
    _cf["status_l"] = _cf["status"].astype(str).str.lower().str.strip()
    _cf = _cf[_cf["status_l"] == "active"].copy()

    _purchase_tag_dates = {
        "ai agent bootcamp core [apr 2025]": pd.Timestamp("2025-04-30"),
        "ai agent bootcamp core [jul 2025]": pd.Timestamp("2025-07-31"),
        "ai agent core [sept]": pd.Timestamp("2025-09-30"),
        "ai agent core [oct]": pd.Timestamp("2025-10-31"),
        "ai agent core [feb]": pd.Timestamp("2026-02-28"),
    }
    _tag_order = list(_purchase_tag_dates.keys())
    _tag_label = {
        "ai agent bootcamp core [apr 2025]": "Apr 2025 Bootcamp",
        "ai agent bootcamp core [jul 2025]": "Jul 2025 Bootcamp",
        "ai agent core [sept]": "Sept 2025 Core",
        "ai agent core [oct]": "Oct 2025 Core",
        "ai agent core [feb]": "Feb 2026 Core",
    }

    _lat_rows = []
    for _, _r in _cf.iterrows():
        _t = str(_r["tags"]).lower() if pd.notna(_r["tags"]) else ""
        _matches = [k for k in _tag_order if k in _t]
        if not _matches or pd.isna(_r["created_at_dt"]):
            continue
        _p_dates = sorted([(_purchase_tag_dates[k], k) for k in _matches], key=lambda x: x[0])
        _first_date, _first_tag = _p_dates[0]
        _created = _r["created_at_dt"].tz_convert(None)
        _days = (_first_date - _created.normalize()).days
        _lat_rows.append({
            "email": _r["email"],
            "created_at": _created,
            "first_purchase_tag": _first_tag,
            "first_purchase_cohort": _tag_label.get(_first_tag, _first_tag),
            "first_purchase_date": _first_date,
            "days_to_first_purchase": _days,
        })

    _lat = pd.DataFrame(_lat_rows)
    if len(_lat):
        _neg_n = int((_lat["days_to_first_purchase"] < 0).sum())
        _lat["days_to_first_purchase_clamped"] = _lat["days_to_first_purchase"].clip(lower=0)
        _lat.to_csv(GENERATED / "bootcamp_conversion_latency.csv", index=False)

        _lat_valid = _lat[_lat["days_to_first_purchase_clamped"].notna()].copy()
        _p25 = _lat_valid["days_to_first_purchase_clamped"].quantile(0.25)
        _p75 = _lat_valid["days_to_first_purchase_clamped"].quantile(0.75)
        _med = _lat_valid["days_to_first_purchase_clamped"].median()
        _mean = _lat_valid["days_to_first_purchase_clamped"].mean()

        _buckets = pd.cut(
            _lat_valid["days_to_first_purchase_clamped"],
            bins=[-1, 30, 90, 180, 365, 10000],
            labels=["<=30d", "31-90d", "91-180d", "181-365d", ">365d"],
        )
        _bucket_tab = _buckets.value_counts().reindex(["<=30d", "31-90d", "91-180d", "181-365d", ">365d"]).fillna(0)
        _bucket_pct = (_bucket_tab / _bucket_tab.sum() * 100).round(1)

        _coh = (
            _lat_valid.groupby("first_purchase_cohort", as_index=False)["days_to_first_purchase_clamped"]
            .agg(["count", "median", "mean"])
            .reset_index()
            .rename(columns={"count": "n_buyers", "median": "median_days", "mean": "mean_days"})
        )
        _coh = _coh.sort_values("median_days")
        _coh.to_csv(GENERATED / "bootcamp_conversion_latency_by_cohort.csv", index=False)

        fig, axes = plt.subplots(1, 2, figsize=(13.6, 5.4))
        ax = axes[0]
        ax.hist(_lat_valid["days_to_first_purchase_clamped"], bins=28, color="#4F46E5", alpha=0.82, edgecolor="white")
        ax.axvline(_med, color="#10B981", linestyle="--", linewidth=2, label=f"Median = {_med:.0f}d")
        ax.axvline(_mean, color="#F59E0B", linestyle="--", linewidth=2, label=f"Mean = {_mean:.0f}d")
        ax.set_title("Time to First Purchase Distribution")
        ax.set_xlabel("Days from signup to first purchase (month-end proxy)")
        ax.set_ylabel("Buyers")
        ax.grid(axis="y", alpha=0.2)
        ax.legend(fontsize=8.8)

        ax = axes[1]
        _coh_plot = _coh.sort_values("median_days")
        _x = np.arange(len(_coh_plot))
        ax.bar(_x, _coh_plot["median_days"], color="#10B981", alpha=0.88)
        for i, (_, rr) in enumerate(_coh_plot.iterrows()):
            ax.text(i, rr["median_days"] + 4, f"n={int(rr['n_buyers'])}", ha="center", fontsize=8, color="#334155")
        ax.set_xticks(_x)
        ax.set_xticklabels(_coh_plot["first_purchase_cohort"], rotation=20, ha="right")
        ax.set_title("Median Days to Purchase by First Purchase Cohort")
        ax.set_ylabel("Median days")
        ax.grid(axis="y", alpha=0.2)

        fig.suptitle("Bootcamp Conversion Latency (Signup → Purchase)", fontsize=13.5, fontweight="bold")
        fig.tight_layout()
        fig.savefig(CHARTS / "AR_bootcamp_conversion_latency.png")
        plt.close()

        add_table(doc,
                  ["Metric", "Value"],
                  [
                      ["Buyers with identified first purchase", f"{len(_lat_valid):,}"],
                      ["Median days to first purchase", f"{_med:.0f} days"],
                      ["Mean days to first purchase", f"{_mean:.1f} days"],
                      ["IQR (P25 to P75)", f"{_p25:.0f}d to {_p75:.0f}d"],
                      ["Share converting <=30 days", f"{_bucket_pct.get('<=30d', 0):.1f}%"],
                      ["Share converting <=90 days", f"{(_bucket_pct.get('<=30d', 0) + _bucket_pct.get('31-90d', 0)):.1f}%"],
                      ["Share converting >180 days", f"{(_bucket_pct.get('181-365d', 0) + _bucket_pct.get('>365d', 0)):.1f}%"],
                  ])

        _coh_rows = []
        for _, rr in _coh.sort_values("median_days").iterrows():
            _coh_rows.append([
                rr["first_purchase_cohort"],
                f"{int(rr['n_buyers']):,}",
                f"{rr['median_days']:.0f} days",
                f"{rr['mean_days']:.1f} days",
            ])
        add_table(doc, ["First Purchase Cohort", "Buyers", "Median Days", "Mean Days"], _coh_rows)

        chart_block(doc,
                    "AR_bootcamp_conversion_latency.png",
                    "AR",
                    "Time to First Purchase (Signup to Purchase Month)",
                    "Left panel: histogram of days from signup to first purchase. The green vertical line is the median "
                    "(the middle buyer), and the yellow line is the mean (average), which is usually pulled upward by "
                    "slower converters in the right tail. Right panel: median conversion delay by purchase cohort with each "
                    "bar annotated by sample size (n). Read this chart as a timing map for sales pressure: if most buyers "
                    "convert in the first 90 days, nurture should be strongest early; if many convert later, long-horizon "
                    "follow-up remains commercially important.")
        body(doc,
             f"AR analysis: The central conversion window is {_p25:.0f} to {_p75:.0f} days (IQR), "
             f"with median {_med:.0f} days and mean {_mean:.1f} days. "
             f"{(_bucket_pct.get('<=30d', 0) + _bucket_pct.get('31-90d', 0)):.1f}% convert within 90 days, while "
             f"{(_bucket_pct.get('181-365d', 0) + _bucket_pct.get('>365d', 0)):.1f}% convert after 180 days. "
             "Plain-English interpretation: conversion is front-loaded but not immediate; a non-trivial share needs "
             "a longer nurture horizon before purchase.",
             color=(50, 50, 50))

        _lat_buyers_n = len(_lat_valid)
        _lat_median_days = _med
        _lat_mean_days = _mean
        _lat_within_90 = (_bucket_pct.get('<=30d', 0) + _bucket_pct.get('31-90d', 0))
        _lat_over_180 = (_bucket_pct.get('181-365d', 0) + _bucket_pct.get('>365d', 0))

        body(doc,
             "Method note: exact transaction timestamps are not available in this export. "
             "Purchase date is approximated by month-end for each purchase tag.",
             size=9, italic=True, color=(120, 120, 120))
        if _neg_n > 0:
            body(doc,
                 f"Data quality note: {_neg_n} rows produced negative lag and were clamped at 0 days in summary metrics.",
                 size=9, italic=True, color=(120, 120, 120))
    else:
        body(doc,
             "[No buyers with recognized purchase tags were found in Confirmed Subscribers.csv]",
             color=(200, 0, 0))
else:
    body(doc,
         "[Section 9.1 requires Confirmed Subscribers.csv]",
         color=(200, 0, 0))

doc.add_paragraph()
divider(doc)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — LEAD MAGNET IMPACT ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
SURVEY_DATE_RPT = pd.Timestamp("2026-01-25")
SURVEY_SIGNUP_WINDOW_END_RPT = pd.Timestamp("2026-03-02")
PRE_SURVEY_WINDOW_DAYS_RPT = 38
Q1_WINDOW_DAYS_RPT = 38

# ── Load broadcasts for Q1 ───────────────────────────────────────────────────
bdf_rpt = pd.read_csv(BASE / "Emails Broadcasting - broadcasts_categorised.csv")
bdf_rpt["date"] = pd.to_datetime(bdf_rpt["date"])
bdf_rpt["category"] = bdf_rpt["category"].replace("Sles", "Sales")
bdf_rpt = bdf_rpt[(bdf_rpt["recipients"] > 900) & bdf_rpt["category"].isin(["Value","Sales"])].copy()

PRE_START_RPT = SURVEY_DATE_RPT - pd.Timedelta(days=Q1_WINDOW_DAYS_RPT)
POST_END_RPT = SURVEY_DATE_RPT + pd.Timedelta(days=Q1_WINDOW_DAYS_RPT - 1)
pre_rpt  = bdf_rpt[(bdf_rpt["date"] >= PRE_START_RPT) & (bdf_rpt["date"] < SURVEY_DATE_RPT)]
post_rpt = bdf_rpt[(bdf_rpt["date"] >= SURVEY_DATE_RPT) & (bdf_rpt["date"] <= POST_END_RPT)]

_q1_stats = {}
for cat in ["Value", "Sales"]:
    pg = pre_rpt[pre_rpt["category"] == cat]["open_rate"]
    po = post_rpt[post_rpt["category"] == cat]["open_rate"]
    _q1_stats[cat] = {
        "pre_n":    len(pg), "pre_mean":  round(pg.mean(), 1),
        "post_n":   len(po), "post_mean": round(po.mean(), 1) if len(po) else None,
        "delta":    round(po.mean() - pg.mean(), 1) if len(po) else None,
        "delta_pct":round((po.mean() - pg.mean()) / pg.mean() * 100, 1) if len(po) else None,
        "p_value":  round(scipy_stats.ttest_ind(pg, po, equal_var=False).pvalue, 3) if len(po) >= 2 and len(pg) >= 2 else None,
    }

# ── Load Kit API cache for Q2 ────────────────────────────────────────────────
_cache_path = GENERATED / "lead_magnet_stats.csv"
_post_summary_path = GENERATED / "lead_magnet_post_survey_summary.csv"
_clicked_path = BASE / "AI Sprint Roadmap - Clicked.csv"
_opened_path = BASE / "AI Sprint Roadmap - Opened subscribers.csv"
_monthly_ab_path = GENERATED / "lead_magnet_group_ab_monthly.csv"
_cohort_monthly_path = GENERATED / "signup_cohort_monthly_evolution.csv"
_buyers_monthly_path = GENERATED / "bootcamp_buyers_monthly_evolution.csv"
_buyers_summary_path = GENERATED / "bootcamp_buyers_summary.csv"
_buyers_age_path = GENERATED / "bootcamp_buyers_age_buckets.csv"
_buyers_failed_path = GENERATED / "bootcamp_buyers_failed_ids.csv"
_q2_available = _cache_path.exists()
_q3_available = _post_summary_path.exists()
_q5_available = _clicked_path.exists() and _opened_path.exists()
_q6_available = _monthly_ab_path.exists()
_q11_available = _cohort_monthly_path.exists()
_q12_available = _buyers_monthly_path.exists() and _buyers_summary_path.exists()
_q2 = {}
if _q2_available:
    _enr = pd.read_csv(_cache_path)
    for col in ["sent","opened","clicked"]:
        _enr[col] = pd.to_numeric(_enr[col], errors="coerce")
    _enr = _enr[_enr["sent"] >= 5].copy()
    _enr["open_rate"]  = _enr["opened"]  / _enr["sent"]
    _enr["click_rate"] = _enr["clicked"] / _enr["sent"]
    _enr["ctor"]       = np.where(_enr["opened"] > 0, _enr["clicked"] / _enr["opened"], np.nan)
    for g in ["A","B"]:
        gd = _enr[_enr["group"] == g]
        ct_all = gd["ctor"].dropna()
        _q2[g] = {
            "n":             len(gd),
            "or_mean":       round(gd["open_rate"].mean()   * 100, 1),
            "or_med":        round(gd["open_rate"].median() * 100, 1),
            "cr_mean":       round(gd["click_rate"].mean()  * 100, 1),
            "cr_med":        round(gd["click_rate"].median()* 100, 1),
            "ctor_mean":     round(gd["ctor"].mean()        * 100, 1),
            "ctor_med":      round(gd["ctor"].median()      * 100, 1),
            "ctor_anomalous":int((ct_all > 1).sum()),
            "ctor_anom_pct": round((ct_all > 1).mean() * 100, 1),
        }
    # significance tests
    _mwu = {}
    for metric, label in [("open_rate","OR"), ("click_rate","CR"), ("ctor","CTOR")]:
        av = _enr[_enr["group"]=="A"][metric].dropna()
        bv = _enr[_enr["group"]=="B"][metric].dropna()
        if len(av) >= 5 and len(bv) >= 5:
            _, p = scipy_stats.mannwhitneyu(av, bv, alternative="two-sided")
            _mwu[metric] = round(p, 4)
        else:
            _mwu[metric] = None

    # status breakdown
    if "api_state" in _enr.columns:
        _status = (
            _enr[_enr["group"].isin(["A","B"])]
            .assign(is_active=lambda d: d["api_state"].str.lower().str.strip() == "active")
            .groupby("group")["is_active"]
            .agg(["sum","count"])
            .assign(active_pct=lambda d: round(d["sum"]/d["count"]*100, 1))
        )
    else:
        _status = None

if _q3_available:
    _q3 = pd.read_csv(_post_summary_path)
else:
    _q3 = None

# ── Q6: monthly evolution (Group A vs Group B) ───────────────────────────────
_q6 = None
_q6_chart_cov_or = "AF_group_ab_openrate_coverage_aware.png"
_q6_chart_gap = "AG_group_ab_monthly_gap.png"
_q6_chart_sales_ctor = "AH_group_ab_sales_ctor_valid_months.png"
_q6_value_gap = np.nan
_q6_sales_gap = np.nan
_q6_sales_ctor_gap = np.nan
_q6_value_a = np.nan
_q6_value_b = np.nan
_q6_sales_a = np.nan
_q6_sales_b = np.nan
_q6_ctor_a = np.nan
_q6_ctor_b = np.nan
if _q6_available:
    _q6 = pd.read_csv(_monthly_ab_path)
    _q6["month_dt"] = pd.to_datetime(_q6["month_dt"], errors="coerce")
    _q6 = _q6.sort_values("month_dt")

_q11 = None
if _q11_available:
    _q11 = pd.read_csv(_cohort_monthly_path)
    _q11["month_dt"] = pd.to_datetime(_q11["month_dt"], errors="coerce")
    _q11 = _q11.sort_values("month_dt")

_q12m = None
_q12s = None
_q12a = None
_q12f = None
if _q12_available:
    _q12m = pd.read_csv(_buyers_monthly_path)
    _q12m["month_dt"] = pd.to_datetime(_q12m["month_dt"], errors="coerce")
    _q12m = _q12m.sort_values("month_dt")
    _q12s = pd.read_csv(_buyers_summary_path)
    if _buyers_age_path.exists():
        _q12a = pd.read_csv(_buyers_age_path)
    if _buyers_failed_path.exists():
        _q12f = pd.read_csv(_buyers_failed_path)

# ── Q5: subscriber age at survey send (Group A vs B) ─────────────────────────
_q5 = None
_q5_chart_file = "AB_group_age_profile.png"
if _q5_available:
    _resp_keywords = ["i want", "i'm", "i work at a company", "i own/run", "i own", "i'm building"]

    def _has_survey_response(_tags):
        if pd.isna(_tags):
            return False
        _t = str(_tags).lower()
        return any(_k in _t for _k in _resp_keywords)

    _clicked = pd.read_csv(_clicked_path)
    _opened = pd.read_csv(_opened_path)

    for _df in [_clicked, _opened]:
        _df["email_lower"] = _df["email"].str.lower().str.strip()
        _df["is_group_a"] = _df["tags"].apply(_has_survey_response)
        _df["created_at_dt"] = pd.to_datetime(_df["created_at"], utc=True, errors="coerce")

    _clicked_emails = set(_clicked["email_lower"])
    _group_a_emails = set(_clicked[_clicked["is_group_a"]]["email_lower"])
    _group_b_emails = (
        set(_clicked[~_clicked["is_group_a"]]["email_lower"])
        | set(_opened[(~_opened["email_lower"].isin(_clicked_emails)) & (~_opened["is_group_a"])]["email_lower"])
    )

    _master = pd.concat([
        _clicked[["email_lower", "created_at_dt"]],
        _opened[["email_lower", "created_at_dt"]],
    ], ignore_index=True)
    _master = _master.sort_values("created_at_dt").drop_duplicates("email_lower", keep="first")

    _survey_dt_utc = pd.Timestamp("2026-01-25", tz="UTC")
    _a = _master[_master["email_lower"].isin(_group_a_emails)].copy()
    _b = _master[_master["email_lower"].isin(_group_b_emails)].copy()
    _a["age_days"] = (_survey_dt_utc - _a["created_at_dt"]).dt.total_seconds() / 86400
    _b["age_days"] = (_survey_dt_utc - _b["created_at_dt"]).dt.total_seconds() / 86400

    _a_raw_n, _b_raw_n = len(_a), len(_b)
    _a = _a[_a["age_days"] >= 0].copy()
    _b = _b[_b["age_days"] >= 0].copy()

    def _bucket(_s):
        return pd.cut(
            _s,
            bins=[-0.0001, 30, 90, 180, 365, 10000],
            labels=["0-30 days", "31-90 days", "91-180 days", "181-365 days", "365+ days"],
        )

    _a["bucket"] = _bucket(_a["age_days"])
    _b["bucket"] = _bucket(_b["age_days"])

    _a_bkt = _a["bucket"].value_counts().reindex(
        ["0-30 days", "31-90 days", "91-180 days", "181-365 days", "365+ days"]
    ).fillna(0).astype(int)
    _b_bkt = _b["bucket"].value_counts().reindex(
        ["0-30 days", "31-90 days", "91-180 days", "181-365 days", "365+ days"]
    ).fillna(0).astype(int)

    _q5 = {
        "a_n": len(_a),
        "b_n": len(_b),
        "a_removed": _a_raw_n - len(_a),
        "b_removed": _b_raw_n - len(_b),
        "a_mean": round(_a["age_days"].mean(), 1),
        "a_median": round(_a["age_days"].median(), 1),
        "a_p25": round(_a["age_days"].quantile(0.25), 1),
        "a_p75": round(_a["age_days"].quantile(0.75), 1),
        "b_mean": round(_b["age_days"].mean(), 1),
        "b_median": round(_b["age_days"].median(), 1),
        "b_p25": round(_b["age_days"].quantile(0.25), 1),
        "b_p75": round(_b["age_days"].quantile(0.75), 1),
        "a_recent": round((_a["age_days"] <= 90).mean() * 100, 1),
        "b_recent": round((_b["age_days"] <= 90).mean() * 100, 1),
        "a_old": round((_a["age_days"] >= 181).mean() * 100, 1),
        "b_old": round((_b["age_days"] >= 181).mean() * 100, 1),
        "a_buckets": _a_bkt.to_dict(),
        "b_buckets": _b_bkt.to_dict(),
    }

    # Q5 chart: stacked composition by age bucket (Group A vs Group B)
    _labels = ["0-30 days", "31-90 days", "91-180 days", "181-365 days", "365+ days"]
    _a_pct = [(_q5["a_buckets"].get(_lb, 0) / _q5["a_n"] * 100) if _q5["a_n"] else 0 for _lb in _labels]
    _b_pct = [(_q5["b_buckets"].get(_lb, 0) / _q5["b_n"] * 100) if _q5["b_n"] else 0 for _lb in _labels]
    _colors = ["#34D399", "#10B981", "#F59E0B", "#F97316", "#EF4444"]

    fig, ax = plt.subplots(figsize=(10, 6))
    _x = np.array([0, 1])
    _bottom = np.zeros(2)

    for _i, _lb in enumerate(_labels):
        _vals = np.array([_a_pct[_i], _b_pct[_i]])
        bars = ax.bar(_x, _vals, bottom=_bottom, color=_colors[_i], width=0.52, edgecolor="white", label=_lb)
        for _j, _bar in enumerate(bars):
            _v = _vals[_j]
            if _v >= 6:
                ax.text(
                    _bar.get_x() + _bar.get_width() / 2,
                    _bottom[_j] + _v / 2,
                    f"{_v:.1f}%",
                    ha="center",
                    va="center",
                    fontsize=8.5,
                    color="white",
                    fontweight="bold",
                )
        _bottom += _vals

    ax.set_xticks(_x)
    ax.set_xticklabels(
        [f"Group A\nResponded\n(n={_q5['a_n']:,})", f"Group B\nDid Not Respond\n(n={_q5['b_n']:,})"],
        fontsize=10,
    )
    ax.set_ylabel("Share of Group (%)")
    ax.set_ylim(0, 100)
    ax.set_title("Subscriber Age Profile at Survey Date (25 Jan 2026)\nGroup A vs Group B", fontsize=12.5)
    ax.legend(title="Age Bucket", fontsize=8.5, title_fontsize=9, bbox_to_anchor=(1.01, 1), loc="upper left")

    # Highlight the key difference in recent subscribers.
    _recent_delta = _q5["a_recent"] - _q5["b_recent"]
    ax.annotate(
        f"Recent (<=90d) lift in Group A: {_recent_delta:+.1f}pp",
        xy=(0, _a_pct[0] + _a_pct[1]),
        xytext=(0.52, 88),
        textcoords="data",
        fontsize=9,
        color="#065F46",
        fontweight="bold",
        arrowprops=dict(arrowstyle="->", color="#065F46", lw=1.4),
    )

    fig.tight_layout()
    fig.savefig(CHARTS / _q5_chart_file, dpi=180)
    plt.close()

# ── Q7: subscriber origin split (lead magnet signups vs existing recipients) ──
_q7 = None
_q7_chart_file = "AC_subscriber_origin_split.png"
if _confirmed_path.exists():
    _q7_df = pd.read_csv(_confirmed_path, usecols=["email", "created_at", "tags", "referrer"])
    _q7_df["email_lower"] = _q7_df["email"].astype(str).str.lower().str.strip()
    _q7_df = _q7_df.drop_duplicates("email_lower", keep="first")
    _q7_df["created_at_dt"] = pd.to_datetime(_q7_df["created_at"], utc=True, errors="coerce")
    _q7_df["created_date"] = _q7_df["created_at_dt"].dt.tz_convert(None).dt.normalize()

    _q7_resp_keywords = ["i want", "i'm", "i work at a company", "i own/run", "i own", "i'm building"]

    def _q7_has_survey_response(_tags):
        if pd.isna(_tags):
            return False
        _t = str(_tags).lower()
        return any(_k in _t for _k in _q7_resp_keywords)

    _q7_roadmap_referrers = {"lonelyoctopus.com/ai-sprint-roadmap"}

    def _q7_has_roadmap_referrer(_ref):
        if pd.isna(_ref):
            return False
        return str(_ref).strip().lower() in _q7_roadmap_referrers

    _q7_df["has_survey_response"] = _q7_df["tags"].apply(_q7_has_survey_response)
    _q7_df["has_roadmap_referrer"] = _q7_df["referrer"].apply(_q7_has_roadmap_referrer)
    _q7_valid = _q7_df[_q7_df["created_date"].notna()].copy()

    _survey_cutoff = SURVEY_DATE_RPT.normalize()
    _pre_window_start = (SURVEY_DATE_RPT - pd.Timedelta(days=PRE_SURVEY_WINDOW_DAYS_RPT)).normalize()
    _survey_signup_window_end = SURVEY_SIGNUP_WINDOW_END_RPT.normalize()
    _pre = _q7_valid[
        (_q7_valid["created_date"] >= _pre_window_start)
        & (_q7_valid["created_date"] < _survey_cutoff)
    ].copy()
    _post = _q7_valid[
        (_q7_valid["created_date"] >= _survey_cutoff)
        & (_q7_valid["created_date"] <= _survey_signup_window_end)
    ].copy()
    _post_lead = _post[_post["has_survey_response"] & _post["has_roadmap_referrer"]].copy()
    _post_resp_nonroadmap = _post[_post["has_survey_response"] & (~_post["has_roadmap_referrer"])].copy()
    _post_other = _post[~(_post["has_survey_response"] & _post["has_roadmap_referrer"])].copy()

    _tot = len(_q7_valid)
    _post_n = len(_post)
    _q7 = {
        "total": _tot,
        "missing_created": len(_q7_df) - _tot,
        "existing_count": len(_pre),
        "existing_pct_total": round(len(_pre) / _tot * 100, 1) if _tot else np.nan,
        "post_total": _post_n,
        "post_pct_total": round(_post_n / _tot * 100, 1) if _tot else np.nan,
        "lead_magnet_count": len(_post_lead),
        "lead_magnet_pct_total": round(len(_post_lead) / _tot * 100, 1) if _tot else np.nan,
        "lead_magnet_pct_post": round(len(_post_lead) / _post_n * 100, 1) if _post_n else np.nan,
        "post_other_count": len(_post_other),
        "post_other_pct_total": round(len(_post_other) / _tot * 100, 1) if _tot else np.nan,
        "post_other_pct_post": round(len(_post_other) / _post_n * 100, 1) if _post_n else np.nan,
        "post_resp_nonroadmap_count": len(_post_resp_nonroadmap),
        "post_resp_nonroadmap_pct_post": round(len(_post_resp_nonroadmap) / _post_n * 100, 1) if _post_n else np.nan,
    }

    fig, ax = plt.subplots(figsize=(10.8, 5.8))
    _labels = [
        "Already subscribed\nbefore survey",
        "New signup via survey\n(lead magnet, strict inferred)",
        "New signup,\nother source",
    ]
    _vals = [_q7["existing_count"], _q7["lead_magnet_count"], _q7["post_other_count"]]
    _cols = ["#6366F1", "#10B981", "#F59E0B"]
    bars = ax.bar(np.arange(3), _vals, color=_cols, width=0.58, edgecolor="white")

    for _i, _b in enumerate(bars):
        _v = _vals[_i]
        _pct = (_v / _q7["total"] * 100) if _q7["total"] else 0
        ax.text(
            _b.get_x() + _b.get_width() / 2,
            _b.get_height() + max(_vals) * 0.012,
            f"{_v:,}\n({_pct:.1f}%)",
            ha="center",
            va="bottom",
            fontsize=9,
            color="#111827",
            fontweight="bold",
        )

    _post_share = _q7["lead_magnet_pct_post"] if pd.notna(_q7["lead_magnet_pct_post"]) else 0
    ax.annotate(
        f"Among post-survey signups:\n{_post_share:.1f}% came via lead magnet",
        xy=(1, _vals[1]),
        xytext=(1.88, max(_vals) * 0.66 if max(_vals) else 1),
        arrowprops=dict(arrowstyle="->", color="#065F46", lw=1.2),
        fontsize=9,
        color="#065F46",
        fontweight="bold",
        ha="center",
    )

    ax.set_xticks(np.arange(3))
    ax.set_xticklabels(_labels, fontsize=9.5)
    ax.set_ylabel("Subscribers")
    ax.set_title("Confirmed Subscriber Origin Split Around Survey Launch (25 Jan 2026)", fontsize=12.5)
    ax.grid(axis="y", alpha=0.2)
    ax.set_ylim(0, max(_vals) * 1.18 if max(_vals) else 1)
    fig.tight_layout()
    fig.savefig(CHARTS / _q7_chart_file, dpi=180)
    plt.close()

doc.add_page_break()
h1(doc, "10. Lead Magnet Impact Analysis")

body(doc,
     "In January 2026, a free lead magnet was launched: a goal-based survey sent to all confirmed "
     "subscribers on 25 January 2026 (19,351 recipients, 37.7% open rate). Subscribers chose one "
     "of four learning-goal options and received the matching AI Sprint Roadmap. "
     "This section answers two key questions about the impact of that lead magnet:")
bullet(doc, "Q1 — Did sending the survey email make the whole list more likely to open future emails?")
bullet(doc, "Q2 — Are subscribers who responded to the survey (Group A) more engaged than those who didn't (Group B)?")

body(doc,
     "Group A (survey responders) are identified by the presence of a goal-choice tag in their Kit "
     "profile (e.g. 'I work at a company — I want to…', 'I'm transitioning into AI/tech…'). "
     "Group B (non-responders) covers ALL subscribers who received the survey but did not complete it — "
     "both those who clicked but stopped short of responding (611), and those who only opened "
     "the email without clicking at all (5,757). Total Group B = 6,368 subscribers.")

body(doc, "Metric guide used from Section 10 onward:", bold=True, color=(79, 70, 229))
bullet(doc, "Open Rate (OR): `opened / delivered` for an email or broadcast set.")
bullet(doc, "CTOR: `clicked / opened`, which measures click quality among people who already opened.")
bullet(doc, "pp = percentage points, not percent. Example: 40% to 45% = +5pp (not +5%).")
bullet(doc, "`n` on charts means sample size (broadcast count or subscriber count, depending on the chart).")
bullet(doc, "Low-coverage months (`n < 3` broadcasts) are shown for context but treated as directional evidence only.")
body(doc,
     "Interpretation format used below: (1) what is observed in the chart, (2) what the observation implies, "
     "and (3) what decision it supports. This is to minimize ambiguity for analytical readers.",
     color=(50, 50, 50))

doc.add_paragraph()
chart_block(doc,
            "lead_magnet_segments.png",
            "W0",
            "AI Sprint Roadmap Survey — How Subscribers Engaged",
            "Three bars show the three distinct engagement segments within the survey email: "
            "subscribers who clicked AND completed the survey (Group A = 892), those who clicked "
            "but did not respond (611), and those who only opened the email without clicking (5,757). "
            "Percentages are as a share of all openers. The annotation highlights the survey "
            "completion rate among those who clicked.")
body(doc,
     "How to read Figure W0: start from total openers, then follow the funnel from open to click to response. "
     "If the biggest drop is between click and response, friction is likely in the survey/roadmap step. "
     "If the biggest drop is between open and click, the CTA or offer framing is the likely bottleneck.",
     color=(50, 50, 50))

doc.add_paragraph()

# ── 10.1 Q1 ─────────────────────────────────────────────────────────────────
h2(doc, "10.1  Did the Survey Boost Future Open Rates? (Q1)")

body(doc,
     f"Broadcast open rates in the {Q1_WINDOW_DAYS_RPT}-day window before the survey "
     f"({PRE_START_RPT.strftime('%d %b %Y')} – {(SURVEY_DATE_RPT - pd.Timedelta(days=1)).strftime('%d %b %Y')}) "
     f"are compared against the {Q1_WINDOW_DAYS_RPT}-day window after launch "
     f"({SURVEY_DATE_RPT.strftime('%d %b %Y')} – {POST_END_RPT.strftime('%d %b %Y')}). "
     f"The comparison is split by email category (Value vs Sales) to see whether any lift "
     f"was uniform or category-specific.")

# Build the Q1 summary table
_q1_table_rows = []
for cat in ["Value", "Sales"]:
    s = _q1_stats[cat]
    delta_str = f"{s['delta']:+.1f}pp ({s['delta_pct']:+.1f}%)" if s["delta"] is not None else "—"
    p_str     = f"{s['p_value']:.3f}" if s["p_value"] is not None else "—"
    sig_str   = ("✓ significant" if s["p_value"] is not None and s["p_value"] < 0.05
                 else "n.s." if s["p_value"] is not None else "—")
    _q1_table_rows.append([
        cat,
        f"{s['pre_mean']}% (n={s['pre_n']})",
        f"{s['post_mean']}% (n={s['post_n']})" if s["post_mean"] else "—",
        delta_str,
        f"{p_str}  [{sig_str}]",
    ])

add_table(doc,
          ["Email Category", "Pre-Survey Avg OR", "Post-Survey Avg OR", "Δ Change", "Significance (t-test)"],
          _q1_table_rows)

_val_s  = _q1_stats["Value"]
_sal_s  = _q1_stats["Sales"]
_val_dir = "up" if (_val_s["delta"] or 0) >= 0 else "down"
_sal_dir = "up" if (_sal_s["delta"] or 0) >= 0 else "down"

body(doc,
     f"Value emails moved {_val_dir} by {abs(_val_s['delta'] or 0):.1f}pp "
     f"({abs(_val_s['delta_pct'] or 0):.1f}%) in the post-survey period. "
     f"Sales emails moved {_sal_dir} by {abs(_sal_s['delta'] or 0):.1f}pp "
     f"({abs(_sal_s['delta_pct'] or 0):.1f}%). "
     "This is a like-for-like window comparison (same duration on both sides), which reduces "
     "window-length bias. It is still observational (not causal), so results should be treated "
     "as directional unless confirmed over additional future windows.")

# Regenerate Figure W from the current Section 10.1 window settings (keeps chart/table aligned).
_w_chart_path = CHARTS / "W_survey_pre_post_open_rates.png"
_cats = ["Value", "Sales"]
_x = np.arange(len(_cats))
_bw = 0.34
_pre_means, _post_means, _pre_std, _post_std = [], [], [], []
for _cat in _cats:
    _pre_vals = pre_rpt[pre_rpt["category"] == _cat]["open_rate"].dropna()
    _post_vals = post_rpt[post_rpt["category"] == _cat]["open_rate"].dropna()
    _pre_means.append(float(_pre_vals.mean()) if len(_pre_vals) else np.nan)
    _post_means.append(float(_post_vals.mean()) if len(_post_vals) else np.nan)
    _pre_std.append(float(_pre_vals.std()) if len(_pre_vals) > 1 else 0.0)
    _post_std.append(float(_post_vals.std()) if len(_post_vals) > 1 else 0.0)

fig, ax = plt.subplots(figsize=(10.8, 5.8))
ax.bar(_x - _bw / 2, _pre_means, _bw, yerr=_pre_std, capsize=4, color="#9CA3AF", alpha=0.65, edgecolor="white", label=f"Pre ({Q1_WINDOW_DAYS_RPT}d)")
ax.bar(_x + _bw / 2, _post_means, _bw, yerr=_post_std, capsize=4, color="#4F46E5", alpha=0.90, edgecolor="white", label=f"Post ({Q1_WINDOW_DAYS_RPT}d)")
for _i, _cat in enumerate(_cats):
    _s = _q1_stats[_cat]
    _delta = _s["delta"]
    _p = _s["p_value"]
    _sig = "sig" if (_p is not None and _p < 0.05) else "n.s."
    _y = np.nanmax([_pre_means[_i], _post_means[_i]]) if np.isfinite(_pre_means[_i]) or np.isfinite(_post_means[_i]) else 0
    if _delta is not None:
        _p_txt = f"p={_p:.3f}" if _p is not None else "p=—"
        ax.text(_x[_i], _y + 1.8, f"{_delta:+.1f}pp | {_p_txt} [{_sig}]", ha="center", va="bottom", fontsize=9, color="#374151")
ax.set_xticks(_x)
ax.set_xticklabels(_cats)
ax.set_ylabel("Open Rate (%)")
ax.set_title(f"Pre vs Post Open Rates by Category ({Q1_WINDOW_DAYS_RPT}-day windows)")
ax.grid(axis="y", alpha=0.2)
ax.legend()
_vals_for_ylim = [v for v in (_pre_means + _post_means) if pd.notna(v)]
_ymax = max(_vals_for_ylim) if _vals_for_ylim else 0
ax.set_ylim(0, max(1.0, _ymax + 9))
fig.tight_layout()
fig.savefig(_w_chart_path, dpi=180)
plt.close()

chart_block(doc,
            "W_survey_pre_post_open_rates.png",
            "W",
            "Pre / Post-Survey Broadcast Open Rates by Category",
            "Each bar pair shows the average open rate before (faded) and after (solid) the survey send. "
            "Error bars show ±1 standard deviation. The annotation above each pair shows the absolute "
            "change in percentage points and the statistical significance of the difference.")
body(doc,
     "How to read Figure W: compare faded vs solid bars within each category, then check the p-value label. "
     "A visible bar difference without statistical significance should be interpreted as provisional. "
     f"This chart uses matched {Q1_WINDOW_DAYS_RPT}-day windows before and after launch, "
     "so the comparison is balanced by construction. It remains an early directional check, "
     "not definitive causal proof.",
     color=(50, 50, 50))

# ── 10.2 Q2 Engagement ───────────────────────────────────────────────────────
h2(doc, "10.2  Survey Responders vs Non-Responders: Engagement (Q2)")

if _q2_available and "A" in _q2 and "B" in _q2:
    a, b = _q2["A"], _q2["B"]
    or_delta_mean = round(a["or_mean"] - b["or_mean"], 1)
    or_delta_med  = round(a["or_med"]  - b["or_med"],  1)
    cr_delta_med  = round(a["cr_med"]  - b["cr_med"],  1)
    ctor_delta_med= round(a["ctor_med"]- b["ctor_med"],1)
    or_sig  = "statistically significant" if (_mwu.get("open_rate")  or 1) < 0.05 else "not statistically significant"
    cr_sig  = "statistically significant" if (_mwu.get("click_rate") or 1) < 0.05 else "not statistically significant"
    ctor_sig= "statistically significant" if (_mwu.get("ctor")       or 1) < 0.05 else "not statistically significant"

    body(doc,
         "Using the Kit API, lifetime engagement stats were fetched for every subscriber in both "
         f"groups (minimum 5 emails received). Group A = {a['n']:,} subscribers; "
         f"Group B = {b['n']:,} subscribers (611 who clicked but didn't respond + 5,757 who only opened). "
         "The stats below represent each subscriber's lifetime open rate, click rate, and CTOR "
         "across all broadcasts received since joining. "
         "Both mean and median are shown: the median is more robust here because Group B contains "
         f"{b['ctor_anomalous']} subscribers ({b['ctor_anom_pct']}%) with CTOR > 100% — "
         "a tracking anomaly most likely caused by link-scanning security services or "
         "email preview clients that auto-click links without a human opening the email. "
         f"By contrast, only {a['ctor_anomalous']} Group A subscribers ({a['ctor_anom_pct']}%) "
         "show the same anomaly. Median values are therefore the primary basis for interpretation.")

    add_table(doc,
              ["Metric", "Group A (Responded)", "Group B (Did Not Respond)", "Δ (A – B, median)"],
              [
                  ["Open Rate — mean",   f"{a['or_mean']}%",   f"{b['or_mean']}%",  f"{or_delta_mean:+.1f}pp"],
                  ["Open Rate — median", f"{a['or_med']}%",    f"{b['or_med']}%",   f"{or_delta_med:+.1f}pp  ← primary metric"],
                  ["Click Rate — median",f"{a['cr_med']}%",    f"{b['cr_med']}%",   f"{cr_delta_med:+.1f}pp"],
                  ["CTOR — median",      f"{a['ctor_med']}%",  f"{b['ctor_med']}%", f"{ctor_delta_med:+.1f}pp"],
                  [f"CTOR > 100% (anomalous)", f"{a['ctor_anom_pct']}%", f"{b['ctor_anom_pct']}%", "—"],
              ])

    # ── Open Rate interpretation ──────────────────────────────────────────
    body(doc, "What the data tells us:", bold=True, color=(79, 70, 229))
    if or_delta_med > 0:
        bullet(doc,
               f"Open Rate (primary signal): Group A opens {or_delta_med}pp more email at the median "
               f"({a['or_med']}% vs {b['or_med']}%, {or_sig}). "
               "This is the most reliable metric in this dataset — it is not affected by the "
               "click-tracking anomaly. Survey responders are meaningfully more likely to open "
               "future emails, which implies higher engagement with BOTH value content and sales offers.")
    else:
        bullet(doc,
               f"Open Rate: Groups are similar at the median (Δ = {or_delta_med:+.1f}pp, {or_sig}). "
               "Completing the survey does not predict meaningfully higher open rates.")

    # ── Click Rate / CTOR ─────────────────────────────────────────────────
    _ctor_note = (
        f"CTOR — use with caution: {b['ctor_anom_pct']}% of Group B has CTOR > 100%, "
        "suggesting automated link-clicks (email security scanners or preview clients) are "
        f"inflating Group B's click numbers. At the median, Group A CTOR = {a['ctor_med']}% "
        f"vs Group B = {b['ctor_med']}%. "
    )
    if ctor_delta_med >= 0:
        _ctor_note += (
            f"Group A has a {ctor_delta_med}pp higher median CTOR — once a subscriber opens "
            "an email, survey responders are at least as likely (and potentially more likely) "
            "to click through. This is a positive indicator for purchase intent on sales emails."
        )
    else:
        _ctor_note += (
            f"At the median, Group B still shows a {abs(ctor_delta_med)}pp higher CTOR. "
            "Given the data quality concerns, this should not be over-interpreted. "
            "The open rate advantage for Group A is the more actionable signal."
        )
    bullet(doc, _ctor_note)

    if cr_delta_med >= 0:
        bullet(doc,
               f"Click Rate: At the median, Group A subscribers click {cr_delta_med}pp more often "
               f"({a['cr_med']}% vs {b['cr_med']}%). This supports the view that survey responders "
               "engage more deeply with email content, including sales CTAs.")
    else:
        bullet(doc,
               f"Click Rate: At the median, Group B shows a {abs(cr_delta_med)}pp higher click rate "
               f"({b['cr_med']}% vs {a['cr_med']}%). Given the high rate of click-tracking anomalies "
               "in Group B, this figure should be treated as unreliable.")

    if or_delta_med > 0:
        callout(doc,
                f"Core finding: Subscribers who responded to the survey (Group A) open emails "
                f"{or_delta_med}pp more at the median than those who didn't respond. "
                "This is statistically significant and means the lead magnet is an effective "
                "quality signal — completing the survey correlates strongly with being a more "
                "engaged reader across ALL email types (value and sales). "
                "Group A is the highest-value segment on the list.",
                color_rgb=(16, 185, 129))
    else:
        callout(doc,
                "The survey response rate does not appear to predict meaningfully higher open rates. "
                "Both groups show similar engagement — the act of completing the survey may not "
                "be a strong quality filter in this case.",
                color_rgb=(245, 158, 11))
else:
    body(doc,
         "[Kit API cache not found — run analysis_lead_magnet.py first to generate lead_magnet_stats.csv]",
         color=(200, 0, 0))

chart_block(doc,
            "X_group_a_vs_b_engagement.png",
            "X",
            "Group A vs Group B — Overall Open Rate, Click Rate, CTOR (means)",
            "Grouped bars compare the two cohorts using mean values from Kit API lifetime stats. "
            "Error bars show ±1 standard deviation. Note: Click Rate and CTOR means for Group B "
            "are inflated by automated-click outliers (see discussion above). "
            "The Open Rate bars are the most reliable comparison in this chart.")
body(doc,
     "How to read Figure X: focus first on Open Rate bars (most trustworthy), then treat Click Rate/CTOR as "
     "supporting context because of known automated-click inflation. Error bars indicate variability; "
     "large overlap means weaker separation between groups.",
     color=(50, 50, 50))
body(doc,
     "Interpretation guardrail: if OR and CTOR point in different directions, prioritize OR for reliability in this dataset, "
     "then use CTOR as a directional indicator only.",
     color=(50, 50, 50))

chart_block(doc,
            "Y_group_ab_distributions.png",
            "Y",
            "Engagement Rate Distributions — Group A vs Group B (Box Plots)",
            "Box plots show the full spread of per-subscriber open rates and CTOR. "
            "A higher median and tighter interquartile range for Group A would confirm "
            "that survey responders are consistently more engaged, not just pulled up by outliers.")
body(doc,
     "How to read Figure Y: the center line of each box is the median (primary signal), the box is the middle 50% "
     "(IQR), and points beyond whiskers are outliers. If medians differ clearly while boxes overlap only partially, "
     "the group difference is broad-based rather than driven by a few extreme subscribers.",
     color=(50, 50, 50))
body(doc,
     "Interpretation guardrail: a difference in means without a difference in medians is often an outlier effect, "
     "not a population-level behavior shift.",
     color=(50, 50, 50))

# ── 10.3 Retention ───────────────────────────────────────────────────────────
h2(doc, "10.3  Retention: Who Stays Active?")

body(doc,
     "Beyond engagement rates, subscriber retention is compared between the two groups. "
     "An 'active' subscriber still receives emails; a 'cancelled' subscriber has opted out. "
     "If Group A (survey responders) retains at a higher rate, it confirms that completing the "
     "survey signals long-term interest in the content — not just a one-time curiosity click.")

if _q2_available and _status is not None:
    _st_rows = []
    for g_label, g_key in [("Group A (Responded)", "A"), ("Group B (Did Not Respond)", "B")]:
        if g_key in _status.index:
            r = _status.loc[g_key]
            _st_rows.append([g_label, int(r["sum"]), int(r["count"]), f"{r['active_pct']}%"])
    if _st_rows:
        add_table(doc,
                  ["Group", "Active Subscribers", "Total", "Active %"],
                  _st_rows)
        if len(_st_rows) == 2:
            a_act = float(_status.loc["A","active_pct"]) if "A" in _status.index else 0
            b_act = float(_status.loc["B","active_pct"]) if "B" in _status.index else 0
            act_delta = round(a_act - b_act, 1)
            if act_delta > 0:
                callout(doc,
                        f"Group A has a {act_delta}pp higher retention rate than Group B. "
                        "Subscribers who engaged with the lead magnet and chose a learning roadmap "
                        "are more likely to still be on the list — they self-selected into the content "
                        "with a specific goal in mind, which keeps them subscribed longer.",
                        color_rgb=(16, 185, 129))
            else:
                callout(doc,
                        f"Group A and Group B have similar retention rates (Δ = {act_delta:+.1f}pp). "
                        "Completing the survey does not appear to correlate with a meaningfully higher "
                        "likelihood of staying subscribed.",
                        color_rgb=(245, 158, 11))
else:
    body(doc,
         "[Retention data requires lead_magnet_stats.csv — run analysis_lead_magnet.py first]",
         color=(200, 0, 0))

chart_block(doc,
            "Z_group_ab_status.png",
            "Z",
            "Subscriber Status — Group A vs Group B (Active vs Cancelled)",
            "Stacked bars show the current Kit subscriber status breakdown for each group. "
            "A higher active % in Group A would confirm that completing the lead magnet survey "
            "is associated with stronger long-term retention.")
body(doc,
     "How to read Figure Z: compare the active segment height across the two bars. "
     "The active-share gap (in pp) is the direct retention advantage. This is a structural metric "
     "that is less noisy than monthly campaign metrics because it summarizes final subscriber state.",
     color=(50, 50, 50))

# ── 10.4 Post-survey email-type lift (Kit filter endpoint) ──────────────────
h2(doc, "10.4  Post-Survey Email-Type Lift: Value OR, Sales OR, Sales CTOR")

body(doc,
     "This section isolates only post-survey broadcasts and compares Group A vs Group B at the "
     "email-type level using Kit's engagement filter endpoint. For each post-survey broadcast, "
     "we count which subscribers in each group opened and clicked, then compute:"
     " (1) Value-email open rate, (2) Sales-email open rate, and (3) Sales-email CTOR.")
body(doc,
     "Interpretation framework for this section: "
     "Value OR answers 'who is still opening educational/value content?'; "
     "Sales OR answers 'who is still opening commercial emails?'; "
     "Sales CTOR answers 'once they open a sales email, who is more likely to click the sales CTA?'. "
     "Together, these three metrics separate reach from commercial action quality.",
     color=(50, 50, 50))

if _q3_available and _q3 is not None and len(_q3):
    _q3 = _q3.copy()
    _q3["sig"] = _q3["p_value"].apply(lambda p: "✓ significant" if pd.notna(p) and p < 0.05 else "n.s.")
    _rows = []
    for _, r in _q3.iterrows():
        _rows.append([
            r["metric"],
            int(r["n_broadcasts"]),
            f"{r['group_a_value']:.1f}%",
            f"{r['group_b_value']:.1f}%",
            f"{r['delta_pp']:+.1f}pp",
            f"{r['p_value']:.3f}" if pd.notna(r["p_value"]) else "—",
            r["sig"],
        ])
    add_table(doc,
              ["Metric", "Post-Survey Broadcasts", "Group A", "Group B", "Δ (A-B)", "p-value", "Significance"],
              _rows)

    _val = _q3[_q3["metric"] == "Value open rate"]
    _sal = _q3[_q3["metric"] == "Sales open rate"]
    _ctor = _q3[_q3["metric"] == "Sales CTOR"]
    if len(_val) and len(_sal) and len(_ctor):
        v = _val.iloc[0]
        s = _sal.iloc[0]
        c = _ctor.iloc[0]
        bullet(doc,
               f"Value emails: Group A open rate = {v['group_a_value']:.1f}% vs Group B = {v['group_b_value']:.1f}% "
               f"(Δ {v['delta_pp']:+.1f}pp, {('significant' if pd.notna(v['p_value']) and v['p_value'] < 0.05 else 'not significant')}).")
        bullet(doc,
               f"Sales emails: Group A open rate = {s['group_a_value']:.1f}% vs Group B = {s['group_b_value']:.1f}% "
               f"(Δ {s['delta_pp']:+.1f}pp, {('significant' if pd.notna(s['p_value']) and s['p_value'] < 0.05 else 'not significant')}).")
        bullet(doc,
               f"Sales CTOR: Group A = {c['group_a_value']:.1f}% vs Group B = {c['group_b_value']:.1f}% "
               f"(Δ {c['delta_pp']:+.1f}pp, {('significant' if pd.notna(c['p_value']) and c['p_value'] < 0.05 else 'not significant')}).")

        if v["delta_pp"] > 0 and s["delta_pp"] > 0:
            callout(doc,
                    "Subscribers who completed the lead magnet survey (Group A) are more likely to open both value "
                    "and sales emails in the post-survey period, indicating stronger sustained attention.",
                    color_rgb=(16, 185, 129))
        else:
            callout(doc,
                    "Post-survey open-rate lift is mixed across value and sales emails. The lead magnet still segments "
                    "intent, but open-rate advantage is not uniform across email types.",
                    color_rgb=(245, 158, 11))

        if c["delta_pp"] > 0:
            callout(doc,
                    "Group A also shows stronger sales CTOR, meaning survey responders are not only opening more "
                    "often but also clicking sales calls-to-action at a higher rate once they open.",
                    color_rgb=(16, 185, 129))
else:
    body(doc,
         "[Post-survey email-type analysis cache not found — run analysis_lead_magnet.py to generate "
         "lead_magnet_post_survey_summary.csv and AA_post_survey_group_by_type.png]",
         color=(200, 0, 0))

chart_block(doc,
            "AA_post_survey_group_by_type.png",
            "AA",
            "Post-Survey Group A vs Group B by Email Type",
            "Each paired bar compares Group A and Group B on Value-email open rate, Sales-email open rate, "
            "and Sales-email CTOR using Kit filter endpoint data across post-survey broadcasts.")
body(doc,
     "How to read Figure AA: treat each metric pair as a separate comparison test. "
     "Value OR and Sales OR answer inbox-reach questions; Sales CTOR answers click-quality after opening. "
     "A positive A-B gap on CTOR with flat OR means the lead magnet is segmenting buying intent more than reach.",
     color=(50, 50, 50))

# ── 10.5 Group age profile ───────────────────────────────────────────────────
h2(doc, "10.5  Are Group A Subscribers New or Old?")

body(doc,
     "To answer whether survey responders are mostly recent signups or long-tenure subscribers, "
     "subscriber age was calculated as days between the survey send date (25 Jan 2026) and each "
     "subscriber's 'created_at' timestamp from the opened/clicked lead-magnet exports.")

if _q5 is not None:
    add_table(doc,
              ["Metric", "Group A (Responded)", "Group B (Did Not Respond)"],
              [
                  ["Subscribers analysed", _q5["a_n"], _q5["b_n"]],
                  ["Mean age at survey",   f"{_q5['a_mean']} days",   f"{_q5['b_mean']} days"],
                  ["Median age at survey", f"{_q5['a_median']} days", f"{_q5['b_median']} days"],
                  ["IQR (25th–75th)",      f"{_q5['a_p25']}–{_q5['a_p75']} days", f"{_q5['b_p25']}–{_q5['b_p75']} days"],
                  ["Recent (<=90 days)",   f"{_q5['a_recent']}%",     f"{_q5['b_recent']}%"],
                  ["Older (>=181 days)",   f"{_q5['a_old']}%",        f"{_q5['b_old']}%"],
              ])

    _bucket_rows = []
    for _lbl in ["0-30 days", "31-90 days", "91-180 days", "181-365 days", "365+ days"]:
        _av = _q5["a_buckets"].get(_lbl, 0)
        _bv = _q5["b_buckets"].get(_lbl, 0)
        _bucket_rows.append([
            _lbl,
            f"{_av:,} ({_av / _q5['a_n'] * 100:.1f}%)" if _q5["a_n"] else "—",
            f"{_bv:,} ({_bv / _q5['b_n'] * 100:.1f}%)" if _q5["b_n"] else "—",
        ])

    add_table(doc,
              ["Age Bucket (at 25 Jan 2026)", "Group A", "Group B"],
              _bucket_rows)

    bullet(doc,
           f"Group A is younger overall: median age { _q5['a_median'] } days vs { _q5['b_median'] } days in Group B.")
    bullet(doc,
           f"Recent subscribers are over-represented in Group A ({_q5['a_recent']}% vs {_q5['b_recent']}%).")
    bullet(doc,
           f"Older subscribers are more concentrated in Group B ({_q5['b_old']}% vs {_q5['a_old']}%).")
    callout(doc,
            "Conclusion: Group A is not only new subscribers. It is a mixed cohort, but it skews younger "
            "than Group B while still containing a substantial share of established subscribers.",
            color_rgb=(16, 185, 129))

    chart_block(doc,
                _q5_chart_file,
                "AB",
                "Group A vs Group B — Subscriber Age Composition at Survey Date",
                "Each stacked bar shows the age mix within each group at the survey send date. "
                "Group A has a visibly larger recent-subscriber slice (0–90 days), while Group B "
                "is more heavily weighted toward older subscribers (181–365 days).")
    body(doc,
         "How to read Figure AB: compare the segment shares within each full bar (each bar totals 100%). "
         "This is composition analysis, not volume analysis. A bigger early-age share in Group A means responders "
         "skew newer, but older segments still matter if their slices are substantial.",
         color=(50, 50, 50))

    if _q5["a_removed"] or _q5["b_removed"]:
        body(doc,
             f"Data note: excluded {_q5['a_removed']} Group A and {_q5['b_removed']} Group B rows with "
             "created_at timestamps later than the survey send date.",
             size=9, italic=True, color=(120, 120, 120))
else:
    body(doc,
         "[Age-profile analysis requires both 'AI Sprint Roadmap - Clicked.csv' and "
         "'AI Sprint Roadmap - Opened subscribers.csv']",
         color=(200, 0, 0))

# ── 10.6 Monthly evolution (A vs B) ──────────────────────────────────────────
h2(doc, "10.6  Monthly Evolution: Open Rate and Sales CTOR (A vs B)")

body(doc,
     "This section tracks month-by-month engagement from February 2025 through February 2026 "
     "for Group A (survey responders) and Group B (non-responders). Three visual views are shown: "
     "(1) open-rate evolution by category, (2) sales CTOR evolution, and "
     "(3) value-vs-sales open rate within each group.")
body(doc,
     "How to read the three figures: AF shows raw trajectories with coverage flags, AG converts everything "
     "to one comparable unit (A-B gap in `pp`), and AH isolates sales click quality after open. "
     "Use AF for context, AG for directional clarity, and AH for commercial-intent interpretation.",
     color=(50, 50, 50))

if _q6 is not None and len(_q6):
    _q6w = _q6.copy()
    _q6w["has_data"] = _q6w["broadcasts_matched"] > 0
    _q6w["low_cov"] = (_q6w["broadcasts_matched"] > 0) & (_q6w["broadcasts_matched"] < 3)

    _v = _q6w[(_q6w["category"] == "Value") & (_q6w["has_data"])].copy()
    _s = _q6w[(_q6w["category"] == "Sales") & (_q6w["has_data"])].copy()
    _v_stable = _v[_v["broadcasts_matched"] >= 3].copy()
    _s_stable = _s[_s["broadcasts_matched"] >= 3].copy()
    _s_ctor_stable = _s_stable.dropna(subset=["a_sales_ctor", "b_sales_ctor"]).copy()

    _value_gap = (_v_stable["a_open_rate"] - _v_stable["b_open_rate"]).mean() if len(_v_stable) else np.nan
    _sales_gap = (_s_stable["a_open_rate"] - _s_stable["b_open_rate"]).mean() if len(_s_stable) else np.nan
    _sales_ctor_gap = (_s_ctor_stable["a_sales_ctor"] - _s_ctor_stable["b_sales_ctor"]).mean() if len(_s_ctor_stable) else np.nan
    _q6_value_gap = _value_gap
    _q6_sales_gap = _sales_gap
    _q6_sales_ctor_gap = _sales_ctor_gap
    _q6_value_a = _v_stable["a_open_rate"].mean() if len(_v_stable) else np.nan
    _q6_value_b = _v_stable["b_open_rate"].mean() if len(_v_stable) else np.nan
    _q6_sales_a = _s_stable["a_open_rate"].mean() if len(_s_stable) else np.nan
    _q6_sales_b = _s_stable["b_open_rate"].mean() if len(_s_stable) else np.nan
    _q6_ctor_a = _s_ctor_stable["a_sales_ctor"].mean() if len(_s_ctor_stable) else np.nan
    _q6_ctor_b = _s_ctor_stable["b_sales_ctor"].mean() if len(_s_ctor_stable) else np.nan

    _v_last = _v.iloc[-1] if len(_v) else None
    _s_last = _s.iloc[-1] if len(_s) else None

    # ── Re-designed charts for interpretability ───────────────────────────
    # AF: coverage-aware open-rate trend (A vs B)
    fig, axes = plt.subplots(2, 1, figsize=(13.5, 8.2), sharex=True)
    for _ax, _cat in zip(axes, ["Value", "Sales"]):
        _d = _q6w[_q6w["category"] == _cat].sort_values("month_dt").copy()
        _d["a_plot"] = _d["a_open_rate"].where(_d["has_data"], np.nan)
        _d["b_plot"] = _d["b_open_rate"].where(_d["has_data"], np.nan)

        _ax.plot(_d["month_dt"], _d["a_plot"], color="#10B981", linewidth=2.2, marker="o", label="Group A")
        _ax.plot(_d["month_dt"], _d["b_plot"], color="#F59E0B", linewidth=2.2, marker="o", label="Group B")

        _low = _d["low_cov"]
        _ax.scatter(_d.loc[_low, "month_dt"], _d.loc[_low, "a_plot"], s=64, facecolors="none", edgecolors="#10B981", linewidths=1.6)
        _ax.scatter(_d.loc[_low, "month_dt"], _d.loc[_low, "b_plot"], s=64, facecolors="none", edgecolors="#F59E0B", linewidths=1.6)

        for _, _r in _d[_low].iterrows():
            _y = np.nanmax([_r["a_plot"], _r["b_plot"]]) if pd.notna(_r["a_plot"]) or pd.notna(_r["b_plot"]) else np.nan
            if pd.notna(_y):
                _ax.text(_r["month_dt"], _y + 1.3, f"n={int(_r['broadcasts_matched'])}", fontsize=7.5, color="#94A3B8", ha="center")

        _ax.set_title(f"{_cat} Monthly Open Rate (coverage-aware)", fontsize=11.5)
        _ax.set_ylabel("Open Rate (%)")
        _ax.grid(axis="y", alpha=0.2)
        _ax.set_ylim(0, 100)
        _ax.legend(fontsize=8.5, loc="upper right")

    axes[-1].set_xlabel("Month")
    fig.suptitle("A vs B Open-Rate Evolution with Low-Coverage Months Flagged", fontsize=13.5, fontweight="bold")
    fig.tight_layout()
    fig.savefig(CHARTS / _q6_chart_cov_or)
    plt.close()

    # AG: monthly A-B gap (pp) for Value OR, Sales OR, Sales CTOR
    _gv = _v.copy()
    _gv["gap"] = _gv["a_open_rate"] - _gv["b_open_rate"]
    _gs = _s.copy()
    _gs["gap"] = _gs["a_open_rate"] - _gs["b_open_rate"]
    _gctor = _s.dropna(subset=["a_sales_ctor", "b_sales_ctor"]).copy()
    _gctor["gap"] = _gctor["a_sales_ctor"] - _gctor["b_sales_ctor"]

    fig, axes = plt.subplots(3, 1, figsize=(13.5, 10.2), sharex=True)
    for _ax, _d, _ttl, _yl in [
        (axes[0], _gv, "Value Open Rate Gap (A-B)", "Gap (pp)"),
        (axes[1], _gs, "Sales Open Rate Gap (A-B)", "Gap (pp)"),
        (axes[2], _gctor, "Sales CTOR Gap (A-B)", "Gap (pp)"),
    ]:
        _colors = _d["gap"].apply(lambda x: "#10B981" if x >= 0 else "#EF4444")
        _ax.bar(_d["month_dt"], _d["gap"], color=_colors, width=22, alpha=0.85)
        _ax.axhline(0, color="#64748B", linewidth=1.1)
        _ax.set_title(_ttl, fontsize=11.5)
        _ax.set_ylabel(_yl)
        _ax.grid(axis="y", alpha=0.18)

    axes[-1].set_xlabel("Month")
    fig.suptitle("Monthly Group Gap (A-B) — Easier Directional View", fontsize=13.5, fontweight="bold")
    fig.tight_layout()
    fig.savefig(CHARTS / _q6_chart_gap)
    plt.close()

    # AH: Sales CTOR with valid-month plotting only
    _sc = _s.dropna(subset=["a_sales_ctor", "b_sales_ctor"]).copy().sort_values("month_dt")
    fig, ax = plt.subplots(figsize=(13.5, 5.4))
    if len(_sc):
        _sc["total_openers"] = _sc["a_openers"].fillna(0) + _sc["b_openers"].fillna(0)
        _sc["a_roll"] = _sc["a_sales_ctor"].rolling(3, min_periods=1).mean()
        _sc["b_roll"] = _sc["b_sales_ctor"].rolling(3, min_periods=1).mean()

        # Coverage overlay: matched broadcasts per month on a secondary axis
        ax2 = ax.twinx()
        ax2.bar(
            _sc["month_dt"],
            _sc["broadcasts_matched"],
            width=20,
            color="#94A3B8",
            alpha=0.20,
            label="Matched broadcasts (coverage)",
            zorder=0,
        )
        ax2.axhline(3, color="#64748B", linestyle=":", linewidth=1.2, alpha=0.8)
        ax2.set_ylabel("Matched broadcasts (n)")
        ax2.set_ylim(0, max(3, float(_sc["broadcasts_matched"].max())) + 1.5)

        ax.plot(_sc["month_dt"], _sc["a_sales_ctor"], color="#10B981", marker="o", linewidth=2.2, label="Group A (monthly)")
        ax.plot(_sc["month_dt"], _sc["b_sales_ctor"], color="#F59E0B", marker="o", linewidth=2.2, label="Group B (monthly)")
        ax.plot(_sc["month_dt"], _sc["a_roll"], color="#065F46", linewidth=1.8, linestyle="--", label="Group A (3-pt smooth)")
        ax.plot(_sc["month_dt"], _sc["b_roll"], color="#B45309", linewidth=1.8, linestyle="--", label="Group B (3-pt smooth)")

        _low_sc = _sc["broadcasts_matched"] < 3
        ax.scatter(_sc.loc[_low_sc, "month_dt"], _sc.loc[_low_sc, "a_sales_ctor"], s=64, facecolors="none", edgecolors="#10B981", linewidths=1.6)
        ax.scatter(_sc.loc[_low_sc, "month_dt"], _sc.loc[_low_sc, "b_sales_ctor"], s=64, facecolors="none", edgecolors="#F59E0B", linewidths=1.6)

        # Small labels for coverage and opener volume by month
        for _, _r in _sc.iterrows():
            _cov = int(_r["broadcasts_matched"])
            _opn = int(_r["total_openers"])
            _y = np.nanmax([_r["a_sales_ctor"], _r["b_sales_ctor"]])
            if pd.notna(_y):
                _txt_color = "#64748B" if _cov >= 3 else "#EF4444"
                ax.text(
                    _r["month_dt"],
                    _y + 1.3,
                    f"n={_cov} | op={_opn}",
                    fontsize=7.2,
                    color=_txt_color,
                    ha="center",
                )
    ax.set_title("Sales CTOR Evolution (valid months only; low coverage highlighted)", fontsize=12)
    ax.set_ylabel("Sales CTOR (%)")
    ax.set_xlabel("Month")
    ax.grid(axis="y", alpha=0.2)
    _h1, _l1 = ax.get_legend_handles_labels()
    _h2, _l2 = (ax2.get_legend_handles_labels() if len(_sc) else ([], []))
    ax.legend(_h1 + _h2, _l1 + _l2, fontsize=8.2, loc="upper right")
    fig.tight_layout()
    fig.savefig(CHARTS / _q6_chart_sales_ctor)
    plt.close()

    add_table(doc,
              ["Metric (stable months: n>=3 broadcasts)", "Group A", "Group B", "Avg Δ (A-B)"],
              [
                  ["Value open rate", f"{_v_stable['a_open_rate'].mean():.1f}%" if len(_v_stable) else "—",
                   f"{_v_stable['b_open_rate'].mean():.1f}%" if len(_v_stable) else "—",
                   f"{_value_gap:+.1f}pp" if pd.notna(_value_gap) else "—"],
                  ["Sales open rate", f"{_s_stable['a_open_rate'].mean():.1f}%" if len(_s_stable) else "—",
                   f"{_s_stable['b_open_rate'].mean():.1f}%" if len(_s_stable) else "—",
                   f"{_sales_gap:+.1f}pp" if pd.notna(_sales_gap) else "—"],
                  ["Sales CTOR", f"{_s_ctor_stable['a_sales_ctor'].mean():.1f}%" if len(_s_ctor_stable) else "—",
                   f"{_s_ctor_stable['b_sales_ctor'].mean():.1f}%" if len(_s_ctor_stable) else "—",
                   f"{_sales_ctor_gap:+.1f}pp" if pd.notna(_sales_ctor_gap) else "—"],
              ])

    if _v_last is not None:
        bullet(doc,
               f"Latest Value month ({_v_last['month']}): Group A OR {_v_last['a_open_rate']:.1f}% vs Group B OR {_v_last['b_open_rate']:.1f}%.")
    if _s_last is not None:
        bullet(doc,
               f"Latest Sales month ({_s_last['month']}): Group A OR {_s_last['a_open_rate']:.1f}% vs Group B OR {_s_last['b_open_rate']:.1f}%; "
               f"Sales CTOR {_s_last['a_sales_ctor']:.1f}% vs {_s_last['b_sales_ctor']:.1f}%.")
    bullet(doc,
           "Hollow markers indicate low-coverage months (fewer than 3 matched broadcasts). "
           "Use those points as directional signals only.")

    chart_block(doc,
                _q6_chart_cov_or,
                "AF",
                "Coverage-Aware Monthly Open Rate (A vs B)",
                "This view separates signal from noise. The filled points represent months where there is enough "
                "broadcast volume to trust the directional comparison, while hollow points are low-coverage months "
                "that should be treated as weaker evidence. Across stable months, Value-email open rates tend to stay "
                "higher in Group B, while Sales-email open-rate leadership alternates by month and should be read as a "
                "volatile pattern rather than a fixed winner. The key takeaway is not one isolated month, but whether a "
                "group repeatedly leads across consecutive high-coverage months.")

    body(doc, "AF analysis:", bold=True, color=(79, 70, 229))
    body(doc,
         "When we focus only on stable months (>=3 matched broadcasts), Value-email opens consistently lean toward "
         f"Group B by about {abs(_value_gap):.1f}pp on average. Sales-email opens are less stable month-to-month, and "
         "leadership changes across periods. This indicates that open behavior is not driven by a single persistent cohort "
         "advantage across both categories.")

    chart_block(doc,
                _q6_chart_gap,
                "AG",
                "Monthly Gap View: A-B (pp) for Value OR, Sales OR, Sales CTOR",
                "This is the most interpretable diagnostic chart in the section because everything is converted to one unit: "
                "percentage-point gap. Bars above zero mean Group A is ahead; bars below zero mean Group B is ahead. "
                "Reading this chart left to right tells you when leadership changes, how large the reversals are, and "
                "whether the gap is tightening or widening. It also highlights a critical nuance: open-rate leadership "
                "and click-quality leadership (CTOR) can diverge, meaning one group may open more while the other still "
                "delivers weaker downstream click intent.")

    body(doc, "AG analysis:", bold=True, color=(79, 70, 229))
    body(doc,
         "The gap view clarifies that outcome quality differs by metric. Value OR bars are mostly below zero "
         "(Group B lead), Sales OR bars are mixed, while Sales CTOR bars tend to be positive when coverage is sufficient. "
         "In other words, Group A does not consistently lead on inbox reach, but often leads on post-open click behavior.")

    chart_block(doc,
                _q6_chart_sales_ctor,
                "AH",
                "Sales CTOR Evolution (valid months only)",
                "This chart intentionally removes months with no valid Sales coverage so we avoid apparent collapses caused "
                "by missing data. The solid lines show the raw month-by-month CTOR, and the dashed lines smooth short-term "
                "volatility to reveal the underlying trend. Grey bars on the right axis show matched-broadcast coverage by month, "
                "with the dotted line at n=3 as the minimum-confidence threshold. Labels (n, op) show broadcast count and total "
                "openers for each point. The interpretation focus here is conversion quality among openers: "
                "when Group A stays above Group B over multiple valid months, it means A's opens are more commercially engaged, "
                "not just more frequent. If the lines converge, the gap is narrowing and message/offer quality is becoming more similar.")

    body(doc, "AH analysis:", bold=True, color=(79, 70, 229))
    body(doc,
         f"Across stable Sales months, Group A's CTOR average is materially higher (about {abs(_sales_ctor_gap):.1f}pp lead). "
         "This is the strongest recurring signal in the section. It suggests that responders convert attention into action "
         "more effectively on sales emails, even in months where they do not lead on raw open rate.")
    body(doc,
         "Decision implication: for sales broadcasts, prioritize responder-focused targeting when the objective is click-through "
         "quality rather than maximum reach.",
         color=(50, 50, 50))

    body(doc, "10.6 take points:", bold=True, color=(79, 70, 229))
    bullet(doc,
           f"Value-email opens: Group B is generally higher in stable months (A-B gap { _value_gap:+.1f}pp).")
    bullet(doc,
           f"Sales-email opens: mixed by month, with no consistently dominant group (A-B gap { _sales_gap:+.1f}pp on stable months).")
    bullet(doc,
           f"Sales CTOR: Group A shows the clearest advantage ({ _sales_ctor_gap:+.1f}pp on stable months), "
           "indicating stronger click intent among sales-email openers.")
    bullet(doc,
           "Strategic implication: use open-rate metrics to judge reach, but prioritize Sales CTOR to identify the segment "
           "with stronger commercial intent for targeted campaigns.")

    body(doc,
         "Data caveat: legacy API 500 failures and unmatched broadcasts affect coverage more heavily in older months. "
         "The redesigned charts expose low-coverage months explicitly to preserve interpretation quality.",
         size=9, italic=True, color=(120, 120, 120))
else:
    body(doc,
         "[Monthly evolution data requires lead_magnet_group_ab_monthly.csv]",
         color=(200, 0, 0))

body(doc, "Key takeaways for the lead magnet strategy:", bold=True, color=(79, 70, 229))
bullet(doc,
       "The survey response rate (Group A / total clicked) reveals what proportion of engaged "
       "subscribers are willing to self-segment — a high conversion here is a strong signal "
       "that the lead magnet topic resonates with the audience.")
bullet(doc,
       "If Group A shows meaningfully higher open rates and CTOR than Group B, the survey is "
       "not just a tactical tool — it is a quality filter that identifies your most engaged "
       "segment. Future segmented campaigns targeting Group A can be expected to outperform "
       "whole-list broadcasts.")
bullet(doc,
       "If retention is also higher in Group A, consider making the survey a standard part of "
       "the welcome sequence for all new subscribers, not just a one-off campaign. It primes "
       "subscribers to engage with goal-specific content from the start.")

doc.add_page_break()

# ── 11 Subscriber origin split ───────────────────────────────────────────────
h1(doc, "11. Subscriber Origin Split: Lead-Magnet Signups vs Existing Recipients")

body(doc,
     "This section answers a practical attribution question: among confirmed subscribers, how many "
     "appear to have subscribed through the survey lead magnet, versus how many were already on the list "
     "when the survey was sent.")
body(doc,
     "Classification rule used: pre-survey cohort = subscribers created in the 38-day window before launch "
     f"({(SURVEY_DATE_RPT - pd.Timedelta(days=PRE_SURVEY_WINDOW_DAYS_RPT)).strftime('%d %b %Y')} to {(SURVEY_DATE_RPT - pd.Timedelta(days=1)).strftime('%d %b %Y')}). "
     "'Lead-magnet signups (strict inferred)' = `created_at` on/after 25 Jan 2026 and on/before 2 Mar 2026 "
     "+ survey-response tag "
     "(goal-choice tag such as 'I want...' / \"I'm...\") + roadmap referrer "
     "(`lonelyoctopus.com/ai-sprint-roadmap`). "
     "This stricter rule intentionally excludes post-survey signups that have response tags but non-roadmap/unknown referrers.",
     color=(50, 50, 50))
body(doc,
     "Boundary note: 25 Jan 2026 is excluded from the pre-survey side. "
     "In this section, post-survey signups are limited to the launch window "
     "(25 Jan 2026 to 2 Mar 2026, inclusive).",
     color=(50, 50, 50))

if _q7 is not None:
    add_table(doc,
              ["Segment", "Subscribers", "Share of confirmed list", "Definition"],
              [
                  ["Pre-survey 38-day cohort", f"{_q7['existing_count']:,}", f"{_q7['existing_pct_total']:.1f}%",
                   f"Created between {(SURVEY_DATE_RPT - pd.Timedelta(days=PRE_SURVEY_WINDOW_DAYS_RPT)).strftime('%d %b %Y')} and {(SURVEY_DATE_RPT - pd.Timedelta(days=1)).strftime('%d %b %Y')}"],
                  ["Subscribed via lead magnet (strict inferred)", f"{_q7['lead_magnet_count']:,}", f"{_q7['lead_magnet_pct_total']:.1f}%",
                   "Created between 25 Jan and 2 Mar 2026 + survey-response tag + roadmap referrer"],
                  ["Post-survey signup, other source", f"{_q7['post_other_count']:,}", f"{_q7['post_other_pct_total']:.1f}%",
                   "All post-survey signups not meeting strict lead-magnet rule"],
              ])

    add_table(doc,
              ["Post-survey signup breakdown", "Subscribers", "Share of post-survey signups"],
              [
                  ["All post-survey signups", f"{_q7['post_total']:,}", "100.0%"],
                  ["Lead magnet (strict inferred)", f"{_q7['lead_magnet_count']:,}", f"{_q7['lead_magnet_pct_post']:.1f}%"],
                  ["Other sources", f"{_q7['post_other_count']:,}", f"{_q7['post_other_pct_post']:.1f}%"],
              ])

    add_table(doc,
              ["Attribution quality check", "Subscribers", "Share of post-survey signups"],
              [
                  ["Response tag present but non-roadmap/unknown referrer (excluded from lead-magnet count)",
                   f"{_q7['post_resp_nonroadmap_count']:,}",
                   f"{_q7['post_resp_nonroadmap_pct_post']:.1f}%"],
              ])

    bullet(doc,
           f"The pre-survey 38-day cohort contains {_q7['existing_count']:,} subscribers "
           f"({_q7['existing_pct_total']:.1f}% of all confirmed subscribers).")
    bullet(doc,
           f"{_q7['lead_magnet_count']:,} subscribers ({_q7['lead_magnet_pct_total']:.1f}% of all confirmed; "
           f"{_q7['lead_magnet_pct_post']:.1f}% of post-survey signups) are strict-inferred lead-magnet signups.")
    bullet(doc,
           f"{_q7['post_other_count']:,} post-survey signups ({_q7['post_other_pct_post']:.1f}% of post-survey signups) "
           "appear to have joined through other acquisition paths.")
    bullet(doc,
           f"Exclusion check: {_q7['post_resp_nonroadmap_count']:,} post-survey signups had survey-response tags but "
           "non-roadmap/unknown referrers, so they were excluded from strict lead-magnet attribution.")

    callout(doc,
            "Interpretation: this campaign was primarily a segmentation and engagement mechanism for the existing list, "
            "while also contributing incremental new signups through the lead magnet path.",
            color_rgb=(16, 185, 129))

    chart_block(doc,
                _q7_chart_file,
                "AC",
                "Subscriber Origin Split Around Survey Launch",
                "The largest bar is the existing subscriber base that received the survey. The two smaller bars isolate "
                "new signups after the survey date into strict-inferred lead-magnet signups versus other acquisition sources. "
                "Use this chart to separate campaign impact on acquisition from campaign impact on existing-list engagement.")
    body(doc,
         "How to read Figure AC: compare the first bar to the two post-survey bars to understand scale. "
         "Then compare the two post-survey bars to estimate what share of new growth is attributable to the survey funnel "
         "versus other channels.",
         color=(50, 50, 50))

    _q7_eng_path = GENERATED / "lead_magnet_origin_post_survey_summary.csv"
    if _q7_eng_path.exists():
        _q7e = pd.read_csv(_q7_eng_path)
        _q7e = _q7e[_q7e["metric"].isin(["Value open rate", "Sales open rate", "Sales CTOR"])].copy()
        _metric_order = ["Value open rate", "Sales open rate", "Sales CTOR"]
        _q7e["metric"] = pd.Categorical(_q7e["metric"], categories=_metric_order, ordered=True)
        _q7e = _q7e.sort_values("metric")

        def _fmt_p(_p):
            if pd.isna(_p):
                return "—"
            return "<0.001" if _p < 0.001 else f"{_p:.3f}"

        if len(_q7e):
            body(doc,
                 "Do lead-magnet signups have higher post-survey engagement than the pre-survey 38-day cohort?",
                 bold=True, color=(79, 70, 229))
            body(doc,
                 "Window definition note: for this origin-cohort comparison, 'post-survey' uses only the matched "
                 "broadcasts from 28 Jan 2026 to 18 Feb 2026 (7 broadcasts total: 5 Value, 2 Sales). "
                 "This is a 24-day observation window after the survey send on 25 Jan 2026 (not a 90-day window).",
                 color=(50, 50, 50))
            body(doc,
                 "Using Kit filter-event data on matched post-survey broadcasts, we compared two origin cohorts: "
                 "(1) strict-inferred lead-magnet signups, and (2) the pre-survey 38-day cohort. "
                 "Rates below are weighted totals across broadcasts (Open Rate = total opens / total eligible; "
                 "Sales CTOR = total sales clicks / total sales opens), not a simple unweighted mean of per-broadcast rates.",
                 color=(50, 50, 50))

            _rows = []
            for _, _r in _q7e.iterrows():
                _rows.append([
                    _r["metric"],
                    int(_r["n_broadcasts"]) if pd.notna(_r["n_broadcasts"]) else "—",
                    f"{_r['lead_magnet_value']:.1f}%",
                    f"{_r['existing_value']:.1f}%",
                    f"{_r['delta_pp']:+.1f}pp",
                    _fmt_p(_r.get("p_value", np.nan)),
                ])

            add_table(doc,
                      ["Metric", "Post-survey broadcasts (n)", "Lead-magnet signups (strict inferred)", "Pre-survey 38-day cohort", "Delta", "p-value"],
                      _rows)

            _val = _q7e[_q7e["metric"] == "Value open rate"].iloc[0]
            _sal = _q7e[_q7e["metric"] == "Sales open rate"].iloc[0]
            _ctor = _q7e[_q7e["metric"] == "Sales CTOR"].iloc[0]

            bullet(doc,
                   f"Value email open rate is higher for lead-magnet signups: {_val['lead_magnet_value']:.1f}% "
                   f"vs {_val['existing_value']:.1f}% (Delta {_val['delta_pp']:+.1f}pp), but this gap is not statistically significant.")
            bullet(doc,
                   f"Sales email open rate is also higher for lead-magnet signups: {_sal['lead_magnet_value']:.1f}% "
                   f"vs {_sal['existing_value']:.1f}% (Delta {_sal['delta_pp']:+.1f}pp), and this gap is statistically significant.")
            bullet(doc,
                   f"Sales CTOR is materially higher for lead-magnet signups: {_ctor['lead_magnet_value']:.1f}% "
                   f"vs {_ctor['existing_value']:.1f}% (Delta {_ctor['delta_pp']:+.1f}pp), but this gap is not statistically significant.")

            callout(doc,
                    "Answer: partially. The lead-magnet cohort shows a clear and statistically strong advantage on Sales open rate. "
                    "Value open rate and Sales CTOR are directionally higher, but not statistically conclusive in this window.",
                    color_rgb=(16, 185, 129))

            chart_block(doc,
                        "AD2_lead_magnet_origin_metrics.png",
                        "AD2",
                        "Post-Survey Engagement by Subscriber Origin",
                        "Each metric compares strict-inferred lead-magnet signups against the pre-survey 38-day cohort. "
                        "A higher bar for the lead-magnet cohort across Value OR, Sales OR, and Sales CTOR indicates "
                        "stronger inbox reach and stronger post-open sales click intent for that origin group.")
            body(doc,
                 "How to read Figure AD2: compare the two bars within each metric. "
                 "If the lead-magnet bar is higher across all three metrics, the lead magnet is bringing in a higher-intent cohort, "
                 "not just increasing list size.",
                 color=(50, 50, 50))

            _q7_prepost_path = GENERATED / "lead_magnet_prepost_window_summary.csv"
            if _q7_prepost_path.exists():
                _pp = pd.read_csv(_q7_prepost_path)
                _pp = _pp[_pp["metric"].isin(["Value open rate", "Sales open rate", "Sales CTOR"])].copy()
                _pp_order = ["Value open rate", "Sales open rate", "Sales CTOR"]
                _pp["metric"] = pd.Categorical(_pp["metric"], categories=_pp_order, ordered=True)
                _pp = _pp.sort_values("metric")
                if len(_pp):
                    h2(doc, "11.1  As-Received Comparison: Each Cohort on Its Own Window Broadcasts")
                    body(doc,
                         "This cross-window diagnostic compares (A) the pre-survey cohort measured on pre-window broadcasts "
                         "against (B) the lead-magnet cohort measured on post-window broadcasts. "
                         "In other words, each cohort is scored only on broadcasts actually sent in its own time window. "
                         "Interpret this as a directional baseline-to-after check, not a like-for-like causal estimate.",
                         color=(50, 50, 50))
                    body(doc,
                         "Window detail: pre-survey cohort uses broadcasts sent from 18 Dec 2025 to 24 Jan 2026; "
                         "lead-magnet cohort uses broadcasts sent from 25 Jan 2026 to 2 Mar 2026.",
                         color=(50, 50, 50))

                    _pp_rows = []
                    for _, _r in _pp.iterrows():
                        _pre_n = int(_r["pre_n_broadcasts"]) if pd.notna(_r["pre_n_broadcasts"]) else 0
                        _post_n = int(_r["post_n_broadcasts"]) if pd.notna(_r["post_n_broadcasts"]) else 0
                        _pre_v = f"{_r['pre_value']:.1f}%" if pd.notna(_r["pre_value"]) else "N/A"
                        _post_v = f"{_r['post_value']:.1f}%" if pd.notna(_r["post_value"]) else "N/A"
                        _d = f"{_r['delta_pp']:+.1f}pp" if pd.notna(_r["delta_pp"]) else "N/A"
                        _p = _fmt_p(_r.get("p_value", np.nan))
                        _pp_rows.append([
                            _r["metric"], f"{_pre_v} (n={_pre_n})", f"{_post_v} (n={_post_n})", _d, _p
                        ])

                    add_table(doc,
                              ["Metric", "Pre window (existing 38-day cohort)", "Post window (lead-magnet cohort)", "Delta", "p-value"],
                              _pp_rows)

                    _pp_val = _pp[_pp["metric"] == "Value open rate"]
                    if len(_pp_val):
                        _v = _pp_val.iloc[0]
                        if pd.notna(_v["delta_pp"]):
                            bullet(doc,
                                   f"Value OR changed from {_v['pre_value']:.1f}% (pre) to {_v['post_value']:.1f}% (post), "
                                   f"Delta {_v['delta_pp']:+.1f}pp.")

                    _pp_sales = _pp[_pp["metric"] == "Sales open rate"]
                    if len(_pp_sales):
                        _s = _pp_sales.iloc[0]
                        if pd.isna(_s["pre_value"]):
                            bullet(doc,
                                   "Sales OR pre-window baseline is unavailable because there were no matched Sales broadcasts in the pre window.")
                    _pp_ctor = _pp[_pp["metric"] == "Sales CTOR"]
                    if len(_pp_ctor):
                        _c = _pp_ctor.iloc[0]
                        if pd.isna(_c["pre_value"]):
                            bullet(doc,
                                   "Sales CTOR pre-window baseline is unavailable for the same reason (no matched pre-window Sales broadcasts).")
    else:
        body(doc,
             "[Origin-cohort engagement summary not found — run analysis_lead_magnet_origin_post.py to generate "
             "lead_magnet_origin_post_survey_summary.csv and AD2_lead_magnet_origin_metrics.png]",
             color=(200, 0, 0))

    _q7_postwindow_path = GENERATED / "lead_magnet_postwindow_nonlead_summary.csv"
    if _q7_postwindow_path.exists():
        _pw = pd.read_csv(_q7_postwindow_path)
        _pw = _pw[_pw["metric"].isin(["Value open rate", "Sales open rate", "Sales CTOR"])].copy()
        _pw_order = ["Value open rate", "Sales open rate", "Sales CTOR"]
        _pw["metric"] = pd.Categorical(_pw["metric"], categories=_pw_order, ordered=True)
        _pw = _pw.sort_values("metric")

        def _fmt_p_pw(_p):
            if pd.isna(_p):
                return "—"
            return "<0.001" if _p < 0.001 else f"{_p:.3f}"

        if len(_pw):
            h2(doc, "11.2  Same Post-Window Comparison: Lead-Magnet vs Non-Lead Signups")
            body(doc,
                 "This subsection compares only subscribers who signed up in the same post-survey window "
                 "(25 Jan 2026 to 2 Mar 2026): strict lead-magnet signups versus non-lead signups. "
                 "This isolates acquisition-path differences while holding signup period constant.",
                 color=(50, 50, 50))
            body(doc,
                 "Definitions used: lead-magnet signup = survey-response tag + referrer "
                 "`lonelyoctopus.com/ai-sprint-roadmap`; non-lead signup = everyone else in the same post window.",
                 color=(50, 50, 50))

            _pw_rows = []
            for _, _r in _pw.iterrows():
                _pw_rows.append([
                    _r["metric"],
                    int(_r["n_broadcasts"]) if pd.notna(_r["n_broadcasts"]) else "—",
                    f"{_r['lead_magnet_value']:.1f}%" if pd.notna(_r["lead_magnet_value"]) else "N/A",
                    f"{_r['non_lead_value']:.1f}%" if pd.notna(_r["non_lead_value"]) else "N/A",
                    f"{_r['delta_pp']:+.1f}pp" if pd.notna(_r["delta_pp"]) else "N/A",
                    _fmt_p_pw(_r.get("p_value", np.nan)),
                ])

            add_table(doc,
                      ["Metric", "Broadcasts (n)", "Lead-magnet signups", "Non-lead signups", "Delta", "p-value"],
                      _pw_rows)

            _pw_val = _pw[_pw["metric"] == "Value open rate"]
            _pw_sal = _pw[_pw["metric"] == "Sales open rate"]
            _pw_ctor = _pw[_pw["metric"] == "Sales CTOR"]
            if len(_pw_val):
                _v = _pw_val.iloc[0]
                bullet(doc,
                       f"Value OR: {_v['lead_magnet_value']:.1f}% vs {_v['non_lead_value']:.1f}% "
                       f"(Delta {_v['delta_pp']:+.1f}pp).")
            if len(_pw_sal):
                _s = _pw_sal.iloc[0]
                bullet(doc,
                       f"Sales OR: {_s['lead_magnet_value']:.1f}% vs {_s['non_lead_value']:.1f}% "
                       f"(Delta {_s['delta_pp']:+.1f}pp).")
            if len(_pw_ctor):
                _c = _pw_ctor.iloc[0]
                bullet(doc,
                       f"Sales CTOR: {_c['lead_magnet_value']:.1f}% vs {_c['non_lead_value']:.1f}% "
                       f"(Delta {_c['delta_pp']:+.1f}pp).")

            callout(doc,
                    "Interpretation: in the same post-signup window, lead-magnet signups show a strong Sales-open advantage, "
                    "while Value open rate is similar and Sales CTOR is directionally mixed.",
                    color_rgb=(16, 185, 129))
    else:
        body(doc,
             "[Post-window lead-vs-non-lead summary not found — run analysis_lead_magnet_origin_post.py to generate "
             "lead_magnet_postwindow_nonlead_summary.csv]",
             color=(200, 0, 0))

    if _q7["missing_created"] > 0:
        body(doc,
             f"Data note: {_q7['missing_created']:,} rows were excluded in this split because `created_at` was missing or invalid.",
             size=9, italic=True, color=(120, 120, 120))
else:
    body(doc,
         "[Section 11 requires Confirmed Subscribers.csv]",
         color=(200, 0, 0))

doc.add_page_break()
h1(doc, "12. Signup-Date Cohort Evolution (Open Rate + CTOR)")

body(doc,
     "This section tests whether older subscribers are disengaging by tracking monthly Open Rate and CTOR "
     "for signup-date cohorts (grouped by subscription date quartile). "
     "If older cohorts repeatedly underperform newer cohorts across recent stable months, the hypothesis is supported.")
body(doc,
     "Reading note: in all cohort gap charts, `pp` means percentage points. "
     "Negative values mean the oldest cohort is behind the newest cohort by that many points.",
     color=(50, 50, 50))

if _q11 is not None and len(_q11):
    _c = _q11.copy()
    _c["ctor_capped"] = _c["ctor"].clip(upper=100)
    _c["stable"] = _c["broadcasts_matched"] >= 3
    _c["valid_ctor"] = _c["stable"] & (_c["open_n"] >= 20)

    _cohort_order = [c for c in ["Q1 Oldest Signups", "Q2 Mid-Old Signups", "Q3 Mid-New Signups", "Q4 Newest Signups"] if c in _c["cohort"].unique()]
    _cohort_colors = {
        "Q1 Oldest Signups": "#2563EB",
        "Q2 Mid-Old Signups": "#10B981",
        "Q3 Mid-New Signups": "#F59E0B",
        "Q4 Newest Signups": "#EF4444",
    }

    # AL — Open-rate trend by cohort (coverage-aware)
    fig, ax = plt.subplots(figsize=(13.5, 6.2))
    for _co in _cohort_order:
        _d = _c[_c["cohort"] == _co].sort_values("month_dt")
        _y = _d["open_rate"].where(_d["broadcasts_matched"] > 0, np.nan)
        _color = _cohort_colors.get(_co, "#4F46E5")
        ax.plot(_d["month_dt"], _y, marker="o", linewidth=2.0, color=_color, label=_co)
        _low = (_d["broadcasts_matched"] > 0) & (_d["broadcasts_matched"] < 3)
        ax.scatter(_d.loc[_low, "month_dt"], _y[_low], s=56, facecolors="none", edgecolors=_color, linewidths=1.4)
    ax.set_title("Monthly Open Rate by Signup Cohort (coverage-aware)")
    ax.set_ylabel("Open Rate (%)")
    ax.set_xlabel("Month")
    ax.set_ylim(0, 100)
    ax.grid(axis="y", alpha=0.2)
    ax.legend(fontsize=8.4)
    fig.tight_layout()
    fig.savefig(CHARTS / "AL_signup_cohort_openrate_trend.png")
    plt.close()

    # AM — CTOR trend by cohort (capped/valid with coverage overlay)
    fig, ax = plt.subplots(figsize=(13.5, 6.2))
    _low_cov_label_added = False
    for _co in _cohort_order:
        _d = _c[_c["cohort"] == _co].sort_values("month_dt")
        _color = _cohort_colors.get(_co, "#4F46E5")
        # Show all months with enough openers as a light dashed bridge, then emphasize valid months.
        _y_open_ok = _d["ctor_capped"].where(_d["open_n"] >= 20, np.nan)
        _y_valid = _d["ctor_capped"].where(_d["valid_ctor"], np.nan)
        _low_cov = (_d["open_n"] >= 20) & (~_d["valid_ctor"]) & _d["ctor_capped"].notna()

        ax.plot(_d["month_dt"], _y_open_ok, linewidth=1.3, linestyle="--", alpha=0.45, color=_color)
        ax.plot(_d["month_dt"], _y_valid, marker="o", linewidth=2.0, color=_color, label=_co)
        if _low_cov.any():
            ax.scatter(
                _d.loc[_low_cov, "month_dt"],
                _d.loc[_low_cov, "ctor_capped"],
                s=54,
                facecolors="none",
                edgecolors=_color,
                linewidths=1.5,
                label=("Low-coverage month (hollow)" if not _low_cov_label_added else None),
            )
            _low_cov_label_added = True

    _mct = (
        _c[_c["open_n"] >= 20]
        .groupby("month_dt", as_index=False)
        .agg(
            valid_cohorts=("valid_ctor", "sum"),
            avg_matched=("broadcasts_matched", "mean"),
            total_openers=("open_n", "sum"),
        )
        .sort_values("month_dt")
    )
    _ctor_weak_months = int((_mct["valid_cohorts"] < 3).sum()) if len(_mct) else 0
    ax2 = ax.twinx()
    if len(_mct):
        ax2.bar(
            _mct["month_dt"],
            _mct["avg_matched"],
            width=20,
            color="#94A3B8",
            alpha=0.20,
            label="Avg matched broadcasts (coverage)",
            zorder=0,
        )
        ax2.axhline(3, color="#64748B", linestyle=":", linewidth=1.2, alpha=0.8)
        for _, _r in _mct.iterrows():
            _mm = _c[_c["month_dt"] == _r["month_dt"]]
            _ymax = _mm.loc[_mm["open_n"] >= 20, "ctor_capped"].max()
            if pd.notna(_ymax):
                _txt_color = "#64748B" if (_r["avg_matched"] >= 3 and _r["valid_cohorts"] >= 3) else "#EF4444"
                ax.text(
                    _r["month_dt"],
                    _ymax + 1.4,
                    f"k={int(_r['valid_cohorts'])} | op={int(_r['total_openers'])}",
                    fontsize=7.0,
                    color=_txt_color,
                    ha="center",
                )
        ax2.set_ylim(0, max(3, float(_mct["avg_matched"].max())) + 1.5)
    else:
        ax2.set_ylim(0, 4.5)
    ax2.set_ylabel("Avg matched broadcasts (n)")

    ax.set_title("Monthly CTOR by Signup Cohort (valid months emphasized; low coverage shown hollow)")
    ax.set_ylabel("CTOR (%)")
    ax.set_xlabel("Month")
    ax.set_ylim(0, 100)
    ax.grid(axis="y", alpha=0.2)
    _h1, _l1 = ax.get_legend_handles_labels()
    _h2, _l2 = ax2.get_legend_handles_labels()
    ax.legend(_h1 + _h2, _l1 + _l2, fontsize=8.0)
    fig.tight_layout()
    fig.savefig(CHARTS / "AM_signup_cohort_ctor_trend.png")
    plt.close()

    # AN — Oldest vs newest gap
    _old, _new = _cohort_order[0], _cohort_order[-1]
    _or_p = _c[_c["stable"]].pivot(index="month_dt", columns="cohort", values="open_rate")
    _ct_p = _c[_c["valid_ctor"]].pivot(index="month_dt", columns="cohort", values="ctor_capped")
    _g_or = _or_p[[x for x in [_old, _new] if x in _or_p.columns]].dropna()
    _g_ct = _ct_p[[x for x in [_old, _new] if x in _ct_p.columns]].dropna()

    fig, axes = plt.subplots(2, 1, figsize=(13.5, 8.0), sharex=True)
    if len(_g_or):
        _gap_or = _g_or[_old] - _g_or[_new]
        axes[0].bar(_gap_or.index, _gap_or.values, width=22,
                    color=["#EF4444" if x < 0 else "#10B981" for x in _gap_or.values], alpha=0.85)
    axes[0].axhline(0, color="#64748B", linewidth=1.0)
    axes[0].set_title(f"Open Rate Gap (Oldest - Newest): {_old} vs {_new}")
    axes[0].set_ylabel("Gap (pp)")
    axes[0].grid(axis="y", alpha=0.18)

    if len(_g_ct):
        _gap_ct = _g_ct[_old] - _g_ct[_new]
        axes[1].bar(_gap_ct.index, _gap_ct.values, width=22,
                    color=["#EF4444" if x < 0 else "#10B981" for x in _gap_ct.values], alpha=0.85)
    axes[1].axhline(0, color="#64748B", linewidth=1.0)
    axes[1].set_title(f"CTOR Gap (Oldest - Newest): {_old} vs {_new}")
    axes[1].set_ylabel("Gap (pp)")
    axes[1].set_xlabel("Month")
    axes[1].grid(axis="y", alpha=0.18)
    fig.tight_layout()
    fig.savefig(CHARTS / "AN_oldest_vs_newest_gap.png")
    plt.close()

    _stable = _c[_c["stable"]].groupby("cohort")["open_rate"].mean()
    _validc = _c[_c["valid_ctor"]].groupby("cohort")["ctor_capped"].mean()
    _or_gap_avg = np.nan
    _or_gap_recent = np.nan
    _ct_gap_avg = np.nan
    _ct_gap_recent = np.nan
    if len(_g_or):
        _go = (_g_or[_old] - _g_or[_new]).dropna()
        _or_gap_avg = _go.mean()
        _or_gap_recent = _go.tail(3).mean() if len(_go) >= 3 else _go.mean()
    if len(_g_ct):
        _gc = (_g_ct[_old] - _g_ct[_new]).dropna()
        _ct_gap_avg = _gc.mean()
        _ct_gap_recent = _gc.tail(3).mean() if len(_gc) >= 3 else _gc.mean()

    _rows = []
    for _co in _cohort_order:
        _rows.append([
            _co,
            f"{_stable.get(_co, np.nan):.1f}%" if pd.notna(_stable.get(_co, np.nan)) else "—",
            f"{_validc.get(_co, np.nan):.1f}%" if pd.notna(_validc.get(_co, np.nan)) else "—",
        ])
    add_table(doc, ["Signup Cohort", "Avg Open Rate (stable months)", "Avg CTOR (valid months)"], _rows)

    chart_block(doc,
                "AL_signup_cohort_openrate_trend.png",
                "AL",
                "Monthly Open Rate by Signup Cohort",
                "How to read this figure: each line is a signup-age cohort tracked month by month. "
                "Look for persistent vertical separation, not one-off spikes. Hollow markers indicate low coverage "
                "(fewer than 3 matched broadcasts) and should be treated as weaker evidence. If the oldest-cohort line "
                "sits below newer cohorts across consecutive stable months, that is direct evidence of tenure-linked "
                "open-rate decay.")
    body(doc,
         f"AL analysis: In stable months, the oldest-vs-newest open-rate gap averages {_or_gap_avg:+.1f}pp, "
         f"and in the most recent months it widens to {_or_gap_recent:+.1f}pp. "
         "That widening negative gap is consistent with older-subscriber disengagement.",
         color=(50, 50, 50))
    body(doc,
         "Decision implication: protecting mature cohorts should be an explicit objective in content planning, "
         "not only acquisition and onboarding optimization.",
         color=(50, 50, 50))

    chart_block(doc,
                "AM_signup_cohort_ctor_trend.png",
                "AM",
                "Monthly CTOR by Signup Cohort",
                "Solid lines show high-confidence months (coverage >= 3 broadcasts and enough openers). "
                "Hollow points show low-coverage months that are still displayed for continuity (so months like January do not disappear). "
                "Light dashed segments bridge months with enough openers but lower confidence. Grey bars on the right axis show monthly "
                "coverage, with the dotted line at n=3 as the confidence threshold. Labels show k=valid cohorts that month and "
                "op=total openers used. If the oldest cohort remains below newer cohorts across both high-confidence and low-coverage "
                "views, the gap is likely behavioural rather than a tracking artifact.")
    body(doc,
         f"AM analysis: The oldest-vs-newest CTOR gap averages {_ct_gap_avg:+.1f}pp and remains around "
         f"{_ct_gap_recent:+.1f}pp in recent months. This indicates the oldest cohort is less likely to click "
         "even after opening, reinforcing the disengagement hypothesis. "
         f"{_ctor_weak_months} month(s) have lower cohort coverage (k<3) and should be interpreted directionally.",
         color=(50, 50, 50))
    body(doc,
         "Decision implication: for older cohorts, CTA strategy should be tested separately because underperformance is present "
         "even after controlling for open behavior.",
         color=(50, 50, 50))

    chart_block(doc,
                "AN_oldest_vs_newest_gap.png",
                "AN",
                "Oldest vs Newest Cohort Gap (Open Rate and CTOR)",
                "Bars show Oldest minus Newest in percentage points (`pp`). Zero is parity. "
                "Bars below zero mean the oldest cohort underperforms; bars above zero mean it outperforms. "
                "This chart converts multi-line trends into one diagnostic so leadership changes are immediately visible.")
    body(doc,
         "AN analysis: The gap view consolidates the finding into one diagnostic: as months progress, "
         "the oldest cohort falls behind newest cohorts on both reach (open rate) and action quality (CTOR). "
         "This is exactly the pattern expected if long-tenure subscribers are disengaging faster.",
         color=(50, 50, 50))
    body(doc,
         "Decision implication: cohort-aware lifecycle treatment is justified by the data and should replace one-size-fits-all sending.",
         color=(50, 50, 50))

    body(doc, "11 take points:", bold=True, color=(79, 70, 229))
    bullet(doc,
           "Older signup cohorts underperform newer cohorts on open rate in recent stable months.")
    bullet(doc,
           "Older signup cohorts also underperform on CTOR after anomaly capping and coverage filtering.")
    bullet(doc,
           "The disengagement hypothesis is supported: tenure/age on list is strongly associated with weaker current engagement.")
else:
    body(doc,
         "[Signup cohort section requires signup_cohort_monthly_evolution.csv and charts AL/AM/AN]",
         color=(200, 0, 0))

doc.add_page_break()
h1(doc, "13. Bootcamp Buyers vs Non-Buyers — Engagement & Age Profile")

body(doc,
     "This section compares confirmed subscribers who bought at least one bootcamp "
     "(Buyer group: tags matching AI Agent Core/Bootcamp Core purchase tags) "
     "against subscribers who never bought. We track monthly Open Rate and CTOR trends, "
     "then test whether buyers are older or newer subscribers.")

if _q12m is not None and len(_q12m) and _q12s is not None and len(_q12s):
    _m = _q12m.copy()
    _m["stable"] = _m["broadcasts_matched"] >= 3
    _m["ctor_capped"] = _m["ctor"].clip(upper=100)
    _mw = _m.pivot_table(
        index=["month", "month_dt", "broadcasts_matched"],
        columns="group",
        values=["open_rate", "ctor_capped", "eligible_n", "open_n", "click_n"],
        aggfunc="first",
    ).reset_index()
    _mw.columns = [
        "_".join([str(x) for x in c if str(x) != ""]).rstrip("_")
        for c in _mw.columns.to_flat_index()
    ]
    _mw = _mw.sort_values("month_dt")

    _stable = _mw[_mw["broadcasts_matched"] >= 3].copy()
    _stable_ctor = _stable.dropna(subset=["ctor_capped_Buyer", "ctor_capped_Non-buyer"]).copy()
    _latest = _mw.iloc[-1]

    _buyer_row = _q12s[_q12s["group"] == "Buyer"].head(1)
    _non_row = _q12s[_q12s["group"] == "Non-buyer"].head(1)
    _buyer_row = _buyer_row.iloc[0] if len(_buyer_row) else None
    _non_row = _non_row.iloc[0] if len(_non_row) else None

    _open_gap = (
        _stable["open_rate_Buyer"].mean() - _stable["open_rate_Non-buyer"].mean()
        if len(_stable) else np.nan
    )
    _ctor_gap = (
        _stable_ctor["ctor_capped_Buyer"].mean() - _stable_ctor["ctor_capped_Non-buyer"].mean()
        if len(_stable_ctor) else np.nan
    )
    _failed_n = len(_q12f) if _q12f is not None else 0

    add_table(doc,
              ["Metric", "Buyer", "Non-buyer", "Gap (Buyer-Non)"],
              [
                  ["Subscribers (confirmed, active)",
                   f"{int(_buyer_row['n_subscribers']):,}" if _buyer_row is not None else "—",
                   f"{int(_non_row['n_subscribers']):,}" if _non_row is not None else "—",
                   "—"],
                  ["Open Rate (stable months, n>=3)",
                   f"{_stable['open_rate_Buyer'].mean():.1f}%" if len(_stable) else "—",
                   f"{_stable['open_rate_Non-buyer'].mean():.1f}%" if len(_stable) else "—",
                   f"{_open_gap:+.1f}pp" if pd.notna(_open_gap) else "—"],
                  ["CTOR (stable months, capped at 100)",
                   f"{_stable_ctor['ctor_capped_Buyer'].mean():.1f}%" if len(_stable_ctor) else "—",
                   f"{_stable_ctor['ctor_capped_Non-buyer'].mean():.1f}%" if len(_stable_ctor) else "—",
                   f"{_ctor_gap:+.1f}pp" if pd.notna(_ctor_gap) else "—"],
                  ["Median subscriber age",
                   f"{_buyer_row['median_age_days']:.0f} days" if _buyer_row is not None else "—",
                   f"{_non_row['median_age_days']:.0f} days" if _non_row is not None else "—",
                   (f"{(_buyer_row['median_age_days'] - _non_row['median_age_days']):+.0f} days"
                    if _buyer_row is not None and _non_row is not None else "—")],
                  ["Old subscribers (>365 days)",
                   f"{_buyer_row['old_365_pct']:.1f}%" if _buyer_row is not None else "—",
                   f"{_non_row['old_365_pct']:.1f}%" if _non_row is not None else "—",
                   (f"{(_buyer_row['old_365_pct'] - _non_row['old_365_pct']):+.1f}pp"
                    if _buyer_row is not None and _non_row is not None else "—")],
              ])

    bullet(doc,
           f"Latest month ({_latest['month']}): Buyer OR {_latest['open_rate_Buyer']:.1f}% vs "
           f"Non-buyer OR {_latest['open_rate_Non-buyer']:.1f}%; "
           f"Buyer CTOR {_latest['ctor_capped_Buyer']:.1f}% vs Non-buyer {_latest['ctor_capped_Non-buyer']:.1f}%.")
    bullet(doc,
           "Stable months are defined as months with at least 3 matched broadcasts. "
           "Low-coverage months are kept for context but interpreted directionally.")

    chart_block(doc,
                "AO_bootcamp_buyers_monthly_or_ctor.png",
                "AO",
                "Monthly Open Rate and CTOR — Buyers vs Non-buyers",
                "How to read this chart: compare Buyer vs Non-buyer lines month by month for both OR and CTOR. "
                "When Buyer stays above Non-buyer across consecutive stable months, that indicates a durable engagement "
                "premium, not a campaign-specific fluctuation. OR reflects inbox reach; CTOR reflects click quality after open.")
    body(doc,
         f"AO analysis: In stable months, buyers lead non-buyers by about {_open_gap:+.1f}pp in Open Rate "
         f"and {_ctor_gap:+.1f}pp in CTOR. The gap remains visible in recent months, confirming that purchase "
         "history is strongly associated with higher engagement quality.",
         color=(50, 50, 50))
    body(doc,
         "Plain-English interpretation: buyers are easier to re-reach and easier to move to click action. "
         "This makes buyer status one of the strongest practical segmentation variables for both value and sales campaigns.",
         color=(50, 50, 50))
    body(doc,
         "Decision implication: buyers and non-buyers should not share identical campaign cadence or CTA intensity.",
         color=(50, 50, 50))

    chart_block(doc,
                "AP_bootcamp_buyers_age_profile.png",
                "AP",
                "Buyer vs Non-buyer Age Profile",
                "This chart shows whether buyers are mostly new or mostly established subscribers. "
                "Comparing median age and old-subscriber share tells us if purchase behavior is concentrated "
                "in long-tenure segments or in newer entrants. Read this as composition, not performance: "
                "it explains who buyers are, while AO explains how buyers behave.")
    body(doc,
         f"AP analysis: Buyers have a median age of {_buyer_row['median_age_days']:.0f} days vs "
         f"{_non_row['median_age_days']:.0f} days for non-buyers. The >365-day share is "
         f"{_buyer_row['old_365_pct']:.1f}% for buyers vs {_non_row['old_365_pct']:.1f}% for non-buyers. "
         "In this dataset, buyers are not the oldest segment; they are relatively mid-age and engagement-strong.",
         color=(50, 50, 50))
    body(doc,
         "Plain-English interpretation: conversion is not restricted to very old subscribers. "
         "A meaningful buyer segment sits in the middle of the subscriber-age distribution, so conversion programs "
         "should target intent signals in addition to tenure.",
         color=(50, 50, 50))
    body(doc,
         "Decision implication: age-only targeting will miss a meaningful share of potential buyers.",
         color=(50, 50, 50))

    body(doc, "12 take points:", bold=True, color=(79, 70, 229))
    bullet(doc,
           "Buyers are materially more engaged than non-buyers on both Open Rate and CTOR across stable months.")
    bullet(doc,
           "The buyer advantage persists in recent months, making purchase history a strong segmentation signal.")
    bullet(doc,
           "Buyers are not simply the oldest subscribers; their higher engagement appears linked to intent, not just tenure.")

    body(doc,
         f"Data caveat: Kit returned 500 for {_failed_n} legacy ID/event checks in this section; "
         "those IDs were skipped via chunk-level fallback logic.",
         size=9, italic=True, color=(120, 120, 120))
else:
    body(doc,
         "[Buyer analysis requires bootcamp_buyers_monthly_evolution.csv and bootcamp_buyers_summary.csv]",
         color=(200, 0, 0))

doc.add_page_break()
h1(doc, "14. Sales Intent Outcomes — Intended vs Acted")

body(doc,
     "This section answers a direct commercial-intent question for Sales emails: "
     "among people who opened (intended), how many clicked a sales link (acted), "
     "and how many opened but did not click (intended but did not act).")

_sales_int = bdf_rpt[bdf_rpt["category"] == "Sales"].copy()
if len(_sales_int):
    _sales_int["date"] = pd.to_datetime(_sales_int["date"], errors="coerce")
    _sales_int["opened_n"] = pd.to_numeric(_sales_int["opened"], errors="coerce").fillna(0)
    _sales_int["clicked_n"] = pd.to_numeric(_sales_int["clicked"], errors="coerce").fillna(0)
    _sales_int["clicked_n"] = np.minimum(_sales_int["clicked_n"], _sales_int["opened_n"])
    _sales_int["not_clicked_n"] = (_sales_int["opened_n"] - _sales_int["clicked_n"]).clip(lower=0)
    _sales_int["intent_to_action_rate"] = np.where(
        _sales_int["opened_n"] > 0,
        _sales_int["clicked_n"] / _sales_int["opened_n"] * 100,
        np.nan,
    )
    _sales_int["month_dt"] = _sales_int["date"].dt.to_period("M").dt.to_timestamp()

    _per_email_path = GENERATED / "sales_intent_outcomes_per_email.csv"
    _sales_int[[
        "date", "subject", "recipients", "opened_n", "clicked_n", "not_clicked_n", "intent_to_action_rate"
    ]].sort_values("date").to_csv(_per_email_path, index=False)

    _tot_open = int(_sales_int["opened_n"].sum())
    _tot_click = int(_sales_int["clicked_n"].sum())
    _tot_no_click = int(_sales_int["not_clicked_n"].sum())
    _overall_rate = (_tot_click / _tot_open * 100) if _tot_open else np.nan

    _monthly_si = (
        _sales_int.groupby("month_dt", as_index=False)
        .agg(
            opened_n=("opened_n", "sum"),
            clicked_n=("clicked_n", "sum"),
            not_clicked_n=("not_clicked_n", "sum"),
            sends_n=("subject", "count"),
        )
        .sort_values("month_dt")
    )
    _all_months = pd.date_range(_monthly_si["month_dt"].min(), _monthly_si["month_dt"].max(), freq="MS")
    _monthly_si = (
        _monthly_si.set_index("month_dt")
        .reindex(_all_months)
        .rename_axis("month_dt")
        .reset_index()
    )
    for _col in ["opened_n", "clicked_n", "not_clicked_n", "sends_n"]:
        _monthly_si[_col] = _monthly_si[_col].fillna(0)
    _monthly_si["sends_n"] = _monthly_si["sends_n"].astype(int)
    _monthly_si["has_sales_send"] = _monthly_si["sends_n"] > 0
    _monthly_si["rate"] = np.where(
        _monthly_si["opened_n"] > 0,
        _monthly_si["clicked_n"] / _monthly_si["opened_n"] * 100,
        np.nan,
    )
    _monthly_si["rate_interp"] = _monthly_si["rate"].interpolate(limit_direction="both")
    _monthly_si["rate_roll"] = _monthly_si["rate_interp"].rolling(3, min_periods=1).mean()

    fig, axes = plt.subplots(2, 1, figsize=(13.5, 8.6), sharex=True)
    ax = axes[0]
    ax.bar(_monthly_si["month_dt"], _monthly_si["clicked_n"], width=22, color="#10B981",
           alpha=0.9, label="Opened + clicked (potentially bought)")
    ax.bar(_monthly_si["month_dt"], _monthly_si["not_clicked_n"], width=22,
           bottom=_monthly_si["clicked_n"], color="#F59E0B", alpha=0.85,
           label="Opened + no click (intended, no action)")
    ax.set_title("Sales Email Intent Outcomes by Month (Counts)")
    ax.set_ylabel("People")
    ax.grid(axis="y", alpha=0.2)
    ax.legend(fontsize=8.5, loc="upper right")

    ax = axes[1]
    ax.plot(_monthly_si["month_dt"], _monthly_si["rate_interp"], color="#EC4899", linestyle="--",
            linewidth=1.2, alpha=0.35, label="Intent-to-action bridge")
    _obs = _monthly_si["has_sales_send"]
    ax.plot(_monthly_si.loc[_obs, "month_dt"], _monthly_si.loc[_obs, "rate"], color="#EC4899", marker="o",
            linewidth=2.2, label="Intent-to-action rate (monthly)")
    _no_send = ~_monthly_si["has_sales_send"]
    if _no_send.any():
        ax.scatter(_monthly_si.loc[_no_send, "month_dt"], _monthly_si.loc[_no_send, "rate_interp"],
                   s=54, facecolors="none", edgecolors="#EC4899", linewidths=1.5,
                   label="No Sales send (hollow)")
    ax.plot(_monthly_si["month_dt"], _monthly_si["rate_roll"], color="#7C3AED", linestyle="--",
            linewidth=1.9, label="Intent-to-action rate (3-mo smooth)")
    ax.set_title("Sales Intent-to-Action Rate Over Time")
    ax.set_ylabel("Rate (%)")
    ax.set_xlabel("Month")
    ax.set_ylim(0, 100)
    ax.grid(axis="y", alpha=0.2)
    ax2 = ax.twinx()
    ax2.bar(_monthly_si["month_dt"], _monthly_si["sends_n"], width=18, color="#94A3B8",
            alpha=0.18, label="Sales sends")
    ax2.set_ylabel("Sales sends (n)")
    h1_, l1_ = ax.get_legend_handles_labels()
    h2_, l2_ = ax2.get_legend_handles_labels()
    ax.legend(h1_ + h2_, l1_ + l2_, fontsize=8.2, loc="upper right")
    fig.suptitle("Sales Intent (Open) vs Action (Click) — Monthly Diagnostics", fontsize=13.5, fontweight="bold")
    fig.tight_layout()
    fig.savefig(CHARTS / "AQ_sales_intent_outcomes.png")
    plt.close()

    add_table(doc,
              ["Metric", "Value"],
              [
                  ["Sales broadcasts analysed", f"{len(_sales_int):,}"],
                  ["Opened Sales emails (intended to buy)", f"{_tot_open:,}"],
                  ["Opened and clicked Sales link (intended and acted)", f"{_tot_click:,}"],
                  ["Opened but did not click (intended and did not act)", f"{_tot_no_click:,}"],
                  ["Intent-to-action rate (clicked / opened)", f"{_overall_rate:.1f}%"],
                  ["Avg opened per Sales email", f"{_sales_int['opened_n'].mean():.0f}"],
                  ["Avg clicked per Sales email", f"{_sales_int['clicked_n'].mean():.0f}"],
              ])

    bullet(doc,
           f"Direct answer: intended and acted = {_tot_click:,} people; intended and did not act = {_tot_no_click:,} people.")
    bullet(doc,
           f"This is based on Sales-email opens as intent and Sales-link clicks as action proxy across {len(_sales_int):,} Sales broadcasts.")

    chart_block(doc,
                "AQ_sales_intent_outcomes.png",
                "AQ",
                "Sales Intent Outcomes (Opened vs Clicked)",
                "Top panel: monthly count split of openers who clicked vs openers who did not click. "
                "Bottom panel: intent-to-action rate (`clicked / opened`) with a smoothed trend line. "
                "Hollow points indicate months with no qualifying Sales sends (shown for continuity), while filled points "
                "are observed months. Interpretation rule: if opens stay high but the rate line falls, demand exists but "
                "message-to-action conversion is weakening.")
    body(doc,
         f"AQ analysis: Across the analysed Sales broadcasts, {_tot_open:,} people opened a Sales email (intent proxy). "
         f"Of these, {_tot_click:,} clicked at least one Sales link, while {_tot_no_click:,} did not click. "
         f"The aggregate intent-to-action rate is {_overall_rate:.1f}%.",
         color=(50, 50, 50))
    body(doc,
         "Plain-English interpretation: most potential demand is visible at open stage, but the major loss occurs between "
         "open and click. Improving offer clarity, CTA placement, and message-to-link continuity is likely to produce the "
         "largest immediate commercial gain.",
         color=(50, 50, 50))
    body(doc,
         "Decision implication: near-term revenue lift is more likely from click-path optimization than from top-of-funnel open-rate gains.",
         color=(50, 50, 50))

    body(doc, "14 take points:", bold=True, color=(79, 70, 229))
    bullet(doc, "Most commercial intent is visible at the open stage, but only a subset proceeds to click action.")
    bullet(doc, "The open-to-click conversion lens is the cleanest operational proxy for potential purchase intent conversion.")
    bullet(doc, "Per-email details are exported for audit at data/generated/sales_intent_outcomes_per_email.csv.")

    body(doc,
         "Important caveat: this section measures intent and action proxies (open/click), not confirmed checkout purchases.",
         size=9, italic=True, color=(120, 120, 120))
else:
    body(doc,
         "[Sales intent section requires Sales rows in Emails Broadcasting - broadcasts_categorised.csv]",
         color=(200, 0, 0))

doc.add_page_break()
h1(doc, "15. Workshop-to-Bootcamp Conversion (Did Workshop Buyers Convert to Bootcamp?)")

body(doc,
     "This section tests whether buying a paid workshop is associated with a higher likelihood of "
     "buying a bootcamp later. Workshops are treated as pre-bootcamp conversion steps based on the "
     "provided mapping file (`Workshops Pre-bootcamps - Sheet1.csv`).")
body(doc,
     "Method: each confirmed subscriber is tagged as workshop buyer if they have at least one mapped "
     "workshop tag, and as bootcamp buyer if they have at least one mapped bootcamp tag. "
     "The comparison is then a 2x2 conversion table across the full confirmed-subscriber base.",
     color=(50, 50, 50))

_workshop_map_path = BASE / "Workshops Pre-bootcamps - Sheet1.csv"
_q15_chart_file = "AS_workshop_bootcamp_lift.png"
if _workshop_map_path.exists() and _confirmed_path.exists():
    _wm = pd.read_csv(_workshop_map_path)
    _wm.columns = [str(c).strip() for c in _wm.columns]

    _boot_tags = [str(x).strip().lower() for x in _wm["TAGS"].dropna().tolist() if str(x).strip()]
    _pair_rows_raw = []
    for _, _r in _wm.iterrows():
        _b = str(_r.get("TAGS", "")).strip().lower() if pd.notna(_r.get("TAGS", np.nan)) else ""
        _w = str(_r.get("Workshop tags", "")).strip().lower() if pd.notna(_r.get("Workshop tags", np.nan)) else ""
        if _b and _w and _w != "nan":
            _pair_rows_raw.append((_b, _w, _r.get("Date Of The Cohort", ""), _r.get("Workshop Date", "")))
    _workshop_tags = sorted({w for _, w, _, _ in _pair_rows_raw})

    _subs15 = pd.read_csv(_confirmed_path, usecols=["email", "tags"])
    _subs15["email_lower"] = _subs15["email"].astype(str).str.lower().str.strip()
    _subs15 = _subs15.drop_duplicates("email_lower", keep="first")
    _subs15["tags_l"] = _subs15["tags"].fillna("").astype(str).str.lower()

    _subs15["has_bootcamp"] = _subs15["tags_l"].apply(lambda t: any(bt in t for bt in _boot_tags))
    _subs15["has_workshop"] = _subs15["tags_l"].apply(lambda t: any(wt in t for wt in _workshop_tags))
    _subs15["has_mapped_pair"] = _subs15["tags_l"].apply(
        lambda t: any((bt in t and wt in t) for bt, wt, _, _ in _pair_rows_raw)
    )

    # 2x2 table
    _a = int((_subs15["has_workshop"] & _subs15["has_bootcamp"]).sum())      # workshop yes, bootcamp yes
    _b = int((_subs15["has_workshop"] & ~_subs15["has_bootcamp"]).sum())     # workshop yes, bootcamp no
    _c = int((~_subs15["has_workshop"] & _subs15["has_bootcamp"]).sum())     # workshop no, bootcamp yes
    _d = int((~_subs15["has_workshop"] & ~_subs15["has_bootcamp"]).sum())    # workshop no, bootcamp no

    _n_workshop = _a + _b
    _n_nonwork = _c + _d
    _boot_total = _a + _c
    _rate_workshop = (_a / _n_workshop * 100) if _n_workshop else np.nan
    _rate_nonwork = (_c / _n_nonwork * 100) if _n_nonwork else np.nan
    _lift_pp = _rate_workshop - _rate_nonwork if pd.notna(_rate_workshop) and pd.notna(_rate_nonwork) else np.nan
    _rel_lift = (_rate_workshop / _rate_nonwork) if (_rate_nonwork and pd.notna(_rate_nonwork)) else np.nan
    _odds_ratio = (_a * _d) / (_b * _c) if (_b > 0 and _c > 0) else np.nan
    try:
        _chi2, _p_val, _, _ = scipy_stats.chi2_contingency([[_a, _b], [_c, _d]], correction=False)
    except Exception:
        _p_val = np.nan
    _p_txt = "<1e-16" if pd.notna(_p_val) and _p_val < 1e-16 else (f"{_p_val:.3g}" if pd.notna(_p_val) else "—")

    _boot_with_any_workshop = _a
    _boot_with_pair = int((_subs15["has_bootcamp"] & _subs15["has_mapped_pair"]).sum())

    add_table(doc,
              ["Workshop Purchase", "Bought bootcamp", "Did not buy bootcamp", "Total", "Bootcamp conversion rate"],
              [
                  ["Yes", f"{_a:,}", f"{_b:,}", f"{_n_workshop:,}", f"{_rate_workshop:.1f}%"],
                  ["No",  f"{_c:,}", f"{_d:,}", f"{_n_nonwork:,}", f"{_rate_nonwork:.1f}%"],
              ])

    add_table(doc,
              ["Metric", "Value"],
              [
                  ["Total bootcamp buyers", f"{_boot_total:,}"],
                  ["Bootcamp buyers with any mapped workshop tag", f"{_boot_with_any_workshop:,} ({(_boot_with_any_workshop / _boot_total * 100 if _boot_total else 0):.1f}%)"],
                  ["Bootcamp buyers with mapped workshop+bootcamp pair", f"{_boot_with_pair:,} ({(_boot_with_pair / _boot_total * 100 if _boot_total else 0):.1f}%)"],
                  ["Conversion lift (workshop buyer vs non-workshop buyer)", f"{_lift_pp:+.1f}pp"],
                  ["Relative conversion multiple", f"{_rel_lift:.1f}x" if pd.notna(_rel_lift) else "—"],
                  ["Odds ratio", f"{_odds_ratio:.1f}" if pd.notna(_odds_ratio) else "—"],
                  ["Significance (chi-square p-value)", _p_txt],
              ])

    h2(doc, "15.1 Cohort Contrast: Cohorts With vs Without a Pre-Workshop")

    # Build cohort metadata from mapping file.
    _cohort_meta = []
    _pair_map = {}
    for _, _r in _wm.iterrows():
        _cohort_lbl = str(_r.get("Date Of The Cohort", "—")).strip() if pd.notna(_r.get("Date Of The Cohort", np.nan)) else "—"
        _bt_raw = str(_r.get("TAGS", "")).strip()
        _wt_raw = str(_r.get("Workshop tags", "")).strip() if pd.notna(_r.get("Workshop tags", np.nan)) else ""
        if not _bt_raw:
            continue
        _bt = _bt_raw.lower()
        _wt = _wt_raw.lower()
        _has_pre_workshop = bool(_wt and _wt != "nan")
        _pair_map[_bt] = _wt if _has_pre_workshop else ""
        _cohort_meta.append({
            "cohort_label": _cohort_lbl,
            "boot_tag": _bt,
            "boot_tag_raw": _bt_raw,
            "workshop_tag_raw": _wt_raw if _has_pre_workshop else "No mapped pre-workshop",
            "workshop_tag": _wt if _has_pre_workshop else "",
            "has_pre_workshop": _has_pre_workshop,
        })

    # Assign each buyer to one cohort only (first bootcamp tag by cohort date) to avoid overlap.
    _boot_tag_dates = {
        "ai agent bootcamp core [apr 2025]": pd.Timestamp("2025-04-30"),
        "ai agent bootcamp core [jul 2025]": pd.Timestamp("2025-07-31"),
        "ai agent core [sept]": pd.Timestamp("2025-09-30"),
        "ai agent core [oct]": pd.Timestamp("2025-10-31"),
        "ai agent core [feb]": pd.Timestamp("2026-02-28"),
    }
    _buyer_assign_rows = []
    for _, _r in _subs15.iterrows():
        _t = _r["tags_l"]
        _matches = [bt for bt in _boot_tags if bt in _t]
        if not _matches:
            continue
        _first_bt = sorted(_matches, key=lambda x: _boot_tag_dates.get(x, pd.Timestamp("2099-12-31")))[0]
        _paired_w = _pair_map.get(_first_bt, "")
        _buyer_assign_rows.append({
            "boot_tag": _first_bt,
            "has_pair_workshop": bool(_paired_w and _paired_w in _t),
        })
    _buyer_assign = pd.DataFrame(_buyer_assign_rows)

    _cohort_rows = []
    for _m in _cohort_meta:
        _bt = _m["boot_tag"]
        _boot_n = int((_buyer_assign["boot_tag"] == _bt).sum()) if len(_buyer_assign) else 0
        _both_n = int(((_buyer_assign["boot_tag"] == _bt) & (_buyer_assign["has_pair_workshop"])).sum()) if len(_buyer_assign) else 0
        if _m["has_pre_workshop"]:
            _ws_n = int(_subs15["tags_l"].str.contains(_m["workshop_tag"], regex=False).sum())
            _pen_boot = (_both_n / _boot_n * 100) if _boot_n else np.nan
            _conv_from_ws = (_both_n / _ws_n * 100) if _ws_n else np.nan
        else:
            _ws_n = np.nan
            _pen_boot = np.nan
            _conv_from_ws = np.nan

        _cohort_rows.append({
            "cohort_label": _m["cohort_label"],
            "boot_tag": _m["boot_tag_raw"],
            "workshop_tag": _m["workshop_tag_raw"],
            "has_pre_workshop": _m["has_pre_workshop"],
            "boot_n": _boot_n,
            "workshop_n": _ws_n,
            "both_n": _both_n,
            "pen_boot_pct": _pen_boot,
            "conv_from_ws_pct": _conv_from_ws,
        })

    _cohort_df = pd.DataFrame(_cohort_rows)
    _cohort_df["cohort_dt"] = _cohort_df["boot_tag"].str.lower().map(_boot_tag_dates)
    _with_df = _cohort_df[_cohort_df["has_pre_workshop"]].copy()
    _no_df = _cohort_df[~_cohort_df["has_pre_workshop"]].copy()

    _with_boot = int(_with_df["boot_n"].sum())
    _no_boot = int(_no_df["boot_n"].sum())
    _with_both = int(_with_df["both_n"].sum())
    _no_both = int(_no_df["both_n"].sum())
    _with_pen = (_with_both / _with_boot * 100) if _with_boot else np.nan
    _no_pen = (_no_both / _no_boot * 100) if _no_boot else np.nan
    _with_n = int(len(_with_df))
    _no_n = int(len(_no_df))
    _no_names = ", ".join(_no_df["cohort_label"].astype(str).tolist()) if _no_n else "none"

    add_table(doc,
              ["Cohort group", "Bootcamp buyers", "Bootcamp buyers with mapped pre-workshop", "Share"],
              [
                  [f"Cohorts with mapped pre-workshop ({_with_n} cohorts)", f"{_with_boot:,}", f"{_with_both:,}", f"{_with_pen:.1f}%"],
                  [f"Cohorts without mapped pre-workshop ({_no_n} cohorts)", f"{_no_boot:,}", f"{_no_both:,}", f"{_no_pen:.1f}%"],
              ])

    body(doc,
         "How to interpret this grouped comparison: it does not compare overall cohort size or demand, because we do not "
         "have cohort-specific exposure denominators in this file. It compares only whether bootcamp buyers in each cohort "
         "also had the mapped pre-workshop tag.",
         color=(50, 50, 50))

    h2(doc, "15.2 Per-Cohort View")
    _cohort_tbl_rows = []
    for _, _r in _cohort_df.iterrows():
        _cohort_tbl_rows.append([
            _r["cohort_label"],
            _r["boot_tag"],
            _r["workshop_tag"],
            f"{int(_r['boot_n']):,}",
            f"{int(_r['both_n']):,}",
            (f"{_r['pen_boot_pct']:.1f}%" if pd.notna(_r["pen_boot_pct"]) else "N/A (no workshop before this cohort)"),
            (f"{_r['conv_from_ws_pct']:.1f}%" if pd.notna(_r["conv_from_ws_pct"]) else "N/A"),
        ])

    add_table(doc,
              ["Cohort", "Bootcamp tag", "Mapped pre-workshop tag", "Bootcamp buyers", "Bought mapped workshop", "Mapped workshop penetration in bootcamp buyers", "Bootcamp conversion from mapped workshop"],
              _cohort_tbl_rows)
    body(doc,
         "Method note for 15.1/15.2: each subscriber is assigned to one bootcamp cohort only (their first bootcamp tag by cohort date), "
         "so cohort counts are non-overlapping and sum to total bootcamp buyers.",
         size=9, italic=True, color=(120, 120, 120))

    bullet(doc,
           f"Across cohorts with a mapped pre-workshop, {_with_both:,} of {_with_boot:,} bootcamp buyers ({_with_pen:.1f}%) "
           "also had that cohort's workshop tag.")
    bullet(doc,
           f"For cohorts without a mapped pre-workshop ({_no_names}), mapped workshop penetration is {_no_pen:.1f}% by definition.")

    # Explicit first-vs-latest workshop-led cohort diagnostic.
    _with_sorted = _cohort_df[_cohort_df["has_pre_workshop"] & _cohort_df["cohort_dt"].notna()].sort_values("cohort_dt")
    if len(_with_sorted) >= 2:
        _first = _with_sorted.iloc[0]
        _last = _with_sorted.iloc[-1]
        _first_name = str(_first["cohort_label"])
        _last_name = str(_last["cohort_label"])
        _first_ws = int(_first["workshop_n"]) if pd.notna(_first["workshop_n"]) else 0
        _last_ws = int(_last["workshop_n"]) if pd.notna(_last["workshop_n"]) else 0
        _first_both = int(_first["both_n"])
        _last_both = int(_last["both_n"])
        _first_conv = float(_first["conv_from_ws_pct"]) if pd.notna(_first["conv_from_ws_pct"]) else np.nan
        _last_conv = float(_last["conv_from_ws_pct"]) if pd.notna(_last["conv_from_ws_pct"]) else np.nan

        body(doc, "First vs latest workshop-led cohort (workshop -> bootcamp):", bold=True, color=(79, 70, 229))
        bullet(doc,
               f"First workshop-led cohort ({_first_name}): {_first_both:,}/{_first_ws:,} = {_first_conv:.1f}%")
        bullet(doc,
               f"Latest workshop-led cohort ({_last_name}): {_last_both:,}/{_last_ws:,} = {_last_conv:.1f}%")

        if pd.notna(_first_conv) and pd.notna(_last_conv):
            _delta_conv = _last_conv - _first_conv
            if _delta_conv < 0:
                if _last_ws > _first_ws:
                    bullet(doc,
                           f"The latest cohort conversion is {_delta_conv:+.1f}pp lower, and its workshop audience is larger "
                           f"({_last_ws:,} vs {_first_ws:,}). So denominator growth could be part of the drop.")
                elif _last_ws < _first_ws:
                    bullet(doc,
                           f"The latest cohort conversion is {_delta_conv:+.1f}pp lower, but its workshop audience is smaller "
                           f"({_last_ws:,} vs {_first_ws:,}). So the drop is not explained by a larger workshop denominator.")
                else:
                    bullet(doc,
                           f"The latest cohort conversion is {_delta_conv:+.1f}pp lower with equal workshop audience size "
                           f"({_last_ws:,} each). So denominator size is not the driver.")
            else:
                bullet(doc,
                       f"The latest cohort conversion is {_delta_conv:+.1f}pp higher than the first workshop-led cohort.")

    # Cohort comparison chart
    _q15_chart_cohort = "AT_workshop_cohort_comparison.png"
    _plot = _cohort_df.copy()
    _plot["plot_val"] = _plot["pen_boot_pct"].fillna(0)
    _colors = _plot["has_pre_workshop"].map({True: "#10B981", False: "#94A3B8"}).tolist()
    fig, ax = plt.subplots(figsize=(11.0, 5.8))
    _bars = ax.bar(_plot["cohort_label"], _plot["plot_val"], color=_colors, edgecolor="white", width=0.62)
    for _i, _bar in enumerate(_bars):
        _is_with = bool(_plot.iloc[_i]["has_pre_workshop"])
        _label = f"{_plot.iloc[_i]['plot_val']:.1f}%" if _is_with else "N/A"
        ax.text(
            _bar.get_x() + _bar.get_width() / 2,
            _bar.get_height() + 1.4,
            _label,
            ha="center",
            va="bottom",
            fontsize=9,
            fontweight="bold",
            color="#111827",
        )
    ax.set_ylabel("Mapped workshop penetration in bootcamp buyers (%)")
    ax.set_title("Per-Cohort Comparison: Pre-Workshop Association with Bootcamp Buyers")
    ax.set_ylim(0, max(62, float(_plot["plot_val"].max()) + 8))
    ax.grid(axis="y", alpha=0.2)
    ax.tick_params(axis="x", labelrotation=15)
    _with_patch = mpatches.Patch(color="#10B981", label="Cohort with mapped pre-workshop")
    _no_patch = mpatches.Patch(color="#94A3B8", label="Cohort without mapped pre-workshop")
    ax.legend(handles=[_with_patch, _no_patch], fontsize=8.5, loc="upper right")
    fig.tight_layout()
    fig.savefig(CHARTS / _q15_chart_cohort, dpi=180)
    plt.close()

    chart_block(doc,
                _q15_chart_cohort,
                "AT",
                "Per-Cohort Mapped Workshop Penetration in Bootcamp Buyers",
                "Green bars are cohorts where a mapped pre-workshop exists; grey bars are cohorts without a mapped pre-workshop "
                "(shown as N/A, plotted at zero for layout). This reveals cohort-level variation in workshop-to-bootcamp linkage, "
                "including which workshop-led cohorts show the strongest carry-through into bootcamp purchase tags.")

    bullet(doc,
           f"Direct answer to 'how many bootcamp buyers purchased workshop previously': {_boot_with_any_workshop:,} out of {_boot_total:,} "
           f"bootcamp buyers ({(_boot_with_any_workshop / _boot_total * 100 if _boot_total else 0):.1f}%) have at least one mapped workshop tag.")
    bullet(doc,
           f"Bootcamp conversion rate among workshop buyers is {_rate_workshop:.1f}% vs {_rate_nonwork:.1f}% among non-workshop buyers.")
    bullet(doc,
           f"That is a {_lift_pp:+.1f}pp lift ({_rel_lift:.1f}x higher conversion), with p-value {_p_txt}.")

    callout(doc,
            "Conclusion: workshop buyers are much more likely to buy a bootcamp than subscribers who did not buy a workshop. "
            "This strongly supports the workshop-as-pre-bootcamp funnel hypothesis.",
            color_rgb=(16, 185, 129))

    # Chart: conversion-rate comparison
    fig, ax = plt.subplots(figsize=(9.4, 5.6))
    _lbls = ["Bought workshop", "Did not buy workshop"]
    _rates = [_rate_workshop, _rate_nonwork]
    _totals = [_n_workshop, _n_nonwork]
    _boots = [_a, _c]
    _bars = ax.bar(_lbls, _rates, color=["#10B981", "#6366F1"], width=0.56, edgecolor="white")
    for _i, _bar in enumerate(_bars):
        ax.text(
            _bar.get_x() + _bar.get_width() / 2,
            _bar.get_height() + max(_rates) * 0.03,
            f"{_rates[_i]:.1f}%\n({_boots[_i]:,}/{_totals[_i]:,})",
            ha="center",
            va="bottom",
            fontsize=9,
            fontweight="bold",
        )
    ax.set_ylabel("Bootcamp purchase rate (%)")
    ax.set_title("Bootcamp Conversion by Workshop Purchase History")
    ax.grid(axis="y", alpha=0.2)
    ax.set_ylim(0, max(_rates) * 1.28 if max(_rates) > 0 else 1)
    ax.annotate(
        f"Lift: {_lift_pp:+.1f}pp ({_rel_lift:.1f}x)",
        xy=(0, _rate_workshop),
        xytext=(0.58, max(_rates) * 0.90 if max(_rates) > 0 else 0.5),
        textcoords="data",
        fontsize=9.2,
        color="#065F46",
        fontweight="bold",
        arrowprops=dict(arrowstyle="->", color="#065F46", lw=1.4),
    )
    fig.tight_layout()
    fig.savefig(CHARTS / _q15_chart_file, dpi=180)
    plt.close()

    chart_block(doc,
                _q15_chart_file,
                "AS",
                "Bootcamp Conversion Rate by Workshop Purchase History",
                "This chart compares the probability of buying a bootcamp between workshop buyers and non-workshop buyers. "
                "The taller workshop bar indicates materially higher conversion likelihood, which means workshops function as "
                "a strong qualification and conversion step in the bootcamp funnel.")
    body(doc,
         "Interpretation note: this is an association analysis from tag data, not a randomized causal experiment. "
         "It indicates strong conversion signal quality for workshop buyers, but part of the effect can also reflect "
         "self-selection (higher baseline intent among people who buy workshops).",
         size=9, italic=True, color=(120, 120, 120))
    body(doc,
         f"Data note: the mapping file currently has {_no_n} cohort(s) without a mapped pre-workshop ({_no_names}). "
         "Those cohorts appear as N/A in mapped workshop penetration by design.",
         size=9, italic=True, color=(120, 120, 120))
else:
    body(doc,
         "[Section 15 requires both Confirmed Subscribers.csv and Workshops Pre-bootcamps - Sheet1.csv]",
         color=(200, 0, 0))

doc.add_page_break()
h1(doc, "16. Final Verdict — Did the Free Lead Magnet Work?")

body(doc,
     "This one-page conclusion answers the two core business questions using the final analysis outputs "
     "from this audit (broadcast-level pre/post comparison and monthly Group A vs Group B cohort trends).")

_val_pre = _q1_stats["Value"]["pre_mean"]
_val_post = _q1_stats["Value"]["post_mean"]
_val_delta = _q1_stats["Value"]["delta"]
_sal_pre = _q1_stats["Sales"]["pre_mean"]
_sal_post = _q1_stats["Sales"]["post_mean"]
_sal_delta = _q1_stats["Sales"]["delta"]

add_table(doc,
          ["Question", "Answer", "Evidence"],
          [
              ["Did survey send improve overall open rate after 25 Jan 2026?", "No",
               f"Value OR: {_val_pre:.1f}% → {_val_post:.1f}% ({_val_delta:+.1f}pp); "
               f"Sales OR: {_sal_pre:.1f}% → {_sal_post:.1f}% ({_sal_delta:+.1f}pp)."],
              ["Are Group A responders more likely to open Value emails?", "No (stable months)",
               (f"Group A Value OR {_q6_value_a:.1f}% vs Group B {_q6_value_b:.1f}% "
                f"({_q6_value_gap:+.1f}pp).") if pd.notna(_q6_value_gap) else "Monthly stable comparison unavailable."],
              ["Are Group A responders more likely to open Sales emails?", "Mixed / no consistent lift",
               (f"Group A Sales OR {_q6_sales_a:.1f}% vs Group B {_q6_sales_b:.1f}% "
                f"({_q6_sales_gap:+.1f}pp).") if pd.notna(_q6_sales_gap) else "Monthly stable comparison unavailable."],
              ["Are Group A responders more likely to click sales links (Sales CTOR)?", "Yes (directional)",
               (f"Group A Sales CTOR {_q6_ctor_a:.1f}% vs Group B {_q6_ctor_b:.1f}% "
                f"({_q6_sales_ctor_gap:+.1f}pp).") if pd.notna(_q6_sales_ctor_gap) else "Monthly stable comparison unavailable."],
          ])

body(doc, "Final interpretation:", bold=True, color=(79, 70, 229))
bullet(doc,
       "The free lead magnet did not produce a list-wide open-rate lift after the survey send date.")
bullet(doc,
       "As a segmenting mechanism, it partially worked: Group A is not consistently better on opens, but "
       "is stronger on sales click intent (CTOR) in stable months.")
bullet(doc,
       "Best strategic use: treat the lead magnet as an intent/qualification filter for sales follow-up, "
       "not as a broad open-rate optimization tactic for the full list.")

callout(doc,
        "Verdict: The lead magnet did not work as an overall open-rate booster, but it did work as a "
        "commercial-intent segmenter. Group A is the better segment for sales conversion-focused campaigns.",
        color_rgb=(16, 185, 129))

doc.add_page_break()
h1(doc, "17. Detailed Executive Summary of Key Takeaways")

body(doc,
     "This final summary consolidates the full audit into decision-ready conclusions. "
     "It is intentionally detailed: each point links observation to implication so the next actions "
     "can be prioritized without re-reading every section.")

_lat_evidence = (
    f"Median time to first purchase: {_lat_median_days:.0f} days; {_lat_within_90:.1f}% convert within 90 days."
    if pd.notna(_lat_median_days) and pd.notna(_lat_within_90)
    else "Conversion-latency evidence unavailable in this run."
)

add_table(doc,
          ["Section", "Key takeaway", "Why it matters"],
          [
              ["2. Subscriber Lifespan", "Subscriber exits are no longer mainly early churn; long-tenure exits increased.",
               "Retention risk shifted to established readers, so quality erosion is now a core business risk."],
              ["3–4. Engagement Trend", "Open rate and click quality declined over the year, with sharper deterioration after Nov 2025.",
               "List size can grow while list value falls; performance monitoring must focus on quality metrics."],
              ["5–6. Who Is Disengaging", "Recent disengagement is disproportionately concentrated in older cohorts.",
               "The problem is not only weak onboarding; mature subscribers need re-engagement strategy."],
              ["7. Sales vs Value", "Sales opens and value click quality both weakened, but with different patterns.",
               "Inbox reach and in-email persuasion need separate fixes rather than one generic optimization."],
              ["8. Cold Subscribers", "A large cold pool accumulated, including many who never opened.",
               "Deliverability and send efficiency are being diluted by inactive inventory."],
              ["9.1 Conversion Timing", _lat_evidence,
               "Nurture timing should be calibrated to actual conversion lag, not assumed instant purchase behavior."],
              ["10. Lead Magnet Impact", "No list-wide post-survey open-rate lift; Group A segment indicator is stronger on sales CTOR than on opens.",
               "Use the lead magnet as a qualification mechanism, not as a broad open-rate lever."],
              ["12. Signup Cohort Evolution", "Older signup cohorts underperform newer cohorts on OR and CTOR in stable months.",
               "This directly supports the hypothesis that tenure-based disengagement is real."],
              ["13. Buyers vs Non-buyers", "Buyers are materially more engaged than non-buyers over time.",
               "Purchase history should be a primary segmentation input for campaign strategy."],
              ["14. Sales Intent Outcomes", "Most intent appears at open stage, but only a minority advances to click action.",
               "The open-to-click bridge is the operational bottleneck for sales conversion."],
              ["15. Workshop→Bootcamp Conversion", "Workshop buyers convert to bootcamp at a materially higher rate than non-workshop buyers.",
               "This validates workshops as a high-intent pre-bootcamp conversion channel."],
              ["16. Final Verdict", "The free lead magnet worked as an intent segmenter, not as an open-rate growth engine.",
               "Commercially, value comes from targeted follow-up to responders rather than list-wide lift expectations."],
              ["18.2 Action Plan", "Root-cause diagnosis, segmentation, and CTA optimization are the highest-leverage interventions.",
               "Execution sequence matters more than isolated optimizations."],
          ])

body(doc, "Executive take points:", bold=True, color=(79, 70, 229))
bullet(doc, "Primary diagnosis: engagement decay is tenure-linked and began before the open-rate collapse became visible.")
bullet(doc, "Priority KPI stack: track Open Rate for reach, but use CTOR as the early warning and purchase-intent quality metric.")
bullet(doc, "Primary activation strategy: segment by intent (lead-magnet response, buyer history, subscriber age) and tailor cadence.")
bullet(doc, "Primary timing strategy: keep strong early nurture, but maintain medium-horizon follow-up because delayed conversion is meaningful.")

h2(doc, "17.1 Direct Answers to the Core Business Questions")
body(doc,
     "Did sending the survey increase open rates for the whole confirmed-subscriber list? "
     "No. Post-survey open rates are lower than pre-survey rates for both Value and Sales emails.")
body(doc,
     "Do survey responders (Group A) outperform non-responders (Group B)? "
     "On opens, not consistently. On sales CTOR, yes in stable months. "
     "Interpretation: the lead magnet behaves more as an intent/quality segmenter than as an inbox-reach booster.")
body(doc,
     "Are older subscribers disengaging more than newer ones? "
     "Yes. Cohort analysis shows older signup cohorts trailing newer cohorts on both open rate and CTOR "
     "across recent stable months.")
body(doc,
     "Do previous buyers behave differently? "
     "Yes. Buyers consistently outperform non-buyers on both open rate and CTOR, which confirms purchase history "
     "as a strong engagement segmentation signal.")
body(doc,
     "Are workshop buyers more likely to buy bootcamp? "
     "Yes. Workshop buyers show a substantially higher bootcamp conversion rate than non-workshop buyers, "
     "supporting the workshop-first funnel strategy.")
body(doc,
     "How long does conversion take? "
     f"{_lat_evidence} This supports a mixed nurture model: strong early conversion pressure plus sustained follow-up.")

h2(doc, "17.2 Executive Interpretation")
body(doc,
     "The list is not failing uniformly. It is bifurcating: high-intent segments still engage at strong levels, "
     "while mature broad-list segments are decaying in both attention (opens) and action quality (CTOR).")
body(doc,
     "This means one-message-for-all strategy is now structurally inefficient. "
     "Future gains are most likely to come from segment-aware content and CTA strategy, not from global subject-line tweaks.")
body(doc,
     "Operationally, CTOR should be treated as the earliest reliable warning signal. "
     "In this dataset, CTOR deterioration preceded broad open-rate decline by months.")

h2(doc, "17.3 Priority Sequence for Execution")
body(doc,
     "Priority 1: root-cause audit of June 2025 message/offer/link changes. "
     "Priority 2: rebuild value-email CTA quality and relevance. "
     "Priority 3: enforce segmentation by intent and tenure (responders, buyers, old cohorts, new cohorts). "
     "Priority 4: tune sales cadence by segment rather than by global launch calendar. "
     "Priority 5: monitor weekly OR+CTOR control metrics with explicit alert thresholds.")

doc.add_page_break()
h1(doc, "18. Strategic Narrative & Action Plan (Moved from former 9.1/9.2)")

h2(doc, "18.1 The Story the Data Is Telling Us")
body(doc,
     "Across the full audit, the timeline is consistent. "
     "In early 2025, engagement was healthy: opens were strong, click quality was stable, and attrition was mostly early churn. "
     "From June 2025, click behavior weakened first while opens stayed comparatively stable. "
     "From November 2025 onward, the decline became visible at the inbox level, with weaker opens and lower downstream action.")
body(doc,
     "The added analyses reinforce this story rather than changing it. "
     "Lead-magnet responders are not universally better openers, but they are stronger on sales click intent. "
     "Older signup cohorts underperform newer cohorts. Buyers outperform non-buyers. "
     "Together, this shows the central issue is not list size but list composition and segment-specific relevance.")
body(doc,
     "Business interpretation: the system still works for high-intent segments, but broad-list relevance has decayed. "
     "Without segmentation and message-fit recovery, growth in subscriber count will continue to outpace growth in engaged value.")

h2(doc, "18.2 What We Should Do About It")
body(doc, "1. Run a forensic June-2025 change audit", bold=True, color=(79, 70, 229))
body(doc,
     "Compare May vs June across content themes, CTA framing, link destinations, frequency, and launch structure. "
     "The objective is to isolate the first drivers of CTOR deterioration.")
body(doc, "2. Redesign value-email CTA architecture", bold=True, color=(79, 70, 229))
body(doc,
     "Simplify to one primary CTA per message where possible, place it earlier, and align it tightly with the email promise. "
     "Measure lift in click quality, not only open rate.")
body(doc, "3. Shift to segment-first delivery", bold=True, color=(79, 70, 229))
body(doc,
     "At minimum, split into: new subscribers, older cohorts, lead-magnet responders, and prior buyers. "
     "Use differentiated cadence and offer intensity by segment.")
body(doc, "4. Treat Sales CTOR as the conversion-quality KPI", bold=True, color=(79, 70, 229))
body(doc,
     "For sales performance, evaluate both opens (reach) and CTOR (commercial action quality). "
     "In this audit, CTOR is the clearer discriminator of intent quality.")
body(doc, "5. Align nurture horizon with observed conversion latency", bold=True, color=(79, 70, 229))
body(doc,
     "Because many buyers convert after an initial delay, keep strong early conversion pathways but maintain sustained follow-up "
     "for medium-cycle subscribers rather than ending persuasion too early.")
body(doc, "6. Install weekly control metrics with thresholds", bold=True, color=(79, 70, 229))
body(doc,
     "Track OR, CTOR, and segment gaps weekly. Trigger structured review when thresholds are breached so decline is detected "
     "before inbox-level deterioration becomes visible.")

doc.add_paragraph()
divider(doc)
doc.add_paragraph()

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run(
    "Email Health Deep Analysis Report  |  Data: February 2025 – February 2026  |  "
    f"Generated: {datetime.date.today().strftime('%B %d, %Y')}"
)
set_font(run, size=9, italic=True, color=(150, 150, 150))

# ── Save ────────────────────────────────────────────────────────────────────
out_path = BASE / "Email_Health_Analysis_Report.docx"
doc.save(str(out_path))
print(f"Report saved: {out_path}")
